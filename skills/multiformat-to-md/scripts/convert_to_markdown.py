# 脚本说明
# !/usr/bin/env python3
# 技能脚本
import argparse
import datetime as dt
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Callable, Iterable, List, Optional, Tuple

# 脚本说明
SUPPORTED_EXTS = {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".mp3", ".mp4"}

ProgressCallback = Callable[[float, float, str], None]

try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except Exception:
    pass


class ProgressPrinter:
    def __init__(self, total_files: int) -> None:
        self.total_files = max(1, total_files)
        self.current_index = 0
        self._last_percent = -1
        self._last_line_len = 0
        self._use_inline = sys.stdout.isatty()

    def start_file(self, path: Path) -> None:
        self.current_index += 1
        print(f"[{self.current_index}/{self.total_files}] 开始处理: {path.name}")

    def update(self, percent: float, message: str = "") -> None:
        if percent is None:
            return
        clamped = max(0, min(100, int(percent)))
        if clamped == self._last_percent:
            return
        self._last_percent = clamped
        bar = self._render_bar(clamped)
        line = f"  {bar} {clamped:3d}% {message}".rstrip()
        if self._use_inline:
            pad = max(0, self._last_line_len - len(line))
            sys.stdout.write("\r" + line + (" " * pad))
            sys.stdout.flush()
            self._last_line_len = len(line)
        else:
            print(line)

    def finish_file(self, status: str) -> None:
        if self._use_inline and self._last_line_len:
            sys.stdout.write("\r" + (" " * self._last_line_len) + "\r")
            sys.stdout.flush()
        self._last_percent = -1
        self._last_line_len = 0
        print(f"[{self.current_index}/{self.total_files}] 完成: {status}")

    @staticmethod
    def _render_bar(percent: int, width: int = 24) -> str:
        filled = int(width * percent / 100)
        return "[" + ("=" * filled) + ("-" * (width - filled)) + "]"


def _make_progress_callback(printer: ProgressPrinter, label: str) -> ProgressCallback:
    def _callback(current: float, total: float, message: str = "") -> None:
        if total <= 0:
            return
        percent = (current / total) * 100
        msg = message or label
        printer.update(percent, msg)

    return _callback


def _get_media_duration_seconds(path: Path) -> Optional[float]:
    ffprobe = "ffprobe"
    cmd = [
        ffprobe,
        "-v",
        "error",
        "-show_entries",
        "format=duration",
        "-of",
        "default=noprint_wrappers=1:nokey=1",
        str(path),
    ]
    try:
        result = subprocess.run(
            cmd,
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
    except (FileNotFoundError, subprocess.CalledProcessError):
        return None
    raw = (result.stdout or "").strip()
    try:
        duration = float(raw)
        return duration if duration > 0 else None
    except Exception:
        return None

# 解析输入
def resolve_inputs(inputs: Iterable[str]) -> List[Path]:
    # 文件处理
    files: List[Path] = []
    for item in inputs:
        p = Path(item).expanduser().resolve()
        if p.is_dir():
            for child in p.rglob("*"):
                if child.is_file() and child.suffix.lower() in SUPPORTED_EXTS:
                    files.append(child)
        elif p.is_file() and p.suffix.lower() in SUPPORTED_EXTS:
            files.append(p)
    unique_sorted = sorted(set(files))
    return unique_sorted


def _meaningful_char_count(text: str) -> int:
    if not text:
        return 0
    # 统计中英文与数字，过滤纯空白/符号
    compact = re.sub(r"\s+", "", text)
    cleaned = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]", "", compact)
    return len(cleaned)


def _cjk_char_count(text: str) -> int:
    if not text:
        return 0
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def _is_likely_garbled_text(text: str) -> bool:
    if not text:
        return False
    compact = re.sub(r"\s+", "", text)
    if len(compact) < 80:
        return False
    if compact.count("\ufffd") >= 3:
        return True

    alpha_tokens = re.findall(r"[A-Za-z]+", text)
    if not alpha_tokens:
        return False
    alpha_chars = sum(len(token) for token in alpha_tokens)
    if alpha_chars < 60:
        return False

    upper_chars = len(re.findall(r"[A-Z]", text))
    short_token_ratio = (
        sum(1 for token in alpha_tokens if len(token) <= 2) / max(len(alpha_tokens), 1)
    )
    upper_ratio = upper_chars / max(alpha_chars, 1)
    cjk_chars = _cjk_char_count(text)
    cjk_ratio = cjk_chars / max(len(compact), 1)

    if cjk_ratio >= 0.05:
        return False
    if upper_ratio >= 0.75 and alpha_chars >= 120:
        return True
    return upper_ratio >= 0.58 and short_token_ratio >= 0.50


def _normalize_tesseract_lang(lang: str) -> str:
    parts = [item.strip() for item in (lang or "").split("+") if item.strip()]
    return "+".join(parts)


def _get_tesseract_lang_setting() -> str:
    raw = (
        os.getenv("OCR_LANG")
        or os.getenv("TESSERACT_LANG")
        or "chi_sim+eng"
    )
    normalized = _normalize_tesseract_lang(raw)
    return normalized or "chi_sim+eng"


def _get_tesseract_config_setting() -> str:
    config = (os.getenv("TESSERACT_CONFIG") or "").strip()
    return config or "--oem 3 --psm 6"


def _get_tesseract_lang_candidates(pytesseract_module: object) -> List[str]:
    preferred = _get_tesseract_lang_setting()
    candidates: List[str] = []
    for lang in [preferred, "chi_sim+eng", "chi_sim", "eng"]:
        normalized = _normalize_tesseract_lang(lang)
        if normalized and normalized not in candidates:
            candidates.append(normalized)

    available_langs: Optional[set] = None
    try:
        getter = getattr(pytesseract_module, "get_languages", None)
        if callable(getter):
            available_langs = set(getter(config=""))
    except Exception:
        available_langs = None

    if not available_langs:
        return candidates

    filtered: List[str] = []
    for lang in candidates:
        parts = [item for item in lang.split("+") if item]
        if all(part in available_langs for part in parts):
            filtered.append(lang)

    if filtered:
        return filtered
    if "eng" in available_langs:
        return ["eng"]
    return candidates


def _score_ocr_text(text: str) -> int:
    return _meaningful_char_count(text) + (_cjk_char_count(text) * 2)


def _ocr_image_to_text(pytesseract_module: object, image: object) -> str:
    config = _get_tesseract_config_setting()
    candidates = _get_tesseract_lang_candidates(pytesseract_module)
    best_text = ""
    best_score = -1
    errors: List[str] = []

    for lang in candidates:
        try:
            text = (
                getattr(pytesseract_module, "image_to_string")(
                    image, lang=lang, config=config
                )
                or ""
            ).strip()
        except Exception as exc:
            errors.append(f"{lang}: {exc}")
            continue
        score = _score_ocr_text(text)
        if score > best_score:
            best_text = text
            best_score = score
        if _cjk_char_count(text) >= 8 and _meaningful_char_count(text) >= 20:
            return text

    if best_text:
        return best_text

    try:
        text = (getattr(pytesseract_module, "image_to_string")(image, config=config) or "").strip()
        if text:
            return text
    except Exception as exc:
        errors.append(f"default: {exc}")

    lang_hint = _get_tesseract_lang_setting()
    hint = f"当前 OCR 语言配置: {lang_hint}。"
    if "chi_sim" in lang_hint:
        hint += " 请确认已安装 tesseract 中文语言包 chi_sim。"
    detail = " | ".join(errors[:3])
    raise RuntimeError(
        "OCR 失败。请确认已安装 Tesseract OCR 并在 PATH 中可用。"
        + (" " + hint if hint else "")
        + (f" 细节: {detail}" if detail else "")
    )


def _extract_pdf_by_fitz_text(
    path: Path,
    total_pages: int,
    progress_callback: Optional[ProgressCallback] = None,
) -> Tuple[List[str], int, int, int]:
    try:
        import fitz  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("缺少依赖: pymupdf（fitz）") from exc

    doc = fitz.open(str(path))
    try:
        fitz_pages = len(doc)
        pages = fitz_pages if fitz_pages > 0 else total_pages
        chunks: List[str] = []
        non_empty_pages = 0
        meaningful_total = 0
        garbled_pages = 0
        for idx in range(fitz_pages):
            page = doc.load_page(idx)
            text = (page.get_text("text") or "").strip()
            page_meaningful = _meaningful_char_count(text)
            if page_meaningful >= 10:
                non_empty_pages += 1
            if text:
                meaningful_total += page_meaningful
            if _is_likely_garbled_text(text):
                garbled_pages += 1
            chunks.append(f"## 第 {idx + 1} 页\n\n{text}\n")
            if progress_callback:
                progress_callback(
                    total_pages + idx + 1,
                    max(total_pages * 2, 1),
                    f"PyMuPDF 提取第 {idx + 1}/{pages} 页",
                )
        return chunks, non_empty_pages, meaningful_total, garbled_pages
    finally:
        doc.close()


def _pdf_text_coverage_ok(total_pages: int, non_empty_pages: int) -> bool:
    pages = max(total_pages, 1)
    ratio = non_empty_pages / pages
    if pages <= 3:
        return non_empty_pages >= 1
    if pages <= 8:
        return non_empty_pages >= 2 and ratio >= 0.25
    return non_empty_pages >= 3 and ratio >= 0.30


def _probe_pdf_page_count(path: Path) -> int:
    try:
        import fitz  # type: ignore
    except Exception:
        return 1
    try:
        doc = fitz.open(str(path))
        count = len(doc)
        doc.close()
        return max(count, 1)
    except Exception:
        return 1


def _recover_pdf_text_with_fallback(
    path: Path,
    total_pages: int,
    reason: str,
    progress_callback: Optional[ProgressCallback] = None,
) -> str:
    fitz_error: Optional[Exception] = None
    if progress_callback:
        progress_callback(
            total_pages,
            max(total_pages * 2, 1),
            f"{reason}，尝试 PyMuPDF 文本提取",
        )
    try:
        fitz_chunks, fitz_non_empty_pages, fitz_meaningful_total, fitz_garbled_pages = _extract_pdf_by_fitz_text(
            path=path,
            total_pages=total_pages,
            progress_callback=progress_callback,
        )
        fitz_garbled_ratio = (fitz_garbled_pages / total_pages) if total_pages > 0 else 1.0
        if (
            fitz_meaningful_total >= 20
            and fitz_garbled_ratio < 0.40
            and _pdf_text_coverage_ok(total_pages, fitz_non_empty_pages)
        ):
            return "\n".join(fitz_chunks).strip()
    except Exception as exc:
        fitz_error = exc

    if progress_callback:
        progress_callback(
            total_pages,
            max(total_pages * 2, 1),
            f"{reason}，启动 OCR 回退",
        )
    try:
        ocr_chunks, ocr_non_empty_pages, ocr_meaningful_total = _extract_pdf_by_ocr(
            path=path,
            total_pages=total_pages,
            progress_callback=progress_callback,
        )
    except Exception as ocr_exc:
        if fitz_error is not None:
            raise RuntimeError(
                f"PDF 回退失败（PyMuPDF 与 OCR 均失败）。PyMuPDF: {fitz_error}; OCR: {ocr_exc}"
            ) from ocr_exc
        raise

    if ocr_non_empty_pages == 0 or ocr_meaningful_total < 20:
        raise RuntimeError(
            "PDF 文本提取与 OCR 回退后仍为空或信息量极低，"
            "请检查文件清晰度或 OCR 语言包配置。"
        )
    return "\n".join(ocr_chunks).strip()


def _extract_pdf_by_ocr(
    path: Path,
    total_pages: int,
    progress_callback: Optional[ProgressCallback] = None,
) -> Tuple[List[str], int, int]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("缺少依赖: pillow 和/或 pytesseract（PDF OCR 回退需要）") from exc

    def _run_ocr_on_pages_with_fitz() -> Tuple[List[str], int, int]:
        try:
            import fitz  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise RuntimeError("缺少依赖: pymupdf（fitz）") from exc

        doc = fitz.open(str(path))
        try:
            fitz_pages = len(doc)
            pages = fitz_pages if fitz_pages > 0 else total_pages
            chunks: List[str] = []
            non_empty_pages = 0
            meaningful_total = 0
            for idx in range(fitz_pages):
                page = doc.load_page(idx)
                # 2x 缩放提升 OCR 识别率
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                mode = "RGB" if pix.n >= 3 else "L"
                image = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
                text = _ocr_image_to_text(pytesseract, image)
                image.close()

                if text:
                    non_empty_pages += 1
                    meaningful_total += _meaningful_char_count(text)
                chunks.append(f"## 第 {idx + 1} 页\n\n{text}\n")
                if progress_callback:
                    progress_callback(
                        total_pages + idx + 1,
                        max(total_pages * 2, 1),
                        f"OCR PDF 第 {idx + 1}/{pages} 页",
                    )
            return chunks, non_empty_pages, meaningful_total
        finally:
            doc.close()

    def _run_ocr_on_pages_with_pdfium() -> Tuple[List[str], int, int]:
        try:
            import pypdfium2 as pdfium  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise RuntimeError("缺少依赖: pypdfium2") from exc

        pdf = pdfium.PdfDocument(str(path))
        try:
            pdfium_pages = len(pdf)
            pages = pdfium_pages if pdfium_pages > 0 else total_pages
            chunks: List[str] = []
            non_empty_pages = 0
            meaningful_total = 0
            for idx in range(pdfium_pages):
                page = pdf[idx]
                bitmap = page.render(scale=2.0)
                image = bitmap.to_pil()
                text = _ocr_image_to_text(pytesseract, image)
                image.close()
                try:
                    bitmap.close()
                except Exception:
                    pass
                try:
                    page.close()
                except Exception:
                    pass

                if text:
                    non_empty_pages += 1
                    meaningful_total += _meaningful_char_count(text)
                chunks.append(f"## 第 {idx + 1} 页\n\n{text}\n")
                if progress_callback:
                    progress_callback(
                        total_pages + idx + 1,
                        max(total_pages * 2, 1),
                        f"OCR PDF 第 {idx + 1}/{pages} 页",
                    )
            return chunks, non_empty_pages, meaningful_total
        finally:
            try:
                pdf.close()
            except Exception:
                pass

    fitz_error: Optional[Exception] = None
    try:
        return _run_ocr_on_pages_with_fitz()
    except Exception as exc:
        fitz_error = exc

    pdfium_error: Optional[Exception] = None
    try:
        return _run_ocr_on_pages_with_pdfium()
    except Exception as exc:
        pdfium_error = exc

    detail = []
    if fitz_error is not None:
        detail.append(f"fitz: {fitz_error}")
    if pdfium_error is not None:
        detail.append(f"pypdfium2: {pdfium_error}")
    raise RuntimeError(
        "PDF OCR 回退失败。请安装 Tesseract 并确保在 PATH 中可用；"
        "同时安装 pymupdf（fitz）或 pypdfium2 以渲染 PDF 页面。"
        + (f" 细节: {' | '.join(detail)}" if detail else "")
    )


# 提取 PDF
def extract_pdf(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    # 尝试导入
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return _recover_pdf_text_with_fallback(
            path=path,
            total_pages=_probe_pdf_page_count(path),
            reason="缺少 pypdf，使用回退链路",
            progress_callback=progress_callback,
        )

    reader = PdfReader(str(path))
    chunks: List[str] = []
    non_empty_pages = 0
    meaningful_total = 0
    low_quality_pages = 0
    garbled_pages = 0
    page_meaningful_min = 10
    total_pages = len(reader.pages)
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_text = text.strip()
        page_meaningful = _meaningful_char_count(page_text)
        if page_meaningful >= page_meaningful_min:
            non_empty_pages += 1
        else:
            low_quality_pages += 1
        if page_text:
            meaningful_total += page_meaningful
        if _is_likely_garbled_text(page_text):
            garbled_pages += 1
        chunks.append(f"## 第 {idx} 页\n\n{page_text}\n")
        if progress_callback:
            progress_callback(idx, total_pages, f"解析 PDF 第 {idx}/{total_pages} 页")

    low_quality_ratio = (low_quality_pages / total_pages) if total_pages > 0 else 1.0
    garbled_ratio = (garbled_pages / total_pages) if total_pages > 0 else 1.0
    # 规则：若出现 60% 页无法通过 pypdf 获取有效文本，则整份文档转全量 OCR
    if low_quality_ratio >= 0.60:
        return _recover_pdf_text_with_fallback(
            path=path,
            total_pages=total_pages,
            reason=f"检测到 {low_quality_pages}/{total_pages} 页文本层不足",
            progress_callback=progress_callback,
        )

    # 规则：若文本层疑似乱码（常见于中文 PDF 被错误映射），切换回退链路
    if garbled_ratio >= 0.40 and garbled_pages >= 1:
        return _recover_pdf_text_with_fallback(
            path=path,
            total_pages=total_pages,
            reason=f"检测到 {garbled_pages}/{total_pages} 页文本层疑似乱码",
            progress_callback=progress_callback,
        )

    if non_empty_pages == 0 or meaningful_total < 20:
        return _recover_pdf_text_with_fallback(
            path=path,
            total_pages=total_pages,
            reason="文本层信息不足",
            progress_callback=progress_callback,
        )

    return "\n".join(chunks).strip()

# 提取 DOCX
def extract_docx(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("缺少依赖: python-docx") from exc

    d = docx.Document(str(path))
    if progress_callback:
        progress_callback(0.2, 1.0, "解析 Word 文档")
    # 行列表
    lines: List[str] = []

    for p in d.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    for t_idx, table in enumerate(d.tables, start=1):
        lines.append(f"\n## 表格 {t_idx}")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
    if progress_callback:
        progress_callback(1.0, 1.0, "Word 解析完成")
    return "\n".join(lines).strip()

# 提取 DOC
def extract_doc(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    antiword = "antiword"
    try:
        result = subprocess.run(
            [antiword, str(path)],
            check=True,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
    except FileNotFoundError as exc:
        raise RuntimeError("缺少系统工具: antiword（用于 .doc）") from exc
    except subprocess.CalledProcessError as exc:
        stderr = (exc.stderr or "").strip()
        raise RuntimeError(f"antiword 失败: {stderr}") from exc
    if progress_callback:
        progress_callback(1.0, 1.0, "DOC 解析完成")
    return result.stdout.strip()

# 提取图片
def extract_image(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("缺少依赖: pillow 和/或 pytesseract") from exc

    try:
        if progress_callback:
            progress_callback(0.1, 1.0, "OCR 识别中")
        with Image.open(path) as image:
            text = _ocr_image_to_text(pytesseract, image)
    except Exception as exc:
        raise RuntimeError(
            "OCR 失败。请确认已安装 Tesseract OCR 并在 PATH 中可用。"
        ) from exc
    if progress_callback:
        progress_callback(1.0, 1.0, "OCR 完成")
    return text.strip()

# 转写音频
def _faster_whisper_available() -> bool:
    try:
        import faster_whisper  # type: ignore  # noqa: F401
        return True
    except Exception:
        return False


def _select_whisper_engine() -> str:
    engine = (os.getenv("WHISPER_ENGINE") or "").strip().lower()
    if engine in {"faster", "faster-whisper"}:
        return "faster"
    if engine in {"openai", "openai-whisper", "whisper"}:
        return "openai"
    return "faster" if _faster_whisper_available() else "openai"


def _transcribe_with_faster_whisper(
    path: Path,
    model_name: str = "base",
    progress_callback: Optional[ProgressCallback] = None,
    duration_seconds: Optional[float] = None,
) -> str:
    try:
        from faster_whisper import WhisperModel  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("缺少依赖: faster-whisper") from exc

    device = "cuda" if _torch_cuda_available() else "cpu"
    compute_type = "float16" if device == "cuda" else "int8"
    model = WhisperModel(model_name, device=device, compute_type=compute_type)
    segments, _info = model.transcribe(str(path))
    texts = []
    last_report = -1
    for seg in segments:
        if seg.text:
            texts.append(seg.text)
        if progress_callback and duration_seconds:
            current = min(duration_seconds, float(getattr(seg, "end", 0.0)))
            percent = int((current / duration_seconds) * 100)
            if percent != last_report:
                last_report = percent
                progress_callback(current, duration_seconds, "转写中")
    return "".join(texts).strip()


def _torch_cuda_available() -> bool:
    try:
        import torch  # type: ignore
        return bool(torch.cuda.is_available())
    except Exception:
        return False


def _transcribe_with_openai_whisper(
    path: Path,
    model_name: str = "base",
    progress_callback: Optional[ProgressCallback] = None,
) -> str:
    try:
        import whisper  # type: ignore
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("缺少依赖: openai-whisper") from exc

    if progress_callback:
        progress_callback(0.1, 1.0, "转写中")
    model = whisper.load_model(model_name)
    result = model.transcribe(str(path), fp16=_torch_cuda_available())
    if progress_callback:
        progress_callback(1.0, 1.0, "转写完成")
    return (result.get("text") or "").strip()


def transcribe_audio(
    path: Path,
    model_name: str = "base",
    progress_callback: Optional[ProgressCallback] = None,
    duration_seconds: Optional[float] = None,
) -> str:
    engine = _select_whisper_engine()
    if engine == "faster":
        return _transcribe_with_faster_whisper(
            path,
            model_name=model_name,
            progress_callback=progress_callback,
            duration_seconds=duration_seconds,
        )
    return _transcribe_with_openai_whisper(
        path,
        model_name=model_name,
        progress_callback=progress_callback,
    )

# 转写视频
def transcribe_video(
    path: Path,
    model_name: str = "base",
    progress_callback: Optional[ProgressCallback] = None,
) -> str:
    ffmpeg = "ffmpeg"
    duration = _get_media_duration_seconds(path)
    with tempfile.TemporaryDirectory() as tmp:
        audio_path = Path(tmp) / "audio.wav"
        cmd = [
            ffmpeg,
            "-y",
            "-i",
            str(path),
            "-vn",
            "-ac",
            "1",
            "-ar",
            "16000",
            str(audio_path),
        ]
        try:
            if progress_callback:
                progress_callback(0.01, 1.0, "提取音频中")
            subprocess.run(cmd, check=True, capture_output=True, text=True)
        except FileNotFoundError as exc:
            raise RuntimeError("缺少系统工具: ffmpeg（用于 .mp4）") from exc
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or "").strip()
            raise RuntimeError(f"ffmpeg 失败: {stderr}") from exc
        return transcribe_audio(
            audio_path,
            model_name=model_name,
            progress_callback=progress_callback,
            duration_seconds=duration,
        )

# 转换文件
def convert_file(
    path: Path,
    whisper_model: str,
    progress_callback: Optional[ProgressCallback] = None,
) -> str:
    ext = path.suffix.lower()
    if progress_callback:
        progress_callback(0.0, 1.0, f"处理中: {path.name}")
    if ext == ".pdf":
        return extract_pdf(path, progress_callback=progress_callback)
    if ext == ".docx":
        return extract_docx(path, progress_callback=progress_callback)
    if ext == ".doc":
        return extract_doc(path, progress_callback=progress_callback)
    if ext in {".png", ".jpg", ".jpeg"}:
        return extract_image(path, progress_callback=progress_callback)
    if ext == ".mp3":
        duration = _get_media_duration_seconds(path)
        return transcribe_audio(
            path,
            model_name=whisper_model,
            progress_callback=progress_callback,
            duration_seconds=duration,
        )
    if ext == ".mp4":
        return transcribe_video(
            path,
            model_name=whisper_model,
            progress_callback=progress_callback,
        )
    raise RuntimeError(f"不支持的文件扩展名: {ext}")

# 构建 Markdown
def build_markdown(path: Path, content: str, status: str, error: str = "") -> str:
    ts = dt.datetime.now(dt.timezone.utc).isoformat()
    header = [
        "---",
        f'source_file: "{path.as_posix()}"',
        f"converted_at_utc: {ts}",
        f"status: {status}",
        "---",
        "",
        f"# {path.name}",
        "",
    ]
    if error:
        header.extend(["## Error", "", error.strip(), ""])
    header.extend(["## Content", "", content.strip(), ""])
    return "\n".join(header)

# 写入输出
def write_output(input_file: Path, output_dir: Path, markdown: str) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    target = output_dir / f"{input_file.stem}.md"
    target.write_text(markdown, encoding="utf-8")
    return target

# 主函数
def main() -> int:
    parser = argparse.ArgumentParser(
        description="将 PDF/Word/图片/音视频转换为适合 LLM 摄取的 Markdown。"
    )
    parser.add_argument(
        "--input",
        nargs="+",
        required=True,
        help="输入文件和/或目录。",
    )
    parser.add_argument("--output", required=True, help="Markdown 输出目录。")
    parser.add_argument(
        "--whisper-model",
        default="base",
        help="用于 mp3/mp4 的 Whisper 模型名（如 tiny、base、small）。",
    )
    args = parser.parse_args()

    files = resolve_inputs(args.input)
    if not files:
        print("未在提供的输入中找到支持的文件。", file=sys.stderr)
        return 1

    output_dir = Path(args.output).expanduser().resolve()
    failed = 0
    printer = ProgressPrinter(total_files=len(files))
    for f in files:
        printer.start_file(f)
        progress_cb = _make_progress_callback(printer, f"处理中: {f.name}")
        try:
            content = convert_file(
                f,
                whisper_model=args.whisper_model,
                progress_callback=progress_cb,
            )
            md = build_markdown(f, content=content, status="ok")
        except Exception as exc:
            failed += 1
            md = build_markdown(
                f,
                content="",
                status="failed",
                error=str(exc),
            )
        printer.update(100, "写入 Markdown")
        out = write_output(f, output_dir, md)
        status = "ok" if "status: ok" in md[:1200] else "failed"
        printer.finish_file(status)
        print(f"[已写入] {out}")

    print(f"已处理 {len(files)} 个文件；失败: {failed}")
    return 0 if failed == 0 else 2

if __name__ == "__main__":
    raise SystemExit(main())
