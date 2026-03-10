#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path
from typing import List, Optional, Tuple

from multiformat_gateway import (
    DOC_IMAGE_EXTENSIONS,
    ProgressCallback,
    ProgressPrinter,
    _make_progress_callback,
    build_markdown,
    resolve_inputs,
    write_output,
)


def _meaningful_char_count(text: str) -> int:
    if not text:
        return 0
    compact = re.sub(r"\s+", "", text)
    cleaned = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]", "", compact)
    return len(cleaned)


def _cjk_char_count(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text or ""))


def _is_likely_garbled_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if len(compact) < 80:
        return False
    if compact.count("\ufffd") >= 3:
        return True

    alpha_tokens = re.findall(r"[A-Za-z]+", text or "")
    if not alpha_tokens:
        return False
    alpha_chars = sum(len(token) for token in alpha_tokens)
    if alpha_chars < 60:
        return False

    upper_chars = len(re.findall(r"[A-Z]", text or ""))
    short_ratio = sum(1 for token in alpha_tokens if len(token) <= 2) / max(len(alpha_tokens), 1)
    upper_ratio = upper_chars / max(alpha_chars, 1)
    cjk_ratio = _cjk_char_count(text) / max(len(compact), 1)
    if cjk_ratio >= 0.05:
        return False
    if upper_ratio >= 0.75 and alpha_chars >= 120:
        return True
    return upper_ratio >= 0.58 and short_ratio >= 0.50


def _normalize_tesseract_lang(lang: str) -> str:
    return "+".join([x.strip() for x in (lang or "").split("+") if x.strip()])


def _get_tesseract_lang_setting() -> str:
    raw = os.getenv("OCR_LANG") or os.getenv("TESSERACT_LANG") or "chi_sim+eng"
    normalized = _normalize_tesseract_lang(raw)
    return normalized or "chi_sim+eng"


def _get_tesseract_config_setting() -> str:
    return (os.getenv("TESSERACT_CONFIG") or "--oem 3 --psm 6").strip()


def _get_tesseract_lang_candidates(pytesseract_module: object) -> List[str]:
    preferred = _get_tesseract_lang_setting()
    candidates: List[str] = []
    for lang in [preferred, "chi_sim+eng", "chi_sim", "eng"]:
        n = _normalize_tesseract_lang(lang)
        if n and n not in candidates:
            candidates.append(n)

    available_langs: Optional[set] = None
    try:
        getter = getattr(pytesseract_module, "get_languages", None)
        if callable(getter):
            available_langs = set(getter(config=""))
    except Exception:
        available_langs = None

    if not available_langs:
        return candidates

    filtered: List[str] = []
    for lang in candidates:
        parts = [x for x in lang.split("+") if x]
        if all(part in available_langs for part in parts):
            filtered.append(lang)

    if filtered:
        return filtered
    if "eng" in available_langs:
        return ["eng"]
    return candidates


def _score_ocr_text(text: str) -> int:
    return _meaningful_char_count(text) + (_cjk_char_count(text) * 2)


def _ocr_image_to_text(pytesseract_module: object, image: object) -> str:
    config = _get_tesseract_config_setting()
    candidates = _get_tesseract_lang_candidates(pytesseract_module)
    best_text = ""
    best_score = -1

    for lang in candidates:
        try:
            text = (getattr(pytesseract_module, "image_to_string")(image, lang=lang, config=config) or "").strip()
        except Exception:
            continue
        score = _score_ocr_text(text)
        if score > best_score:
            best_text = text
            best_score = score
        if _cjk_char_count(text) >= 8 and _meaningful_char_count(text) >= 20:
            return text

    if best_text:
        return best_text

    text = (getattr(pytesseract_module, "image_to_string")(image, config=config) or "").strip()
    return text


def _extract_pdf_by_fitz_text(path: Path, progress_callback: Optional[ProgressCallback] = None) -> Tuple[List[str], int, int, int]:
    try:
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("Missing dependency: pymupdf (fitz)") from exc

    chunks: List[str] = []
    non_empty = 0
    meaningful_total = 0
    garbled_pages = 0

    with fitz.open(str(path)) as doc:
        total_pages = len(doc)
        for idx in range(total_pages):
            page = doc[idx]
            text = (page.get_text("text") or "").strip()
            if text:
                non_empty += 1
                meaningful_total += _meaningful_char_count(text)
                if _is_likely_garbled_text(text):
                    garbled_pages += 1
            chunks.append(text)
            if progress_callback:
                progress_callback(idx + 1, total_pages, "Extract PDF text")

    return chunks, non_empty, meaningful_total, garbled_pages


def _pdf_text_coverage_ok(total_pages: int, non_empty_pages: int) -> bool:
    if total_pages <= 0:
        return False
    ratio = non_empty_pages / total_pages
    if total_pages <= 3:
        return non_empty_pages >= 1
    if total_pages <= 8:
        return non_empty_pages >= 2 and ratio >= 0.25
    return non_empty_pages >= 3 and ratio >= 0.30


def _probe_pdf_page_count(path: Path) -> int:
    try:
        from pypdf import PdfReader  # type: ignore
        return len(PdfReader(str(path)).pages)
    except Exception:
        return 0


def _extract_pdf_by_ocr(path: Path, total_pages: int, progress_callback: Optional[ProgressCallback] = None) -> Tuple[List[str], int, int]:
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
        import fitz  # type: ignore
    except Exception as exc:
        raise RuntimeError("Missing OCR dependency: pillow/pytesseract/pymupdf") from exc

    chunks: List[str] = []
    non_empty = 0
    meaningful_total = 0

    with fitz.open(str(path)) as doc:
        pages = len(doc)
        total = pages if pages > 0 else total_pages
        for idx in range(pages):
            page = doc[idx]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = _ocr_image_to_text(pytesseract, image)
            if text:
                non_empty += 1
                meaningful_total += _meaningful_char_count(text)
            chunks.append(text.strip())
            if progress_callback:
                progress_callback(idx + 1, max(total, 1), "OCR PDF page")

    return chunks, non_empty, meaningful_total


def _recover_pdf_text_with_fallback(path: Path, total_pages: int, reason: str, progress_callback: Optional[ProgressCallback] = None) -> str:
    try:
        chunks, non_empty, meaningful_total = _extract_pdf_by_ocr(path, total_pages=total_pages, progress_callback=progress_callback)
        if chunks and meaningful_total > 0 and _pdf_text_coverage_ok(max(total_pages, len(chunks)), non_empty):
            pages = []
            for idx, text in enumerate(chunks, start=1):
                pages.append(f"## 第 {idx} 页\n\n{text.strip()}")
            return "\n\n".join(pages).strip()
        raise RuntimeError("OCR fallback produced low-quality content")
    except Exception as exc:
        raise RuntimeError(f"{reason}; fallback failed: {exc}") from exc


def extract_pdf(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    total_pages = _probe_pdf_page_count(path)
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return _recover_pdf_text_with_fallback(path, total_pages, "Missing pypdf", progress_callback)

    reader = PdfReader(str(path))
    pages = len(reader.pages)
    texts: List[str] = []
    low_quality_pages = 0

    for idx, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        texts.append(text)
        if not text or _meaningful_char_count(text) < 10 or _is_likely_garbled_text(text):
            low_quality_pages += 1
        if progress_callback:
            progress_callback(idx, max(pages, 1), "Extract PDF text")

    if pages > 0 and (low_quality_pages / pages) >= 0.60:
        return _recover_pdf_text_with_fallback(path, pages, "pypdf extraction quality is low", progress_callback)

    parts = [f"## 第 {idx} 页\n\n{txt}" for idx, txt in enumerate(texts, start=1)]
    out = "\n\n".join(parts).strip()
    if _meaningful_char_count(out) < 20:
        return _recover_pdf_text_with_fallback(path, pages, "pypdf extraction is almost empty", progress_callback)
    return out


def extract_docx(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise RuntimeError("Missing dependency: python-docx") from exc

    d = docx.Document(str(path))
    lines: List[str] = []
    total = len(d.paragraphs) + len(d.tables)
    step = 0

    for p in d.paragraphs:
        step += 1
        t = (p.text or "").strip()
        if t:
            lines.append(t)
        if progress_callback:
            progress_callback(step, max(total, 1), "Extract DOCX paragraph")

    for table in d.tables:
        step += 1
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
        if progress_callback:
            progress_callback(step, max(total, 1), "Extract DOCX table")

    return "\n".join(lines).strip()


def extract_doc(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    antiword = "antiword"
    cmd = [antiword, str(path)]
    try:
        result = subprocess.run(cmd, check=True, capture_output=True, text=True, encoding="utf-8", errors="ignore")
        if progress_callback:
            progress_callback(1, 1, "Extract DOC")
        return (result.stdout or "").strip()
    except FileNotFoundError as exc:
        raise RuntimeError("Missing system tool: antiword (for .doc)") from exc
    except subprocess.CalledProcessError as exc:
        stderr = (exc.stderr or "").strip()
        raise RuntimeError(f"antiword failed: {stderr}") from exc


def extract_image(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except Exception as exc:
        raise RuntimeError("Missing dependency: pillow and/or pytesseract") from exc

    if progress_callback:
        progress_callback(0.1, 1.0, "OCR in progress")
    with Image.open(path) as image:
        text = _ocr_image_to_text(pytesseract, image)
    if progress_callback:
        progress_callback(1.0, 1.0, "OCR done")
    return text.strip()


def convert_doc_or_image(path: Path, progress_callback: Optional[ProgressCallback] = None) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path, progress_callback=progress_callback)
    if ext == ".docx":
        return extract_docx(path, progress_callback=progress_callback)
    if ext == ".doc":
        return extract_doc(path, progress_callback=progress_callback)
    if ext in {".png", ".jpg", ".jpeg"}:
        return extract_image(path, progress_callback=progress_callback)
    raise RuntimeError(f"Unsupported extension for doc/image converter: {ext}")


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert PDF/Word/images to Markdown for LLM ingestion.")
    parser.add_argument("--input", nargs="+", required=True, help="Input files and/or directories.")
    parser.add_argument("--output", required=True, help="Output markdown directory.")
    args = parser.parse_args()

    files = resolve_inputs(args.input, allowed_exts=DOC_IMAGE_EXTENSIONS)
    if not files:
        print("No supported PDF/Word/image files found.", file=sys.stderr)
        return 1

    output_dir = Path(args.output).expanduser().resolve()
    failed = 0
    printer = ProgressPrinter(total_files=len(files))

    for f in files:
        printer.start_file(f)
        progress_cb = _make_progress_callback(printer, f"Processing: {f.name}")
        try:
            content = convert_doc_or_image(f, progress_callback=progress_cb)
            md = build_markdown(f, content=content, status="ok")
        except Exception as exc:
            failed += 1
            md = build_markdown(f, content="", status="failed", error=str(exc))

        printer.update(100, "Writing markdown")
        out = write_output(f, output_dir, md)
        status = "ok" if "status: ok" in md[:1200] else "failed"
        printer.finish_file(status)
        print(f"[written] {out}")

    print(f"Processed {len(files)} files; failed: {failed}")
    return 0 if failed == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())
