#!/usr/bin/env python
# -*- coding: utf-8 -*-

import json
import os
import re
import shutil
import threading
import time
import uuid
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from flask import Blueprint, jsonify, request, send_file, Response, stream_with_context
from werkzeug.utils import secure_filename

from utils.agents import (
    SUPPORTED_EXTENSIONS,
    chat_with_agent,
    chat_with_agent_stream,
    convert_file_to_markdown_via_skill,
    crawl_wechat_articles_from_text,
    list_agent_profiles,
    summarize_markdown_with_llm,
    set_runtime_event_hook,
)


strategy_watch_bp = Blueprint("strategy_watch", __name__)

BASE_DIR = Path(__file__).resolve().parent.parent
STORE_DIR = BASE_DIR / "data_cache" / "strategy_watch"
REFERENCE_DIR = STORE_DIR / "reference"
MEMORY_DIR = STORE_DIR / "memory"
UPLOAD_DIR = REFERENCE_DIR / "uploads"
MARKDOWN_DIR = REFERENCE_DIR / "markdown"
CRAWLED_DIR = REFERENCE_DIR / "crawled"
RESOURCES_FILE = REFERENCE_DIR / "resources.json"
CONVERSATIONS_FILE = MEMORY_DIR / "conversations.json"
STRATEGY_ROOT_DIR = STORE_DIR / "strategy"
STRATEGY_DIR = STRATEGY_ROOT_DIR / "strategies"
STRATEGIES_FILE = STRATEGY_ROOT_DIR / "strategies.json"
STRATEGY_INDEX_FILE = STRATEGY_ROOT_DIR / "strategies_index.json"
MEMORY_PROFILES_FILE = MEMORY_DIR / "memory_profiles.json"
MEMORY_LINKS_FILE = MEMORY_DIR / "memory_links.json"
MEMORY_PORTRAITS_FILE = MEMORY_DIR / "memory_portraits.json"
AGENT_LOGS_DIR = STORE_DIR / "logs"
LEGACY_AGENT_LOGS_FILE = STORE_DIR / "agent_logs.jsonl"

MAX_CONTEXT_CHARS = 120_000
MAX_RESOURCE_CHARS = 30_000
MAX_HISTORY_MESSAGES = 24
MAX_MEMORY_CONTEXT_CHARS = 40_000
DEFAULT_GROUP_NAME = "未分组"
DEFAULT_STRATEGY_VIEW = "basic"
DEFAULT_MODEL_OPTIONS = [
    {"label": "DeepSeek v3.2", "value": "deepseek-v3.2"},
    {"label": "DeepSeek v3.2 Thinking", "value": "deepseek-v3.2-thinking"},
    {"label": "qwen-max", "value": "qwen-max"},
]
ALLOWED_RUNTIME_MODEL_VALUES = {item["value"] for item in DEFAULT_MODEL_OPTIONS}
DEFAULT_CHAT_AGENT_NAME = "dialog_agent"
LEGACY_CRAWLER_AGENT_NAME = "crawler_agent"
ALLOWED_CONVERSATION_MODES = {"dialog", "crawler", "strategy_edit", "strategy_analysis"}
CONVERSATION_MODE_AGENT_MAP = {
    "dialog": "dialog_agent",
    "crawler": "engineer_agent",
    "strategy_edit": "architect_agent",
    "strategy_analysis": "analyst_agent",
}

_DATE_YMD_SEP_PATTERN = re.compile(r"(?<!\d)(20\d{2})[-./年](\d{1,2})[-./月](\d{1,2})(?:日)?(?!\d)")
_DATE_YMD_COMPACT_PATTERN = re.compile(r"(?<!\d)(20\d{2})(\d{2})(\d{2})(?!\d)")
_DATE_YYMD_COMPACT_PATTERN = re.compile(r"(?<!\d)(\d{2})(\d{2})(\d{2})(?!\d)")

_STORE_LOCK = threading.Lock()
_JOB_LOCK = threading.Lock()
_AGENT_LOG_LOCK = threading.Lock()
_JOBS: Dict[str, Dict] = {}


def _canonical_model_name(model_name: str) -> str:
    raw = (model_name or "").strip()
    if not raw:
        return raw
    lowered = raw.lower()
    alias_map = {
        "ds v3.2": "deepseek-v3.2",
        "ds-v3.2": "deepseek-v3.2",
        "ds v3.2 thinking": "deepseek-v3.2-thinking",
        "ds-v3.2-thinking": "deepseek-v3.2-thinking",
        "ds v3.1": "deepseek-v3.1",
        "ds-v3.1": "deepseek-v3.1",
        "deepseek/deepseek-v3.2": "deepseek-v3.2",
        "deepseek/deepseek-v3.2-thinking": "deepseek-v3.2-thinking",
        "deepseek/deepseek-v3.1": "deepseek-v3.1",
        "deepseek/deepseek-v3.1-terminus": "deepseek-v3.1",
    }
    if lowered in alias_map:
        return alias_map[lowered]
    if lowered.startswith("deepseek/"):
        tail = raw.split("/", 1)[1].strip()
        if tail.lower().startswith("deepseek-"):
            return tail
    return raw


def _load_env_files() -> None:
    for filename in (".env", "env"):
        env_path = BASE_DIR / filename
        if not env_path.exists() or not env_path.is_file():
            continue
        try:
            lines = env_path.read_text(encoding="utf-8", errors="ignore").splitlines()
        except Exception:
            continue
        for raw in lines:
            line = raw.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            key = key.strip()
            value = value.strip()
            if key.startswith("export "):
                key = key.replace("export ", "", 1).strip()
            if value.startswith(("'", '"')) and value.endswith(("'", '"')) and len(value) >= 2:
                value = value[1:-1]
            if key and key not in os.environ:
                os.environ[key] = value


def _get_runtime_model_options() -> List[Dict[str, str]]:
    env_raw = (os.getenv("OPENAI_MODEL_OPTIONS") or "").strip()
    parsed: List[Dict[str, str]] = []

    if env_raw:
        for chunk in env_raw.split(","):
            token = chunk.strip()
            if not token:
                continue
            if ":" in token:
                label, value = token.split(":", 1)
                label = label.strip()
                value = _canonical_model_name(value.strip())
                if value and value in ALLOWED_RUNTIME_MODEL_VALUES:
                    parsed.append({"label": label or value, "value": value})
            else:
                normalized = _canonical_model_name(token)
                if normalized in ALLOWED_RUNTIME_MODEL_VALUES:
                    parsed.append({"label": normalized, "value": normalized})

    base = parsed if parsed else list(DEFAULT_MODEL_OPTIONS)

    for extra_value in (
        _canonical_model_name((os.getenv("OPENAI_MODEL") or "").strip()),
        _canonical_model_name((os.getenv("OPENAI_PORTRAIT_MODEL") or "").strip()),
    ):
        if extra_value and extra_value in ALLOWED_RUNTIME_MODEL_VALUES:
            base.append({"label": extra_value, "value": extra_value})

    seen = set()
    out: List[Dict[str, str]] = []
    for item in base:
        value = (item.get("value") or "").strip()
        if not value or value in seen or value not in ALLOWED_RUNTIME_MODEL_VALUES:
            continue
        seen.add(value)
        label = (item.get("label") or value).strip() or value
        out.append({"label": label, "value": value})
    return out


def _resolve_default_chat_agent(known_agents: List[str]) -> str:
    if DEFAULT_CHAT_AGENT_NAME in known_agents:
        return DEFAULT_CHAT_AGENT_NAME
    if known_agents:
        return known_agents[0]
    return DEFAULT_CHAT_AGENT_NAME


def _resolve_conversation_dispatch(payload: Dict[str, Any]) -> Tuple[str, str, bool]:
    profiles = list_agent_profiles()
    known_agents = []
    for item in profiles:
        name = str(item.get("name") or "").strip()
        if name:
            known_agents.append(name)
    known_set = set(known_agents)

    requested_agent = (payload.get("agent_name") or "").strip()
    requested_mode_raw = (
        payload.get("conversation_mode")
        or payload.get("chat_mode")
        or payload.get("mode")
        or ""
    )
    requested_mode = str(requested_mode_raw).strip().lower()
    if requested_mode not in ALLOWED_CONVERSATION_MODES:
        requested_mode = ""

    legacy_crawler = requested_agent == LEGACY_CRAWLER_AGENT_NAME
    if not requested_mode:
        if legacy_crawler:
            requested_mode = "crawler"
        elif requested_agent == CONVERSATION_MODE_AGENT_MAP.get("strategy_edit"):
            requested_mode = "strategy_edit"
        elif requested_agent == CONVERSATION_MODE_AGENT_MAP.get("strategy_analysis"):
            requested_mode = "strategy_analysis"
        else:
            requested_mode = "dialog"

    mapped_agent = CONVERSATION_MODE_AGENT_MAP.get(requested_mode) or ""
    if mapped_agent and mapped_agent in known_set:
        resolved_agent = mapped_agent
    elif requested_agent in known_set:
        resolved_agent = requested_agent
    else:
        resolved_agent = _resolve_default_chat_agent(known_agents)

    run_crawler_direct = requested_mode == "crawler" or legacy_crawler
    return resolved_agent, requested_mode, run_crawler_direct


def _build_conversation_mode_runtime() -> List[Dict[str, str]]:
    mode_labels = {
        "dialog": "对话",
        "crawler": "爬虫",
        "strategy_edit": "策略编辑",
        "strategy_analysis": "策略分析",
    }
    profiles = list_agent_profiles()
    known_agents = []
    for item in profiles:
        name = str(item.get("name") or "").strip()
        if name:
            known_agents.append(name)
    fallback_agent = _resolve_default_chat_agent(known_agents)
    known_set = set(known_agents)

    out: List[Dict[str, str]] = []
    for mode_key in ("dialog", "crawler", "strategy_edit", "strategy_analysis"):
        mapped_agent = CONVERSATION_MODE_AGENT_MAP.get(mode_key) or ""
        if mapped_agent not in known_set:
            mapped_agent = fallback_agent
        out.append(
            {
                "id": mode_key,
                "label": mode_labels.get(mode_key, mode_key),
                "agent_name": mapped_agent,
            }
        )
    return out


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _safe_markdown_stem(filename: str) -> str:
    stem = Path(filename or "").stem.strip()
    if not stem:
        return "file"
    # Windows-invalid chars and control chars
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", stem)
    stem = stem.strip(" .")
    if not stem:
        return "file"
    return stem[:80]


def _markdown_filename(resource_id: str, original_name: str) -> str:
    safe_stem = _safe_markdown_stem(original_name)
    return f"{safe_stem}__{resource_id}.md"


def _ai_summary_filename(original_name: str) -> str:
    safe_stem = _safe_markdown_stem(original_name)
    return f"{safe_stem}__ai_summary.md"


def _resource_ai_summary_path(resource: Dict) -> Path:
    original_name = (resource.get("original_name") or "").strip() or (resource.get("id") or "resource")
    return (MARKDOWN_DIR / _ai_summary_filename(original_name)).resolve()


def _resolve_markdown_path_from_resource(resource: Dict) -> Tuple[Optional[Path], str]:
    markdown_rel = (resource.get("markdown_relpath") or "").strip()
    if not markdown_rel:
        return None, "资源 markdown 路径不存在。"
    markdown_path = (BASE_DIR / markdown_rel).resolve()
    try:
        markdown_path.relative_to(MARKDOWN_DIR.resolve())
    except Exception:
        return None, "非法 markdown 路径。"
    if not markdown_path.exists() or not markdown_path.is_file():
        return None, "资源 markdown 文件不存在。"
    return markdown_path, ""


def _build_resource_markdown_download_name(resource: Dict) -> str:
    stem = _safe_markdown_stem((resource.get("original_name") or "").strip() or (resource.get("id") or "resource"))
    return f"{stem}.md"


def _compose_user_content_with_prompt_template(user_content: str, prompt_template: str) -> str:
    content = (user_content or "").strip()
    template = (prompt_template or "").strip()
    if not template:
        return content
    return (
        "请严格参考以下提示词模板组织回答：\n"
        f"{template}\n\n"
        "用户本次问题：\n"
        f"{content}"
    )


def _ensure_store() -> None:
    STORE_DIR.mkdir(parents=True, exist_ok=True)
    REFERENCE_DIR.mkdir(parents=True, exist_ok=True)
    MEMORY_DIR.mkdir(parents=True, exist_ok=True)
    AGENT_LOGS_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    MARKDOWN_DIR.mkdir(parents=True, exist_ok=True)
    CRAWLED_DIR.mkdir(parents=True, exist_ok=True)
    STRATEGY_ROOT_DIR.mkdir(parents=True, exist_ok=True)
    STRATEGY_DIR.mkdir(parents=True, exist_ok=True)


def _read_json(path: Path, default):
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def _write_json(path: Path, data) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _safe_log_bucket_name(name: str) -> str:
    raw = (name or "").strip()
    if not raw:
        return "general"
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", raw).strip(" .")
    if not safe:
        return "general"
    return safe[:120]


def _resolve_agent_log_path(payload: Optional[Dict[str, Any]]) -> Path:
    data = payload or {}
    bucket_name = ""
    for key in (
        "conversation_title",
        "strategy_name",
        "conversation_name",
        "strategy_title",
        "conversation_id",
        "strategy_id",
        "entrypoint",
    ):
        val = str(data.get(key) or "").strip()
        if val:
            bucket_name = val
            break
    safe_name = _safe_log_bucket_name(bucket_name)
    return AGENT_LOGS_DIR / f"{safe_name}.jsonl"


def _append_agent_log(event_type: str, payload: Optional[Dict[str, Any]] = None, level: str = "info") -> None:
    record = {
        "timestamp": _now_iso(),
        "event_type": (event_type or "").strip() or "agent_log",
        "level": (level or "info").strip() or "info",
        "payload": payload or {},
    }
    line = json.dumps(record, ensure_ascii=False)
    log_path = _resolve_agent_log_path(record.get("payload") or {})
    with _AGENT_LOG_LOCK:
        log_path.parent.mkdir(parents=True, exist_ok=True)
        with log_path.open("a", encoding="utf-8") as f:
            f.write(line + "\n")


def _on_agent_runtime_event(event: Dict[str, Any]) -> None:
    item = dict(event or {})
    event_name = str(item.get("event_type") or "agent_runtime_event")
    _append_agent_log(event_name, item)


set_runtime_event_hook(_on_agent_runtime_event)


def _safe_date_str(year: int, month: int, day: int) -> str:
    try:
        return datetime(year, month, day, tzinfo=timezone.utc).date().isoformat()
    except Exception:
        return ""


def _extract_first_date(text: str) -> Tuple[str, str]:
    if not text:
        return "", ""
    match = _DATE_YMD_SEP_PATTERN.search(text)
    if match:
        date_str = _safe_date_str(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        if date_str:
            return date_str, match.group(0)
    match = _DATE_YMD_COMPACT_PATTERN.search(text)
    if match:
        date_str = _safe_date_str(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        if date_str:
            return date_str, match.group(0)
    match = _DATE_YYMD_COMPACT_PATTERN.search(text)
    if match:
        yy = int(match.group(1))
        year = (2000 + yy) if yy <= 69 else (1900 + yy)
        date_str = _safe_date_str(year, int(match.group(2)), int(match.group(3)))
        if date_str:
            return date_str, match.group(0)
    return "", ""


def _scan_dir_snapshot(path: Path, recursive: bool, max_files: int = 4000) -> Dict[str, Any]:
    result: Dict[str, Any] = {
        "exists": path.exists() and path.is_dir(),
        "path": str(path),
        "file_count": 0,
        "size_bytes": 0,
        "latest_mtime_iso": "",
        "latest_mtime_date": "",
        "latest_inferred_data_date": "",
        "truncated": False,
    }
    if not result["exists"]:
        return result

    latest_mtime = 0.0
    latest_data_date = ""
    inspected = 0

    if recursive:
        iterator = path.rglob("*")
    else:
        iterator = path.glob("*")

    for item in iterator:
        if not item.is_file():
            continue
        inspected += 1
        if inspected > max_files:
            result["truncated"] = True
            break
        result["file_count"] += 1
        try:
            st = item.stat()
            result["size_bytes"] += int(st.st_size)
            if st.st_mtime > latest_mtime:
                latest_mtime = st.st_mtime
        except Exception:
            pass
        date_from_name, _ = _extract_first_date(item.name)
        if date_from_name and date_from_name > latest_data_date:
            latest_data_date = date_from_name

    if latest_mtime > 0:
        latest_mtime_dt = datetime.fromtimestamp(latest_mtime, tz=timezone.utc)
        result["latest_mtime_iso"] = latest_mtime_dt.isoformat()
        result["latest_mtime_date"] = latest_mtime_dt.date().isoformat()
    result["latest_inferred_data_date"] = latest_data_date
    return result


def _detect_date_column(schema_items: List[Tuple[str, Any]]) -> str:
    if not schema_items:
        return ""
    preferred_names = ("日期", "date", "trade_date", "时间", "datetime", "timestamp")
    lowered_map = {str(name).strip().lower(): str(name) for name, _ in schema_items}
    for key in preferred_names:
        hit = lowered_map.get(key.lower())
        if hit:
            return hit
    for name, dtype in schema_items:
        dt = str(dtype or "")
        if "Date" in dt or "Datetime" in dt:
            return str(name)
    return ""


def _schema_preview_for_parquet(path: Path, max_columns: int = 18) -> Dict[str, Any]:
    out: Dict[str, Any] = {
        "exists": path.exists() and path.is_file(),
        "col_count": 0,
        "schema_preview": [],
        "date_col": "",
        "date_min": "",
        "date_max": "",
        "error": "",
    }
    if not out["exists"]:
        return out
    try:
        import polars as pl  # type: ignore

        lf = pl.scan_parquet(str(path))
        schema = dict(getattr(lf, "schema", {}) or {})
        schema_items = list(schema.items())
        out["col_count"] = len(schema_items)
        preview_items = schema_items[:max_columns]
        out["schema_preview"] = [f"{name}:{dtype}" for name, dtype in preview_items]
        if len(schema_items) > max_columns:
            out["schema_preview"].append(f"...(+{len(schema_items) - max_columns} cols)")

        date_col = _detect_date_column(schema_items)
        out["date_col"] = date_col
        if date_col:
            try:
                date_range = (
                    lf.select(
                        pl.col(date_col).min().alias("__min_date__"),
                        pl.col(date_col).max().alias("__max_date__"),
                    )
                    .collect()
                    .to_dicts()
                )
                if date_range:
                    row = date_range[0] or {}
                    if row.get("__min_date__") is not None:
                        out["date_min"] = str(row.get("__min_date__"))
                    if row.get("__max_date__") is not None:
                        out["date_max"] = str(row.get("__max_date__"))
            except Exception:
                pass
    except Exception as exc:
        out["error"] = str(exc)
    return out


def _build_dialog_data_cache_schema_lines(data_cache_root: Path) -> List[str]:
    targets = [
        "indices/index_daily_metadata.parquet",
        "indices/index_minute_metadata.parquet",
        "stock_daily/stock_daily_metadata.parquet",
        "other/market_metadata.parquet",
        "other/market_states.parquet",
        "sectors/sectors_ths.parquet",
        "sectors/sectors_dc.parquet",
    ]
    lines = ["Core metadata schemas (column:dtype preview):"]
    for rel in targets:
        summary = _schema_preview_for_parquet(data_cache_root / rel)
        if not summary.get("exists"):
            lines.append(f"- {rel}: missing")
            continue
        if summary.get("error"):
            lines.append(f"- {rel}: read_error={summary.get('error')}")
            continue

        details = [f"cols={summary.get('col_count', 0)}"]
        date_col = str(summary.get("date_col") or "").strip()
        date_min = str(summary.get("date_min") or "").strip()
        date_max = str(summary.get("date_max") or "").strip()
        if date_col:
            details.append(f"date_col={date_col}")
        if date_min or date_max:
            details.append(f"date_range={date_min or 'unknown'} -> {date_max or 'unknown'}")
        schema_preview = ", ".join(summary.get("schema_preview") or [])
        lines.append(f"- {rel}: {'; '.join(details)}; schema={schema_preview or 'unknown'}")
    return lines


def _build_runtime_clock_context() -> str:
    now_utc = datetime.now(timezone.utc)
    now_local = datetime.now().astimezone()
    today_local = now_local.date()
    yesterday_local = today_local - timedelta(days=1)

    lines = [
        "Runtime clock context (auto-executed):",
        f"- runtime_now_utc: {now_utc.isoformat()}",
        f"- runtime_now_local: {now_local.isoformat()}",
        f"- runtime_today_local: {today_local.isoformat()}",
        f"- runtime_yesterday_local: {yesterday_local.isoformat()}",
        "- date_judgement_rule: 如果用户给出的日期 <= runtime_today_local，不能判定为“未来日期”；若数据缺失，应说明“缓存/数据源缺失”而非“未来不可得”。",
    ]
    return "\n".join(lines)


def _build_dialog_data_cache_context(user_content: str) -> str:
    data_cache_root = BASE_DIR / "data_cache"
    targets: List[Tuple[str, Path, bool]] = [
        ("data_cache_root", data_cache_root, False),
        ("indices", data_cache_root / "indices", False),
        ("stock_daily", data_cache_root / "stock_daily", False),
        ("stock_minute", data_cache_root / "stock_minute", False),
        ("other", data_cache_root / "other", False),
        ("strategy_watch", data_cache_root / "strategy_watch", True),
    ]

    snapshots: List[Tuple[str, Dict[str, Any]]] = []
    for name, path, recursive in targets:
        snapshots.append((name, _scan_dir_snapshot(path, recursive=recursive)))

    all_dates = []
    for _, snap in snapshots:
        for key in ("latest_inferred_data_date", "latest_mtime_date"):
            value = str(snap.get(key) or "").strip()
            if value:
                all_dates.append(value)
    latest_available_date = max(all_dates) if all_dates else ""

    requested_date, _token = _extract_first_date(user_content or "")
    request_hint = ""
    if requested_date and latest_available_date and requested_date > latest_available_date:
        request_hint = (
            f"用户请求日期: {requested_date}；缓存观测到的最新日期: {latest_available_date}。"
            "请求日期可能超出当前缓存范围，回答时应明确该限制并避免编造。"
        )

    lines = [
        "Data cache inspector snapshot (auto-executed):",
        f"- cache_root: {data_cache_root.as_posix()}",
        f"- latest_available_data_date: {latest_available_date or 'unknown'}",
    ]
    if request_hint:
        lines.append(f"- request_vs_cache: {request_hint}")

    for name, snap in snapshots:
        if not snap.get("exists"):
            lines.append(f"- {name}: missing")
            continue
        lines.append(
            "- {name}: files={files}, size_bytes={size}, latest_mtime={mtime}, mtime_date={mdate}, inferred_date={idata}{tail}".format(
                name=name,
                files=snap.get("file_count", 0),
                size=snap.get("size_bytes", 0),
                mtime=snap.get("latest_mtime_iso") or "unknown",
                mdate=snap.get("latest_mtime_date") or "unknown",
                idata=snap.get("latest_inferred_data_date") or "unknown",
                tail=", truncated=true" if snap.get("truncated") else "",
            )
        )
    lines.extend(["", *_build_dialog_data_cache_schema_lines(data_cache_root)])

    lines.append("Use these facts first when answering data-availability questions.")
    return "\n".join(lines)


def _normalize_date_only(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return ""
    lowered = raw.replace("Z", "+00:00")
    try:
        dt_obj = datetime.fromisoformat(lowered)
        return dt_obj.date().isoformat()
    except Exception:
        pass
    date_str, _ = _extract_first_date(raw)
    return date_str


def _normalize_datetime_to_iso(value: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return ""
    lowered = raw.replace("Z", "+00:00").replace("/", "-")
    try:
        return datetime.fromisoformat(lowered).astimezone(timezone.utc).isoformat()
    except Exception:
        pass

    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
        try:
            parsed = datetime.strptime(lowered, fmt)
            if fmt == "%Y-%m-%d":
                parsed = parsed.replace(hour=0, minute=0, second=0)
            return parsed.replace(tzinfo=timezone.utc).isoformat()
        except Exception:
            continue
    return ""


def _parse_frontmatter_map(text: str) -> Dict[str, str]:
    if not text.startswith("---"):
        return {}
    parts = text.split("---", 2)
    if len(parts) < 3:
        return {}
    front = parts[1]
    result: Dict[str, str] = {}
    for raw in front.splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or ":" not in line:
            continue
        key, val = line.split(":", 1)
        key = key.strip().lower()
        val = val.strip().strip('"').strip("'")
        if key:
            result[key] = val
    return result


def _read_markdown_text_by_relpath(rel_path: str) -> str:
    target = (rel_path or "").strip()
    if not target or "://" in target:
        return ""
    md_path = BASE_DIR / target
    if not md_path.exists() or not md_path.is_file():
        return ""
    return md_path.read_text(encoding="utf-8", errors="ignore")


def _enrich_resource_temporal_fields(record: Dict, markdown_text: str = "") -> bool:
    changed = False
    record.setdefault("content_time", "")
    record.setdefault("content_time_type", "unknown")
    record.setdefault("content_time_confidence", 0.0)
    record.setdefault("content_time_evidence", "")
    record.setdefault("published_at", "")

    if not isinstance(record.get("content_time_confidence"), (int, float)):
        record["content_time_confidence"] = 0.0
        changed = True

    normalized_published = _normalize_datetime_to_iso(record.get("published_at", ""))
    if normalized_published != (record.get("published_at") or ""):
        record["published_at"] = normalized_published
        changed = True

    normalized_content_time = _normalize_date_only(record.get("content_time", ""))
    if normalized_content_time != (record.get("content_time") or ""):
        record["content_time"] = normalized_content_time
        changed = True

    if not record.get("content_time"):
        inferred_date, hit = _extract_first_date(record.get("original_name") or "")
        if inferred_date:
            record["content_time"] = inferred_date
            record["content_time_type"] = "inferred_filename"
            record["content_time_confidence"] = 0.70
            record["content_time_evidence"] = f"filename:{hit}"
            changed = True

    text = markdown_text or ""
    if text:
        front = _parse_frontmatter_map(text)

        if not record.get("published_at"):
            for key in ("published_at", "publish_time", "published_time", "publishdate", "date"):
                candidate = _normalize_datetime_to_iso(front.get(key, ""))
                if candidate:
                    record["published_at"] = candidate
                    changed = True
                    break

        better_content_time = ""
        better_type = ""
        better_evidence = ""
        better_conf = 0.0

        for key in ("content_time", "recorded_at", "publish_date", "published_at", "date"):
            raw = front.get(key, "")
            normalized = _normalize_date_only(raw)
            if not normalized:
                continue
            if key in {"published_at", "publish_date"}:
                better_type = "published"
                better_conf = 0.95
            elif key in {"recorded_at"}:
                better_type = "recorded"
                better_conf = 0.90
            else:
                better_type = "inferred_text"
                better_conf = 0.82
            better_content_time = normalized
            better_evidence = f"frontmatter:{key}"
            break

        if not better_content_time:
            sample = "\n".join(
                [
                    front.get("title", ""),
                    text[:5000],
                ]
            )
            inferred_date, hit = _extract_first_date(sample)
            if inferred_date:
                better_content_time = inferred_date
                better_type = "inferred_text"
                better_conf = 0.55
                better_evidence = f"content:{hit}"

        current_type = (record.get("content_time_type") or "unknown").strip() or "unknown"
        priority = {
            "unknown": 0,
            "inferred_text": 1,
            "inferred_filename": 2,
            "recorded": 3,
            "published": 4,
        }
        if better_content_time:
            should_replace = not record.get("content_time")
            if not should_replace:
                should_replace = priority.get(better_type, 0) >= priority.get(current_type, 0)
            if should_replace:
                record["content_time"] = better_content_time
                record["content_time_type"] = better_type
                record["content_time_confidence"] = better_conf
                record["content_time_evidence"] = better_evidence
                changed = True

    if not record.get("content_time_type"):
        record["content_time_type"] = "unknown"
        changed = True

    if not record.get("content_time"):
        if record.get("content_time_confidence") != 0.0:
            record["content_time_confidence"] = 0.0
            changed = True
        if record.get("content_time_type") != "unknown":
            record["content_time_type"] = "unknown"
            changed = True

    return changed


def _normalize_resource_record(record: Dict) -> Dict:
    group_id = (record.get("group_id") or "").strip()
    group_name = (record.get("group_name") or "").strip()
    if group_name and not group_id:
        group_id = f"legacy_{group_name}"
    if not group_name:
        group_name = DEFAULT_GROUP_NAME
    if not group_id:
        group_id = f"legacy_{DEFAULT_GROUP_NAME}"
    record["group_id"] = group_id
    record["group_name"] = group_name
    record["strategy_id"] = (record.get("strategy_id") or "").strip()
    record["strategy_name"] = (record.get("strategy_name") or "").strip()
    record.setdefault("progress", 0)
    record.setdefault("progress_message", "")
    _enrich_resource_temporal_fields(record)
    return record


def _set_job(job_id: str, **fields) -> Dict:
    with _JOB_LOCK:
        job = _JOBS.get(job_id, {})
        job.update(fields)
        _JOBS[job_id] = job
        snapshot = dict(job)

    should_log = any(key in fields for key in ("status", "progress", "message", "error", "started_at", "finished_at"))
    if should_log:
        _append_agent_log(
            "todo_job_progress",
            {
                "job_id": job_id,
                "status": snapshot.get("status"),
                "progress": snapshot.get("progress"),
                "message": snapshot.get("message"),
                "error": snapshot.get("error"),
                "resource_id": snapshot.get("resource_id"),
                "strategy_id": snapshot.get("strategy_id"),
                "strategy_name": snapshot.get("strategy_name"),
                "created_at": snapshot.get("created_at"),
                "started_at": snapshot.get("started_at"),
                "finished_at": snapshot.get("finished_at"),
            },
        )
    return snapshot


def _get_job(job_id: str) -> Dict:
    with _JOB_LOCK:
        return dict(_JOBS.get(job_id) or {})


def _list_jobs() -> List[Dict]:
    with _JOB_LOCK:
        jobs = list(_JOBS.values())
    return sorted(jobs, key=lambda x: x.get("created_at", ""), reverse=True)


def _read_agent_logs(limit: int = 200, bucket: str = "") -> List[Dict[str, Any]]:
    if limit <= 0:
        return []
    candidates: List[Path] = []
    if bucket:
        candidates = [AGENT_LOGS_DIR / f"{_safe_log_bucket_name(bucket)}.jsonl"]
    else:
        if AGENT_LOGS_DIR.exists() and AGENT_LOGS_DIR.is_dir():
            candidates.extend(sorted(AGENT_LOGS_DIR.glob("*.jsonl")))
        if LEGACY_AGENT_LOGS_FILE.exists() and LEGACY_AGENT_LOGS_FILE.is_file():
            candidates.append(LEGACY_AGENT_LOGS_FILE)
    if not candidates:
        return []

    out: List[Dict[str, Any]] = []
    with _AGENT_LOG_LOCK:
        for path in candidates:
            if not path.exists() or not path.is_file():
                continue
            try:
                with path.open("r", encoding="utf-8", errors="ignore") as f:
                    lines = f.readlines()
            except Exception:
                continue
            for raw in lines[-limit:]:
                text = (raw or "").strip()
                if not text:
                    continue
                try:
                    item = json.loads(text)
                    if isinstance(item, dict):
                        out.append(item)
                except Exception:
                    continue
    out.sort(key=lambda x: str(x.get("timestamp") or ""), reverse=True)
    return out[:limit]


def _update_resource_fields(resource_id: str, **fields) -> None:
    with _STORE_LOCK:
        payload = _load_resources_payload()
        resources = payload.get("resources", [])
        updated = False
        for item in resources:
            if item.get("id") == resource_id:
                item.update(fields)
                updated = True
                break
        if updated:
            payload["resources"] = resources
            _save_resources_payload(payload)


def _process_resource_job(
    job_id: str,
    resource_id: str,
    upload_path: Path,
    markdown_path: Path,
    whisper_model: str,
) -> None:
    _set_job(
        job_id,
        status="running",
        progress=0,
        message="开始处理",
        started_at=_now_iso(),
    )
    _update_resource_fields(
        resource_id,
        status="processing",
        progress=0,
        progress_message="开始处理",
        error="",
    )

    last_percent = -1
    last_ts = 0.0

    def _progress_cb(current: int, total: int, message: str) -> None:
        nonlocal last_percent, last_ts
        if total <= 0:
            return
        percent = int(max(0, min(100, (current / total) * 100)))
        now_ts = time.time()
        if percent == last_percent and (now_ts - last_ts) < 0.6:
            return
        last_percent = percent
        last_ts = now_ts
        msg = (message or "").strip()
        _set_job(job_id, progress=percent, message=msg)
        _update_resource_fields(resource_id, progress=percent, progress_message=msg)

    status = "failed"
    error = ""
    try:
        status, error = convert_file_to_markdown_via_skill(
            input_file=upload_path,
            output_markdown=markdown_path,
            whisper_model=whisper_model,
            progress_callback=_progress_cb,
        )
    except Exception as exc:
        status = "failed"
        error = str(exc)

    final_status = "ok" if status == "ok" else "failed"
    final_message = "完成" if final_status == "ok" else (error or "转换失败")
    _set_job(
        job_id,
        status=final_status,
        progress=100,
        message=final_message,
        error=error,
        finished_at=_now_iso(),
    )
    _update_resource_fields(
        resource_id,
        status=final_status,
        progress=100,
        progress_message=final_message,
        error=error,
    )
    try:
        if upload_path.exists() and upload_path.is_file():
            upload_path.unlink()
    except Exception:
        pass


def _load_resources_payload() -> Dict:
    _ensure_store()
    raw = _read_json(RESOURCES_FILE, {"resources": [], "active_resource_id": "", "groups": []})
    if isinstance(raw, list):
        raw = {"resources": raw, "active_resource_id": "", "groups": []}
    resources = raw.get("resources", [])
    manual_groups = _normalize_group_records(raw.get("groups") if isinstance(raw.get("groups"), list) else [])
    normalized: List[Dict] = []
    touched = False
    for item in resources:
        if isinstance(item, dict):
            normalized_item = _normalize_resource_record(item)
            md_rel = (normalized_item.get("markdown_relpath") or "").strip()
            source_type = (normalized_item.get("source_type") or "").strip().lower()
            content_time = (normalized_item.get("content_time") or "").strip()
            content_time_type = (normalized_item.get("content_time_type") or "").strip().lower()
            published_at = (normalized_item.get("published_at") or "").strip()
            should_probe_markdown = (
                (not content_time)
                or (content_time_type == "unknown")
                or (source_type == "url" and not published_at)
            )
            if should_probe_markdown and md_rel:
                md_text = _read_markdown_text_by_relpath(md_rel)
                if md_text and _enrich_resource_temporal_fields(normalized_item, markdown_text=md_text):
                    touched = True
            normalized.append(normalized_item)
    active_resource_id = (raw.get("active_resource_id") or "").strip()
    if touched:
        _save_resources_payload(
            {"resources": normalized, "active_resource_id": active_resource_id, "groups": manual_groups}
        )
    return {"resources": normalized, "active_resource_id": active_resource_id, "groups": manual_groups}


def _save_resources_payload(payload: Dict) -> None:
    _ensure_store()
    _write_json(
        RESOURCES_FILE,
        {
            "resources": payload.get("resources", []),
            "active_resource_id": (payload.get("active_resource_id") or "").strip(),
            "groups": _normalize_group_records(payload.get("groups") if isinstance(payload.get("groups"), list) else []),
        },
    )


def _load_resources() -> List[Dict]:
    return _load_resources_payload().get("resources", [])


def _save_resources(resources: List[Dict]) -> None:
    payload = _load_resources_payload()
    payload["resources"] = resources
    valid_ids = {item.get("id") for item in resources}
    if payload.get("active_resource_id") not in valid_ids:
        payload["active_resource_id"] = ""
    _save_resources_payload(payload)


def _load_conversations() -> List[Dict]:
    _ensure_store()
    payload = _read_json(CONVERSATIONS_FILE, {"conversations": []})
    return payload.get("conversations", [])


def _save_conversations(conversations: List[Dict]) -> None:
    _ensure_store()
    _write_json(CONVERSATIONS_FILE, {"conversations": conversations})


def _normalize_strategy_id(strategy_id: str, allow_empty: bool = False) -> str:
    raw = (strategy_id or "").strip()
    if not raw:
        if allow_empty:
            return ""
        raw = f"strategy_{uuid.uuid4().hex[:12]}"
    return "".join(ch if (ch.isalnum() or ch in {"-", "_"}) else "_" for ch in raw)


def _normalize_strategy_record(strategy: Dict) -> Dict:
    item = dict(strategy or {})
    item["id"] = _normalize_strategy_id(item.get("id", ""))
    item["name"] = (item.get("name") or "").strip() or item["id"]
    item["description"] = (item.get("description") or "").strip()
    item["view_type"] = (item.get("view_type") or DEFAULT_STRATEGY_VIEW).strip() or DEFAULT_STRATEGY_VIEW
    item["config"] = item.get("config") if isinstance(item.get("config"), dict) else {}
    item["created_at"] = (item.get("created_at") or _now_iso()).strip()
    item["updated_at"] = (item.get("updated_at") or item["created_at"]).strip()
    return item


def _safe_strategy_filename_component(name: str) -> str:
    raw = (name or "").strip()
    if not raw:
        return ""
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", raw).strip(" .")
    safe = re.sub(r"\s+", "_", safe)
    return safe[:96]


def _default_strategy_file_path(strategy_id: str) -> Path:
    return STRATEGY_DIR / f"{_normalize_strategy_id(strategy_id)}.json"


def _strategy_storage_filename(strategy: Dict) -> str:
    sid = _normalize_strategy_id((strategy or {}).get("id", ""))
    safe_name = _safe_strategy_filename_component((strategy or {}).get("name", ""))
    stem = safe_name or sid or "strategy"
    return f"{stem}.json"


def _resolve_strategy_file_from_index(strategy_id: str, strategy_files: Dict[str, str]) -> Path:
    sid = _normalize_strategy_id(strategy_id)
    mapped = strategy_files.get(sid)
    candidates: List[Path] = []
    if isinstance(mapped, str) and mapped.strip():
        raw = mapped.strip().replace("\\", "/")
        rel_path = Path(raw)
        if raw.lower().endswith(".json"):
            if not rel_path.is_absolute() and ".." not in rel_path.parts:
                parts = [p for p in rel_path.parts if p]
                if len(parts) == 1:
                    # Bare filename -> always resolve inside strategy/strategies/
                    candidates.append(STRATEGY_DIR / parts[0])
                elif parts and parts[0].lower() == "strategies":
                    # Relative path explicitly under strategies/
                    candidates.append(STRATEGY_ROOT_DIR / rel_path)

    candidates.append(_default_strategy_file_path(sid))
    for path in candidates:
        if path.exists():
            return path
    return candidates[0]


def _load_strategies_payload() -> Dict:
    _ensure_store()
    strategies: List[Dict] = []
    active_strategy_id = ""

    # New layout: one JSON file per strategy + index file
    if STRATEGY_INDEX_FILE.exists():
        raw_index = _read_json(STRATEGY_INDEX_FILE, {"strategy_ids": [], "active_strategy_id": ""})
        strategy_ids = raw_index.get("strategy_ids", [])
        strategy_files = raw_index.get("strategy_files") if isinstance(raw_index.get("strategy_files"), dict) else {}
        seen = set()

        if isinstance(strategy_ids, list):
            for sid in strategy_ids:
                sid = _normalize_strategy_id(str(sid))
                if sid in seen:
                    continue
                path = _resolve_strategy_file_from_index(sid, strategy_files)
                rec = _read_json(path, None)
                # Fallback to default id-based filename when mapped filename is stale.
                if (not isinstance(rec, dict) or not rec.get("id")) and path != _default_strategy_file_path(sid):
                    rec = _read_json(_default_strategy_file_path(sid), None)
                if isinstance(rec, dict) and rec.get("id"):
                    normalized = _normalize_strategy_record(rec)
                    strategies.append(normalized)
                    seen.add(normalized["id"])

        # Backfill: include strategy files that are not present in index
        for path in STRATEGY_DIR.glob("*.json"):
            rec = _read_json(path, None)
            if not isinstance(rec, dict) or not rec.get("id"):
                continue
            normalized = _normalize_strategy_record(rec)
            if normalized["id"] in seen:
                continue
            strategies.append(normalized)
            seen.add(normalized["id"])

        active_strategy_id = _normalize_strategy_id(raw_index.get("active_strategy_id", ""), allow_empty=True)
        valid_ids = {item["id"] for item in strategies}
        if active_strategy_id not in valid_ids:
            active_strategy_id = strategies[0]["id"] if strategies else ""

        # Heal index/layout if needed
        _save_strategies_payload({"strategies": strategies, "active_strategy_id": active_strategy_id})
        return {"strategies": strategies, "active_strategy_id": active_strategy_id}

    # Legacy layout fallback: migrate from aggregated strategies.json
    raw = _read_json(STRATEGIES_FILE, {"strategies": [], "active_strategy_id": ""})
    if isinstance(raw, list):
        raw = {"strategies": raw, "active_strategy_id": ""}

    raw_strategies = raw.get("strategies", [])
    if isinstance(raw_strategies, list):
        for item in raw_strategies:
            if isinstance(item, dict):
                strategies.append(_normalize_strategy_record(item))

    active_strategy_id = _normalize_strategy_id(raw.get("active_strategy_id", ""), allow_empty=True)
    valid_ids = {item["id"] for item in strategies}
    if active_strategy_id not in valid_ids:
        active_strategy_id = strategies[0]["id"] if strategies else ""

    if strategies or active_strategy_id:
        _save_strategies_payload({"strategies": strategies, "active_strategy_id": active_strategy_id})

    return {"strategies": strategies, "active_strategy_id": active_strategy_id}


def _save_strategies_payload(payload: Dict) -> None:
    _ensure_store()
    raw_list = payload.get("strategies", [])
    if not isinstance(raw_list, list):
        raw_list = []

    # Keep input order while removing duplicated ids
    seen = set()
    strategies: List[Dict] = []
    for item in raw_list:
        if not isinstance(item, dict):
            continue
        normalized = _normalize_strategy_record(item)
        sid = normalized["id"]
        if sid in seen:
            continue
        strategies.append(normalized)
        seen.add(sid)

    active_strategy_id = _normalize_strategy_id(payload.get("active_strategy_id", ""), allow_empty=True)
    valid_ids = {item["id"] for item in strategies}
    if active_strategy_id not in valid_ids:
        active_strategy_id = strategies[0]["id"] if strategies else ""

    current_files = {path.name for path in STRATEGY_DIR.glob("*.json")}
    target_files = set()
    ordered_ids = []
    strategy_files: Dict[str, str] = {}
    used_filenames = set()

    for item in strategies:
        sid = item["id"]
        ordered_ids.append(sid)
        filename = _strategy_storage_filename(item)
        if filename in used_filenames:
            safe_name = _safe_strategy_filename_component(item.get("name") or "")
            filename = f"{safe_name or sid}__{sid}.json"
        used_filenames.add(filename)
        strategy_files[sid] = f"strategies/{filename}"
        path = STRATEGY_DIR / filename
        target_files.add(filename)
        _write_json(path, item)

    # Remove stale strategy files
    for filename in current_files - target_files:
        stale_path = STRATEGY_DIR / filename
        try:
            stale_path.unlink()
        except Exception:
            pass

    # New index file
    _write_json(
        STRATEGY_INDEX_FILE,
        {
            "strategy_ids": ordered_ids,
            "strategy_files": strategy_files,
            "active_strategy_id": active_strategy_id,
        },
    )

    # Backward-compatible mirror for existing external readers
    _write_json(
        STRATEGIES_FILE,
        {
            "strategies": strategies,
            "active_strategy_id": active_strategy_id,
        },
    )


def _normalize_memory_subscription(item: Dict) -> Optional[Dict]:
    if not isinstance(item, dict):
        return None
    group_id = (item.get("group_id") or "").strip()
    if not group_id:
        return None
    first_bound_at = (item.get("first_bound_at") or _now_iso()).strip()
    last_sync_at = (item.get("last_sync_at") or "").strip()
    return {
        "group_id": group_id,
        "first_bound_at": first_bound_at,
        "last_sync_at": last_sync_at,
    }


def _normalize_memory_profile_record(profile: Dict) -> Dict:
    raw_id = (profile.get("id") or "").strip()
    normalized_id = _normalize_strategy_id(raw_id, allow_empty=True) or _gen_id("mp")
    subs_raw = profile.get("group_subscriptions") if isinstance(profile.get("group_subscriptions"), list) else []
    normalized_subs: List[Dict] = []
    seen_subs = set()
    for item in subs_raw:
        sub = _normalize_memory_subscription(item)
        if not sub:
            continue
        gid = sub["group_id"]
        if gid in seen_subs:
            continue
        seen_subs.add(gid)
        normalized_subs.append(sub)
    created_at = (profile.get("created_at") or _now_iso()).strip()
    updated_at = (profile.get("updated_at") or created_at).strip()
    return {
        "id": normalized_id,
        "name": (profile.get("name") or "").strip() or normalized_id,
        "description": (profile.get("description") or "").strip(),
        "source_blogger": (profile.get("source_blogger") or "").strip(),
        "created_at": created_at,
        "updated_at": updated_at,
        "group_subscriptions": normalized_subs,
    }


def _load_memory_profiles_payload() -> Dict:
    _ensure_store()
    raw = _read_json(MEMORY_PROFILES_FILE, {"profiles": [], "active_profile_id": ""})
    if isinstance(raw, list):
        raw = {"profiles": raw, "active_profile_id": ""}

    profiles: List[Dict] = []
    seen = set()
    for item in (raw.get("profiles") or []):
        if not isinstance(item, dict):
            continue
        normalized = _normalize_memory_profile_record(item)
        pid = normalized["id"]
        if pid in seen:
            continue
        seen.add(pid)
        profiles.append(normalized)

    active_profile_id = (raw.get("active_profile_id") or "").strip()
    valid_ids = {item["id"] for item in profiles}
    if active_profile_id not in valid_ids:
        active_profile_id = ""
    return {"profiles": profiles, "active_profile_id": active_profile_id}


def _save_memory_profiles_payload(payload: Dict) -> None:
    _ensure_store()
    raw_list = payload.get("profiles")
    if not isinstance(raw_list, list):
        raw_list = []
    profiles: List[Dict] = []
    seen = set()
    for item in raw_list:
        if not isinstance(item, dict):
            continue
        normalized = _normalize_memory_profile_record(item)
        pid = normalized["id"]
        if pid in seen:
            continue
        seen.add(pid)
        profiles.append(normalized)
    active_profile_id = (payload.get("active_profile_id") or "").strip()
    valid_ids = {item["id"] for item in profiles}
    if active_profile_id not in valid_ids:
        active_profile_id = ""
    _write_json(
        MEMORY_PROFILES_FILE,
        {
            "profiles": profiles,
            "active_profile_id": active_profile_id,
        },
    )


def _normalize_memory_link(item: Dict) -> Optional[Dict]:
    if not isinstance(item, dict):
        return None
    profile_id = (item.get("profile_id") or "").strip()
    resource_id = (item.get("resource_id") or "").strip()
    if not profile_id or not resource_id:
        return None
    bind_source = (item.get("bind_source") or "manual").strip().lower()
    if bind_source not in {"manual", "group"}:
        bind_source = "manual"
    group_id = (item.get("group_id") or "").strip()
    added_at = (item.get("added_at") or _now_iso()).strip()
    return {
        "profile_id": profile_id,
        "resource_id": resource_id,
        "bind_source": bind_source,
        "group_id": group_id,
        "added_at": added_at,
    }


def _load_memory_links() -> List[Dict]:
    _ensure_store()
    raw = _read_json(MEMORY_LINKS_FILE, {"links": []})
    if isinstance(raw, list):
        raw = {"links": raw}
    links: List[Dict] = []
    seen = set()
    for item in (raw.get("links") or []):
        normalized = _normalize_memory_link(item)
        if not normalized:
            continue
        key = (normalized["profile_id"], normalized["resource_id"])
        if key in seen:
            continue
        seen.add(key)
        links.append(normalized)
    return links


def _save_memory_links(links: List[Dict]) -> None:
    _ensure_store()
    normalized_links: List[Dict] = []
    seen = set()
    for item in (links or []):
        normalized = _normalize_memory_link(item)
        if not normalized:
            continue
        key = (normalized["profile_id"], normalized["resource_id"])
        if key in seen:
            continue
        seen.add(key)
        normalized_links.append(normalized)
    _write_json(MEMORY_LINKS_FILE, {"links": normalized_links})


def _default_memory_portrait(profile_id: str) -> Dict:
    return {
        "profile_id": profile_id,
        "methodology": "",
        "tactics": "",
        "views": "",
        "operations": "",
        "risk_rules": "",
        "style_constraints": "",
        "evidence_refs": [],
        "updated_at": _now_iso(),
        "updated_by": "",
    }


def _normalize_evidence_ref(item: Dict) -> Optional[Dict]:
    if not isinstance(item, dict):
        return None
    resource_id = (item.get("resource_id") or "").strip()
    quote = (item.get("quote") or "").strip()
    if not resource_id and not quote:
        return None
    return {
        "id": (item.get("id") or _gen_id("ev")).strip(),
        "resource_id": resource_id,
        "topic": (item.get("topic") or "").strip(),
        "quote": quote,
        "source_time": _normalize_date_only(item.get("source_time") or ""),
        "source_time_type": (item.get("source_time_type") or "").strip(),
        "timecode": (item.get("timecode") or "").strip(),
    }


def _normalize_memory_portrait(item: Dict) -> Optional[Dict]:
    if not isinstance(item, dict):
        return None
    profile_id = (item.get("profile_id") or "").strip()
    if not profile_id:
        return None
    normalized = _default_memory_portrait(profile_id)
    for key in ("methodology", "tactics", "views", "operations", "risk_rules", "style_constraints"):
        normalized[key] = (item.get(key) or "").strip()
    refs_raw = item.get("evidence_refs") if isinstance(item.get("evidence_refs"), list) else []
    refs: List[Dict] = []
    seen = set()
    for ref in refs_raw:
        normalized_ref = _normalize_evidence_ref(ref)
        if not normalized_ref:
            continue
        rid = normalized_ref["id"]
        if rid in seen:
            continue
        seen.add(rid)
        refs.append(normalized_ref)
    normalized["evidence_refs"] = refs
    normalized["updated_at"] = (item.get("updated_at") or _now_iso()).strip()
    normalized["updated_by"] = (item.get("updated_by") or "").strip()
    return normalized


def _load_memory_portraits() -> Dict[str, Dict]:
    _ensure_store()
    raw = _read_json(MEMORY_PORTRAITS_FILE, {"portraits": []})
    if isinstance(raw, list):
        raw = {"portraits": raw}
    portraits: Dict[str, Dict] = {}
    for item in (raw.get("portraits") or []):
        normalized = _normalize_memory_portrait(item)
        if not normalized:
            continue
        portraits[normalized["profile_id"]] = normalized
    return portraits


def _save_memory_portraits(portraits: Dict[str, Dict]) -> None:
    _ensure_store()
    if not isinstance(portraits, dict):
        portraits = {}
    records: List[Dict] = []
    for profile_id, value in portraits.items():
        rec = dict(value or {})
        rec["profile_id"] = profile_id
        normalized = _normalize_memory_portrait(rec)
        if normalized:
            records.append(normalized)
    _write_json(MEMORY_PORTRAITS_FILE, {"portraits": records})


def _find_memory_profile(profile_id: str) -> Optional[Dict]:
    target_id = (profile_id or "").strip()
    if not target_id:
        return None
    payload = _load_memory_profiles_payload()
    return _find_by_id(payload.get("profiles", []), target_id)


def _collect_profile_resource_snippets(
    profile_id: str,
    max_chars: int = 24000,
    max_items: int = 10,
    per_item_chars: int = 2200,
) -> Tuple[str, List[str], List[Dict]]:
    links = [x for x in _load_memory_links() if x.get("profile_id") == profile_id]
    links.sort(key=lambda x: x.get("added_at", ""), reverse=True)
    resources_payload = _load_resources_payload()
    resources = resources_payload.get("resources", [])
    resource_map = {item.get("id"): item for item in resources}

    used_ids: List[str] = []
    skipped: List[Dict] = []
    blocks: List[str] = []
    total = 0
    for link in links:
        if len(used_ids) >= max_items:
            break
        rid = (link.get("resource_id") or "").strip()
        item = resource_map.get(rid)
        if not item:
            skipped.append({"resource_id": rid, "reason": "资源不存在"})
            continue
        if (item.get("status") or "") != "ok":
            skipped.append({"resource_id": rid, "reason": f"资源状态异常: {item.get('status') or 'unknown'}"})
            continue
        md_text = _read_markdown_content(item).strip()
        if not md_text:
            skipped.append({"resource_id": rid, "reason": "markdown 文件不存在或为空"})
            continue
        if not _markdown_has_meaningful_content(md_text):
            skipped.append({"resource_id": rid, "reason": "markdown 信息量不足"})
            continue

        snippet = md_text[:per_item_chars]
        name = item.get("original_name") or rid
        content_time = item.get("content_time") or ""
        time_label = f"（内容时间: {content_time}）" if content_time else ""
        block = f"### 资料: {name}{time_label}\n\n{snippet}"
        projected = total + len(block)
        if projected > max_chars:
            remain = max_chars - total
            if remain <= 0:
                break
            block = block[:remain]
            projected = total + len(block)
        blocks.append(block)
        used_ids.append(rid)
        total = projected
        if total >= max_chars:
            break
    return "\n\n".join(blocks).strip(), used_ids, skipped


def _build_memory_profile_context(memory_profile_id: str) -> Tuple[str, Dict[str, Any]]:
    profile_id = (memory_profile_id or "").strip()
    if not profile_id:
        return "", {"memory_profile_id": "", "used_memory_resource_ids": [], "skipped_memory_refs": []}

    profiles_payload = _load_memory_profiles_payload()
    profiles = profiles_payload.get("profiles", [])
    profile = _find_by_id(profiles, profile_id)
    if not profile:
        return "", {
            "memory_profile_id": "",
            "used_memory_resource_ids": [],
            "skipped_memory_refs": [{"profile_id": profile_id, "reason": "人格不存在"}],
        }

    portraits = _load_memory_portraits()
    portrait = portraits.get(profile_id) or _default_memory_portrait(profile_id)

    lines: List[str] = [
        f"长期记忆人格：{profile.get('name') or profile_id}",
        f"- 人格ID: {profile_id}",
    ]
    description = (profile.get("description") or "").strip()
    if description:
        lines.append(f"- 人格说明: {description}")
    source_blogger = (profile.get("source_blogger") or "").strip()
    if source_blogger:
        lines.append(f"- 来源博主: {source_blogger}")

    sections = [
        ("交易方法论", "methodology"),
        ("交易手法", "tactics"),
        ("观点", "views"),
        ("交易操作", "operations"),
        ("风控规则", "risk_rules"),
        ("风格约束", "style_constraints"),
    ]
    for title, key in sections:
        text = (portrait.get(key) or "").strip()
        if not text:
            continue
        lines.extend(["", f"## {title}", text])

    evidence_refs = portrait.get("evidence_refs") if isinstance(portrait.get("evidence_refs"), list) else []
    if evidence_refs:
        lines.extend(["", "## 证据索引"])
        for idx, ref in enumerate(evidence_refs, start=1):
            resource_id = (ref.get("resource_id") or "").strip() if isinstance(ref, dict) else ""
            quote = (ref.get("quote") or "").strip() if isinstance(ref, dict) else ""
            topic = (ref.get("topic") or "").strip() if isinstance(ref, dict) else ""
            source_time = (ref.get("source_time") or "").strip() if isinstance(ref, dict) else ""
            timecode = (ref.get("timecode") or "").strip() if isinstance(ref, dict) else ""
            meta = " | ".join(
                [x for x in [topic, source_time, timecode, resource_id] if x]
            )
            if quote:
                lines.append(f"{idx}. {quote} {'（' + meta + '）' if meta else ''}")
            elif meta:
                lines.append(f"{idx}. {meta}")

    base_context = "\n".join(lines).strip()
    remain_chars = max(0, MAX_MEMORY_CONTEXT_CHARS - len(base_context) - 2_000)
    snippets, used_ids, skipped = _collect_profile_resource_snippets(
        profile_id,
        max_chars=max(8_000, remain_chars),
        max_items=8,
        per_item_chars=1_800,
    )
    if snippets:
        base_context += "\n\n## 关联资料摘录\n\n" + snippets

    if len(base_context) > MAX_MEMORY_CONTEXT_CHARS:
        base_context = base_context[:MAX_MEMORY_CONTEXT_CHARS]

    return base_context.strip(), {
        "memory_profile_id": profile_id,
        "memory_profile_name": profile.get("name") or profile_id,
        "used_memory_resource_ids": used_ids,
        "skipped_memory_refs": skipped,
    }


def _memory_portrait_sections(portrait: Dict) -> List[Tuple[str, str]]:
    source = portrait if isinstance(portrait, dict) else {}
    return [
        ("交易方法论", str(source.get("methodology") or "").strip()),
        ("交易手法", str(source.get("tactics") or "").strip()),
        ("观点", str(source.get("views") or "").strip()),
        ("交易操作", str(source.get("operations") or "").strip()),
        ("风控规则", str(source.get("risk_rules") or "").strip()),
        ("风格约束", str(source.get("style_constraints") or "").strip()),
    ]


def _merge_portrait_override(base_portrait: Dict, override: Any) -> Dict:
    merged = dict(base_portrait or {})
    if not isinstance(override, dict):
        return merged
    key_map = {
        "methodology": "methodology",
        "tactics": "tactics",
        "views": "views",
        "operations": "operations",
        "risk_rules": "risk_rules",
        "style_constraints": "style_constraints",
        "交易方法论": "methodology",
        "交易手法": "tactics",
        "观点": "views",
        "交易操作": "operations",
        "风控规则": "risk_rules",
        "风格约束": "style_constraints",
    }
    for raw_key, value in override.items():
        key = key_map.get(str(raw_key).strip())
        if key:
            merged[key] = str(value or "").strip()
    return merged


def _build_memory_portrait_markdown(profile: Dict, portrait: Dict) -> str:
    name = (profile.get("name") or "").strip() or "未命名人格"
    profile_id = (profile.get("id") or "").strip()
    description = (profile.get("description") or "").strip()
    source_blogger = (profile.get("source_blogger") or "").strip()
    updated_at = (portrait.get("updated_at") or "").strip()
    updated_by = (portrait.get("updated_by") or "").strip()

    lines: List[str] = [
        f"# 人物侧写：{name}",
        "",
        f"- 人格ID: {profile_id or '--'}",
        f"- 导出时间(UTC): {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 侧写更新时间: {updated_at or '--'}",
        f"- 侧写来源: {updated_by or '--'}",
    ]
    if description:
        lines.append(f"- 人格说明: {description}")
    if source_blogger:
        lines.append(f"- 来源博主: {source_blogger}")
    lines.append("")

    for title, body in _memory_portrait_sections(portrait):
        lines.append(f"## {title}")
        lines.append(body or "（空）")
        lines.append("")

    evidence_refs = portrait.get("evidence_refs") if isinstance(portrait.get("evidence_refs"), list) else []
    if evidence_refs:
        lines.append("## 证据引用")
        for idx, item in enumerate(evidence_refs, start=1):
            if not isinstance(item, dict):
                continue
            rid = str(item.get("resource_id") or "--").strip() or "--"
            topic = str(item.get("topic") or "--").strip() or "--"
            quote = str(item.get("quote") or "").strip()
            source_time = str(item.get("source_time") or "").strip()
            source_type = str(item.get("source_time_type") or "").strip()
            timecode = str(item.get("timecode") or "").strip()
            lines.append(f"{idx}. [{topic}] 资源: {rid}")
            if quote:
                lines.append(f"   - 引文: {quote}")
            if source_time:
                lines.append(f"   - 时间: {source_time} ({source_type or 'unknown'})")
            if timecode:
                lines.append(f"   - 片段: {timecode}")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def _resolve_portrait_pdf_font() -> Tuple[str, Optional[Path]]:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/msyh.ttf"),
        Path("C:/Windows/Fonts/msyhbd.ttc"),
        Path("C:/Windows/Fonts/msyhbd.ttf"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return "PortraitFont", path
    return "Helvetica", None


def _render_memory_portrait_docx(profile: Dict, portrait: Dict) -> bytes:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("缺少 Word 导出依赖 python-docx") from exc

    doc = Document()
    name = (profile.get("name") or "").strip() or "未命名人格"
    doc.add_heading(f"人物侧写：{name}", level=1)
    doc.add_paragraph(f"人格ID：{(profile.get('id') or '').strip() or '--'}")
    doc.add_paragraph(f"导出时间(UTC)：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"侧写更新时间：{(portrait.get('updated_at') or '').strip() or '--'}")
    doc.add_paragraph(f"侧写来源：{(portrait.get('updated_by') or '').strip() or '--'}")
    if (profile.get("description") or "").strip():
        doc.add_paragraph(f"人格说明：{(profile.get('description') or '').strip()}")
    if (profile.get("source_blogger") or "").strip():
        doc.add_paragraph(f"来源博主：{(profile.get('source_blogger') or '').strip()}")
    doc.add_paragraph("")

    for title, body in _memory_portrait_sections(portrait):
        doc.add_heading(title, level=2)
        for line in (body or "（空）").splitlines():
            doc.add_paragraph(line)

    evidence_refs = portrait.get("evidence_refs") if isinstance(portrait.get("evidence_refs"), list) else []
    if evidence_refs:
        doc.add_heading("证据引用", level=2)
        for idx, item in enumerate(evidence_refs, start=1):
            if not isinstance(item, dict):
                continue
            rid = str(item.get("resource_id") or "--").strip() or "--"
            topic = str(item.get("topic") or "--").strip() or "--"
            quote = str(item.get("quote") or "").strip()
            source_time = str(item.get("source_time") or "").strip()
            source_type = str(item.get("source_time_type") or "").strip()
            timecode = str(item.get("timecode") or "").strip()
            doc.add_paragraph(f"{idx}. [{topic}] 资源: {rid}")
            if quote:
                doc.add_paragraph(f"引文: {quote}")
            if source_time:
                doc.add_paragraph(f"时间: {source_time} ({source_type or 'unknown'})")
            if timecode:
                doc.add_paragraph(f"片段: {timecode}")

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def _render_memory_portrait_pdf(profile: Dict, portrait: Dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise RuntimeError("缺少 PDF 导出依赖 reportlab") from exc

    font_name, font_path = _resolve_portrait_pdf_font()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
        except Exception:
            font_name = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("portrait_title", parent=styles["Heading1"], fontName=font_name, fontSize=16, leading=20, spaceAfter=10)
    h2_style = ParagraphStyle("portrait_h2", parent=styles["Heading2"], fontName=font_name, fontSize=13, leading=18, spaceAfter=6)
    body_style = ParagraphStyle("portrait_body", parent=styles["Normal"], fontName=font_name, fontSize=10.5, leading=16, spaceAfter=3)
    meta_style = ParagraphStyle("portrait_meta", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=13, textColor="#5a6570", spaceAfter=2)

    def _safe(text: str) -> str:
        return str(text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story: List[Any] = []
    name = (profile.get("name") or "").strip() or "未命名人格"
    story.append(Paragraph(_safe(f"人物侧写：{name}"), title_style))
    story.append(Paragraph(_safe(f"人格ID：{(profile.get('id') or '').strip() or '--'}"), meta_style))
    story.append(Paragraph(_safe(f"导出时间(UTC)：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}"), meta_style))
    story.append(Paragraph(_safe(f"侧写更新时间：{(portrait.get('updated_at') or '').strip() or '--'}"), meta_style))
    story.append(Paragraph(_safe(f"侧写来源：{(portrait.get('updated_by') or '').strip() or '--'}"), meta_style))
    if (profile.get("description") or "").strip():
        story.append(Paragraph(_safe(f"人格说明：{(profile.get('description') or '').strip()}"), meta_style))
    if (profile.get("source_blogger") or "").strip():
        story.append(Paragraph(_safe(f"来源博主：{(profile.get('source_blogger') or '').strip()}"), meta_style))
    story.append(Spacer(1, 8))

    for title, body in _memory_portrait_sections(portrait):
        story.append(Paragraph(_safe(title), h2_style))
        text = body or "（空）"
        for line in text.splitlines():
            story.append(Paragraph(_safe(line), body_style))
        story.append(Spacer(1, 4))

    evidence_refs = portrait.get("evidence_refs") if isinstance(portrait.get("evidence_refs"), list) else []
    if evidence_refs:
        story.append(Paragraph(_safe("证据引用"), h2_style))
        for idx, item in enumerate(evidence_refs, start=1):
            if not isinstance(item, dict):
                continue
            rid = str(item.get("resource_id") or "--").strip() or "--"
            topic = str(item.get("topic") or "--").strip() or "--"
            quote = str(item.get("quote") or "").strip()
            source_time = str(item.get("source_time") or "").strip()
            source_type = str(item.get("source_time_type") or "").strip()
            timecode = str(item.get("timecode") or "").strip()
            story.append(Paragraph(_safe(f"{idx}. [{topic}] 资源: {rid}"), body_style))
            if quote:
                story.append(Paragraph(_safe(f"引文: {quote}"), body_style))
            if source_time:
                story.append(Paragraph(_safe(f"时间: {source_time} ({source_type or 'unknown'})"), body_style))
            if timecode:
                story.append(Paragraph(_safe(f"片段: {timecode}"), body_style))
        story.append(Spacer(1, 4))

    out = BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36, title=name, author="strategy_watch")
    doc.build(story)
    return out.getvalue()


def _bind_resources_to_profile(
    profile_id: str,
    resource_ids: List[str],
    bind_source: str = "manual",
    group_id: str = "",
) -> Dict[str, Any]:
    target_ids = [str(x).strip() for x in (resource_ids or []) if str(x).strip()]
    target_ids = list(dict.fromkeys(target_ids))

    resources_payload = _load_resources_payload()
    resource_map = {item.get("id"): item for item in resources_payload.get("resources", [])}
    links = _load_memory_links()
    existing = {(x.get("profile_id"), x.get("resource_id")) for x in links}

    added = 0
    skipped: List[Dict] = []
    for rid in target_ids:
        item = resource_map.get(rid)
        if not item:
            skipped.append({"resource_id": rid, "reason": "资源不存在"})
            continue
        key = (profile_id, rid)
        if key in existing:
            skipped.append({"resource_id": rid, "reason": "已在人格长期记忆中"})
            continue
        links.append(
            {
                "profile_id": profile_id,
                "resource_id": rid,
                "bind_source": bind_source,
                "group_id": group_id,
                "added_at": _now_iso(),
            }
        )
        existing.add(key)
        added += 1

    _save_memory_links(links)
    return {"added": added, "skipped": skipped}


def _upsert_profile_subscription(profile: Dict, group_id: str, sync_now: bool = False) -> bool:
    group_id = (group_id or "").strip()
    if not group_id:
        return False
    subs = profile.get("group_subscriptions") if isinstance(profile.get("group_subscriptions"), list) else []
    now = _now_iso()
    for sub in subs:
        if not isinstance(sub, dict):
            continue
        if (sub.get("group_id") or "").strip() == group_id:
            if sync_now:
                sub["last_sync_at"] = now
            profile["group_subscriptions"] = subs
            return False
    subs.append(
        {
            "group_id": group_id,
            "first_bound_at": now,
            "last_sync_at": now if sync_now else "",
        }
    )
    profile["group_subscriptions"] = subs
    return True


def _sync_profile_group_resources(profile: Dict, group_id: str) -> Dict[str, Any]:
    gid = (group_id or "").strip()
    if not gid:
        return {"added": 0, "skipped": [{"group_id": gid, "reason": "缺少 group_id"}]}
    resources_payload = _load_resources_payload()
    group_resource_ids = [
        (item.get("id") or "").strip()
        for item in resources_payload.get("resources", [])
        if (item.get("group_id") or "").strip() == gid and (item.get("id") or "").strip()
    ]
    result = _bind_resources_to_profile(
        profile_id=profile.get("id") or "",
        resource_ids=group_resource_ids,
        bind_source="group",
        group_id=gid,
    )
    _upsert_profile_subscription(profile, gid, sync_now=True)
    return result


def _memory_profile_view(profile: Dict, links: List[Dict], portraits: Dict[str, Dict]) -> Dict:
    pid = (profile.get("id") or "").strip()
    linked_ids = [item.get("resource_id") for item in links if item.get("profile_id") == pid]
    portrait = portraits.get(pid) or _default_memory_portrait(pid)
    subscription_ids = []
    for sub in (profile.get("group_subscriptions") or []):
        if isinstance(sub, dict) and (sub.get("group_id") or "").strip():
            subscription_ids.append((sub.get("group_id") or "").strip())
    return {
        **profile,
        "linked_resource_count": len(set(linked_ids)),
        "group_subscription_count": len(set(subscription_ids)),
        "portrait_updated_at": portrait.get("updated_at") or "",
    }


def _extract_json_block(text: str) -> Dict:
    if not text:
        return {}
    cleaned = text.strip()
    import re

    # Try plain JSON first.
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        pass

    # Handle fenced code block content like ```json ... ```.
    if cleaned.startswith("```"):
        lines = cleaned.splitlines()
        if len(lines) >= 3 and lines[0].startswith("```") and lines[-1].strip().startswith("```"):
            body = "\n".join(lines[1:-1]).strip()
            try:
                parsed = json.loads(body)
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                pass

    fenced = re.search(r"```json\\s*({[\\s\\S]*?})\\s*```", text, re.IGNORECASE)
    if fenced:
        try:
            return json.loads(fenced.group(1))
        except Exception:
            return {}

    generic = re.search(r"({[\\s\\S]*})", text)
    if generic:
        try:
            return json.loads(generic.group(1))
        except Exception:
            return {}

    return {}


def _normalize_portrait_draft_payload(parsed: Dict, reply: str = "") -> Dict:
    if not isinstance(parsed, dict):
        parsed = {}

    key_alias = {
        "交易方法论": "methodology",
        "交易手法": "tactics",
        "观点": "views",
        "交易操作": "operations",
        "风控规则": "risk_rules",
        "风格约束": "style_constraints",
        "evidence_refs": "evidence_refs",
    }
    normalized: Dict[str, Any] = {}
    for raw_key, value in parsed.items():
        key = key_alias.get(str(raw_key).strip(), str(raw_key).strip())
        normalized[key] = value

    # Some models put the full JSON object string inside "methodology".
    nested_text = normalized.get("methodology")
    has_other_sections = any(
        str(normalized.get(k) or "").strip()
        for k in ("tactics", "views", "operations", "risk_rules", "style_constraints")
    )
    if isinstance(nested_text, str) and nested_text.strip() and not has_other_sections:
        nested = _extract_json_block(nested_text)
        if isinstance(nested, dict) and nested:
            nested_norm = _normalize_portrait_draft_payload(nested, "")
            if any(
                str(nested_norm.get(k) or "").strip()
                for k in ("tactics", "views", "operations", "risk_rules", "style_constraints")
            ):
                normalized = nested_norm

    if not normalized and reply:
        fallback = _extract_json_block(reply)
        if isinstance(fallback, dict) and fallback:
            normalized = _normalize_portrait_draft_payload(fallback, "")

    return normalized


def _build_visualization_prompt(
    goal: str,
    view_type: str = "",
    existing_widgets: Optional[List[Dict[str, Any]]] = None,
) -> str:
    view_hint = view_type or DEFAULT_STRATEGY_VIEW
    existing = existing_widgets if isinstance(existing_widgets, list) else []
    existing_json = json.dumps(existing[:20], ensure_ascii=False, indent=2)
    return f"""
你是A股可视化规划助理。请根据目标生成“看盘视图配置”的JSON，只返回JSON，不要附加解释。

目标：{goal}
偏好视图类型：{view_hint}
当前已存在 widgets（不可删除、不可覆盖、不可修改）：
```json
{existing_json}
```

允许的widget.type：
- market_sentiment_chart（市场情绪单图）
- index_kline（指数K线）
- sector_kline（板块K线）
- stock_kline（个股K线）
- market_volume（市场量能对比）

market_sentiment_chart.params 约束：
- chart_key 仅允许：
  red_ratio_and_amount / limit_up_count / ground_ceiling_count / continuous_limit_up / change_distribution
- days_back 为10~240整数，默认30

index_kline.params 约束：
- index_name 例如“上证指数”
- days_range 为20~500整数，默认60

sector_kline.params 约束：
- sector_name 例如“半导体”
- days_range 为20~500整数，默认60

stock_kline.params 约束：
- stock_code 为6位数字字符串，例如“600519”
- days 为20~500整数，默认120

layout 字段要求（支持手动拖拽布局）：
- layout.x / layout.y / layout.w / layout.h 均为数字（像素）
- 建议 w>=300, h>=240

输出规则（必须遵守）：
1) widgets 仅返回“新增图表”，不要重复当前已存在图表（按 type+params 判重）。
2) 不要删除、重写或替换现有图表。
3) name / description / view_type 如无需修改请原样返回空字符串或当前值。

JSON结构示例：
{{
  "name": "策略名称",
  "description": "策略说明",
  "view_type": "basic",
  "widgets": [
    {{
      "id": "sentiment-red-ratio",
      "type": "market_sentiment_chart",
      "title": "红盘率与成交额",
      "params": {{ "chart_key": "red_ratio_and_amount", "days_back": 30 }},
      "layout": {{ "x": 0, "y": 0, "w": 640, "h": 420 }}
    }},
    {{
      "id": "index-kline",
      "type": "index_kline",
      "title": "上证指数 K线",
      "params": {{ "index_name": "上证指数", "days_range": 60 }},
      "layout": {{ "x": 0, "y": 434, "w": 720, "h": 520 }}
    }}
  ]
}}
""".strip()


def _safe_number(value: Any, default: float) -> float:
    try:
        num = float(value)
        if num != num:  # NaN
            return float(default)
        return num
    except Exception:
        return float(default)


def _normalize_strategy_widget_for_merge(widget: Dict[str, Any], index: int = 0) -> Optional[Dict[str, Any]]:
    if not isinstance(widget, dict):
        return None

    widget_type = str(widget.get("type") or "").strip()
    if not widget_type:
        return None

    widget_id = str(widget.get("id") or "").strip() or _gen_id("widget")
    params = widget.get("params") if isinstance(widget.get("params"), dict) else {}
    title = str(widget.get("title") or "").strip() or widget_type
    layout_src = widget.get("layout") if isinstance(widget.get("layout"), dict) else {}

    default_w = 620.0 if widget_type in {"index_kline", "sector_kline", "stock_kline"} else 500.0
    default_h = 520.0 if widget_type in {"index_kline", "sector_kline", "stock_kline"} else 420.0
    layout = {
        "x": max(0.0, _safe_number(layout_src.get("x"), 0.0)),
        "y": max(0.0, _safe_number(layout_src.get("y"), float(index * 320))),
        "w": max(300.0, _safe_number(layout_src.get("w"), default_w)),
        "h": max(240.0, _safe_number(layout_src.get("h"), default_h)),
    }

    return {
        "id": widget_id,
        "type": widget_type,
        "title": title,
        "params": params,
        "layout": layout,
    }


def _widget_merge_signature(widget: Dict[str, Any]) -> str:
    if not isinstance(widget, dict):
        return ""
    widget_type = str(widget.get("type") or "").strip()
    params = widget.get("params") if isinstance(widget.get("params"), dict) else {}
    params_json = json.dumps(params, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return f"{widget_type}|{params_json}"


def _merge_strategy_widgets(existing_widgets: List[Dict[str, Any]], new_widgets: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    existing = existing_widgets if isinstance(existing_widgets, list) else []
    incoming = new_widgets if isinstance(new_widgets, list) else []

    merged: List[Dict[str, Any]] = []
    used_ids = set()
    seen_signatures = set()

    def _append_one(raw_widget: Dict[str, Any], index_hint: int) -> None:
        normalized = _normalize_strategy_widget_for_merge(raw_widget, index=index_hint)
        if not normalized:
            return

        wid = str(normalized.get("id") or "").strip() or _gen_id("widget")
        while wid in used_ids:
            wid = f"{wid}_{len(used_ids) + 1}"
        normalized["id"] = wid

        sign = _widget_merge_signature(normalized)
        if sign and sign in seen_signatures:
            return

        used_ids.add(wid)
        if sign:
            seen_signatures.add(sign)
        merged.append(normalized)

    for item in existing:
        _append_one(item, len(merged))
    for item in incoming:
        _append_one(item, len(merged))

    return merged


def _gen_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:12]}"


def _find_by_id(items: List[Dict], item_id: str) -> Dict:
    return next((item for item in items if item.get("id") == item_id), None)


def _read_markdown_content(resource: Dict) -> str:
    md_rel = resource.get("markdown_relpath") or ""
    if not md_rel:
        return ""
    md_path = BASE_DIR / md_rel
    if not md_path.exists():
        return ""
    return md_path.read_text(encoding="utf-8", errors="ignore")


def _strip_likely_page_headings(text: str) -> str:
    lines = text.splitlines()
    kept = []
    for raw in lines:
        line = raw.strip()
        compact = re.sub(r"\s+", "", line)
        if line.startswith("##") and re.search(r"\d+", line) and len(compact) <= 20:
            continue
        kept.append(raw)
    return "\n".join(kept)


def _markdown_has_meaningful_content(md_content: str, min_chars: int = 20) -> bool:
    if not md_content:
        return False
    marker = "\n## Content\n"
    content = md_content.split(marker, 1)[1] if marker in md_content else md_content
    if not content.strip():
        return False
    content = _strip_likely_page_headings(content)
    compact = re.sub(r"\s+", "", content)
    meaningful = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9]", "", compact)
    return len(meaningful) >= min_chars


def _build_strategy_reference_markdown(strategy: Dict) -> str:
    sid = (strategy.get("id") or "").strip()
    name = (strategy.get("name") or "").strip() or sid
    description = (strategy.get("description") or "").strip()
    view_type = (strategy.get("view_type") or DEFAULT_STRATEGY_VIEW).strip() or DEFAULT_STRATEGY_VIEW
    updated_at = (strategy.get("updated_at") or "").strip()
    config = strategy.get("config") if isinstance(strategy.get("config"), dict) else {}
    widgets = config.get("widgets") if isinstance(config, dict) else None

    lines: List[str] = [
        f"# 策略：{name}",
        "",
        "## 展示内容",
        f"- 策略ID：{sid}",
        f"- 视图类型：{view_type}",
        f"- 更新时间：{updated_at or '(unknown)'}",
    ]
    if description:
        lines.append(f"- 说明：{description}")
    else:
        lines.append("- 说明：暂无")

    lines.extend(["", "## 策略逻辑"])
    if isinstance(widgets, list) and widgets:
        for idx, widget in enumerate(widgets, start=1):
            w_type = (widget.get("type") or "").strip() if isinstance(widget, dict) else ""
            w_title = (widget.get("title") or "").strip() if isinstance(widget, dict) else ""
            w_id = (widget.get("id") or "").strip() if isinstance(widget, dict) else ""
            params = widget.get("params") if isinstance(widget, dict) else {}
            params_json = json.dumps(params if isinstance(params, dict) else {}, ensure_ascii=False, separators=(",", ":"))
            lines.append(
                f"{idx}. {w_title or w_type or w_id or 'widget'}"
                f"（type={w_type or 'unknown'}, id={w_id or 'unknown'}, params={params_json}）"
            )
    else:
        lines.append("1. 当前策略未配置 widgets，默认由 view_type 决定展示内容。")

    lines.extend(["", "## 原始配置", "```json", json.dumps(config, ensure_ascii=False, indent=2), "```"])
    return "\n".join(lines).strip()


def _build_resource_context(resource_ids: List[str]) -> Tuple[str, List[str], List[Dict]]:
    payload = _load_resources_payload()
    resources = payload.get("resources", [])
    active_resource_id = payload.get("active_resource_id") or ""
    resource_map = {r.get("id"): r for r in resources}
    strategies_payload = _load_strategies_payload()
    strategies = strategies_payload.get("strategies", [])
    strategy_map = {s.get("id"): s for s in strategies if s.get("id")}

    # Strict mode for chat: if user does not explicitly select resources,
    # do not auto-inject active/all resources into model context.
    selected_ids = list(resource_ids or [])
    if not selected_ids:
        return "", [], []
    total_len = 0
    used_ids: List[str] = []
    skipped: List[Dict] = []
    blocks: List[str] = []

    for rid in selected_ids:
        if isinstance(rid, str) and rid.startswith("strategy_ctx:"):
            strategy_id = rid.split(":", 1)[1].strip()
            strategy = strategy_map.get(strategy_id)
            if not strategy:
                skipped.append({"resource_id": rid, "reason": "策略不存在"})
                continue

            strategy_md = _build_strategy_reference_markdown(strategy)
            snippet = strategy_md[:MAX_RESOURCE_CHARS]
            projected = total_len + len(snippet)
            if projected > MAX_CONTEXT_CHARS:
                remain = MAX_CONTEXT_CHARS - total_len
                if remain <= 0:
                    break
                snippet = snippet[:remain]
                projected = total_len + len(snippet)

            strategy_name = (strategy.get("name") or strategy_id).strip()
            blocks.append(f"## Strategy Reference: {strategy_name}\n\n{snippet}")
            used_ids.append(rid)
            total_len = projected
            if total_len >= MAX_CONTEXT_CHARS:
                break
            continue

        item = resource_map.get(rid)
        if not item:
            skipped.append({"resource_id": rid, "reason": "资源不存在"})
            continue
        if (item.get("status") or "") != "ok":
            skipped.append({"resource_id": rid, "reason": f"资源状态异常: {item.get('status') or 'unknown'}"})
            continue

        md_content = _read_markdown_content(item).strip()
        if not md_content:
            skipped.append({"resource_id": rid, "reason": "markdown 文件不存在或为空"})
            continue
        if not _markdown_has_meaningful_content(md_content):
            skipped.append({"resource_id": rid, "reason": "markdown 信息量不足（疑似空壳转换结果）"})
            continue

        snippet = md_content[:MAX_RESOURCE_CHARS]
        projected = total_len + len(snippet)
        if projected > MAX_CONTEXT_CHARS:
            remain = MAX_CONTEXT_CHARS - total_len
            if remain <= 0:
                break
            snippet = snippet[:remain]
            projected = total_len + len(snippet)

        blocks.append(f"## Resource: {item.get('original_name', rid)}\n\n{snippet}")
        used_ids.append(rid)
        total_len = projected
        if total_len >= MAX_CONTEXT_CHARS:
            break

    return ("\n\n---\n\n".join(blocks)).strip(), used_ids, skipped


def _build_strategy_context(strategy_id: str = "") -> Tuple[str, str]:
    payload = _load_strategies_payload()
    strategies = payload.get("strategies", [])
    strategy_map = {item.get("id"): item for item in strategies if item.get("id")}

    requested_id = (strategy_id or "").strip()
    # Strict mode for strategy context:
    # only inject strategy context when user explicitly provides strategy_id.
    if not requested_id:
        return "", ""

    target = strategy_map.get(requested_id)
    if not target:
        return "", ""

    config = target.get("config") if isinstance(target.get("config"), dict) else {}
    widgets = config.get("widgets") if isinstance(config, dict) else None
    widget_count = len(widgets) if isinstance(widgets, list) else 0
    config_json = json.dumps(config, ensure_ascii=False, separators=(",", ":"))

    lines: List[str] = [
        "策略与看盘上下文（用户显式选择）：",
        f"- strategy_id: {target.get('id') or ''}",
        f"- strategy_count: {len(strategies)}",
    ]
    lines.extend(
        [
            "- 当前策略：",  # keep existing downstream parser compatibility
            f"  - id: {target.get('id') or ''}",
            f"  - name: {target.get('name') or ''}",
            f"  - description: {target.get('description') or ''}",
            f"  - view_type: {target.get('view_type') or DEFAULT_STRATEGY_VIEW}",
            f"  - widget_count: {widget_count}",
            f"  - config_json: {config_json}",
        ]
    )

    return "\n".join(lines), requested_id


_ORCHESTRATION_TASK_AGENTS = {
    "pm_agent",
    "engineer_agent",
    "analyst_agent",
    "visualization_master_agent",
}


def _clip_text(value: Any, max_chars: int = 2200) -> str:
    text = str(value or "").strip()
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n...[truncated]"


def _extract_json_payload(text: str) -> Any:
    raw = str(text or "").strip()
    if not raw:
        return None

    candidates: List[str] = [raw]
    for block in re.findall(r"```json\s*([\s\S]*?)```", raw, flags=re.IGNORECASE):
        piece = block.strip()
        if piece:
            candidates.append(piece)
    for block in re.findall(r"```(?:[a-zA-Z0-9_-]+)?\s*([\s\S]*?)```", raw):
        piece = block.strip()
        if piece:
            candidates.append(piece)

    obj_start = raw.find("{")
    obj_end = raw.rfind("}")
    if obj_start >= 0 and obj_end > obj_start:
        candidates.append(raw[obj_start : obj_end + 1].strip())

    arr_start = raw.find("[")
    arr_end = raw.rfind("]")
    if arr_start >= 0 and arr_end > arr_start:
        candidates.append(raw[arr_start : arr_end + 1].strip())

    seen = set()
    for candidate in candidates:
        key = candidate.strip()
        if not key or key in seen:
            continue
        seen.add(key)
        try:
            return json.loads(key)
        except Exception:
            continue
    return None


def _parse_requirements_ready(pm_text: str) -> bool:
    payload = _extract_json_payload(pm_text)
    if isinstance(payload, dict):
        raw = payload.get("requirements_ready")
        if raw is None and isinstance(payload.get("requirements"), dict):
            raw = payload.get("requirements", {}).get("ready")
        if isinstance(raw, bool):
            return raw
        if isinstance(raw, (int, float)):
            return bool(raw)
        if isinstance(raw, str):
            lowered = raw.strip().lower()
            if lowered in {"true", "1", "yes", "ready", "ok"}:
                return True
            if lowered in {"false", "0", "no", "not_ready", "pending"}:
                return False

    text = str(pm_text or "")
    hits = re.findall(
        r"""["'`]?requirements_ready["'`]?\s*(?:[:=：]|->)\s*["'`]?(true|false|1|0|yes|no|ready|not_ready|pending|ok)["'`]?""",
        text,
        flags=re.IGNORECASE,
    )
    if hits:
        token = str(hits[-1]).strip().lower()
        return token in {"true", "1", "yes", "ready", "ok"}

    # Fallback: tolerate "requirements_ready true" (without separator)
    loose_hits = re.findall(
        r"""["'`]?requirements_ready["'`]?\s+(true|false|1|0|yes|no|ready|not_ready|pending|ok)""",
        text,
        flags=re.IGNORECASE,
    )
    if loose_hits:
        token = str(loose_hits[-1]).strip().lower()
        return token in {"true", "1", "yes", "ready", "ok"}

    return False


def _normalize_text_list(value: Any) -> List[str]:
    if isinstance(value, list):
        out: List[str] = []
        for item in value:
            text = str(item or "").strip()
            if text:
                out.append(text)
        return out
    if isinstance(value, str):
        text = value.strip()
        if not text:
            return []
        parts = re.split(r"(?:\r?\n|[；;])+", text)
        out = []
        for part in parts:
            line = re.sub(r"^\s*(?:[-*•]|\d+[.)、])\s*", "", part.strip())
            if line:
                out.append(line)
        return out
    return []


def _extract_named_list_lines(text: str, field_name: str) -> List[str]:
    raw = str(text or "")
    if not raw.strip():
        return []
    lines = raw.splitlines()
    start = -1
    inline_value = ""
    pattern = re.compile(rf"""^\s*["'`]?{re.escape(field_name)}["'`]?\s*[:：]\s*(.*)$""", flags=re.IGNORECASE)
    for idx, line in enumerate(lines):
        hit = pattern.match(line)
        if hit:
            start = idx
            inline_value = (hit.group(1) or "").strip()
            break
    if start < 0:
        return []

    out: List[str] = []
    if inline_value and inline_value not in {"[]", "null", "None"}:
        if inline_value.startswith("[") and inline_value.endswith("]"):
            try:
                parsed = json.loads(inline_value)
                return _normalize_text_list(parsed)
            except Exception:
                pass
        out.extend(_normalize_text_list(inline_value))

    next_key_pattern = re.compile(r"^\s*[a-zA-Z_][a-zA-Z0-9_]*\s*[:：]")
    for line in lines[start + 1 :]:
        if next_key_pattern.match(line):
            break
        item = re.sub(r"^\s*(?:[-*•]|\d+[.)、])\s*", "", line.strip())
        if item:
            out.append(item)
    return out


def _extract_pm_requirements_package(pm_text: str) -> Dict[str, Any]:
    parsed: Dict[str, Any] = {
        "requirements_ready": _parse_requirements_ready(pm_text),
        "block_reason": "",
        "next_questions": [],
        "unknowns": [],
        "raw_text": str(pm_text or ""),
    }
    payload = _extract_json_payload(pm_text)

    if isinstance(payload, dict):
        parsed["next_questions"] = _normalize_text_list(payload.get("next_questions"))
        parsed["unknowns"] = _normalize_text_list(payload.get("unknowns"))
        block_reason = (
            payload.get("block_reason")
            or payload.get("blocking_reason")
            or payload.get("cannot_proceed_reason")
            or ""
        )
        parsed["block_reason"] = str(block_reason or "").strip()

    if not parsed["next_questions"]:
        parsed["next_questions"] = _extract_named_list_lines(pm_text, "next_questions")
    if not parsed["unknowns"]:
        parsed["unknowns"] = _extract_named_list_lines(pm_text, "unknowns")
    if not parsed["block_reason"]:
        hit = re.search(
            r"""["'`]?block_reason["'`]?\s*[:：]\s*(.+)""",
            str(pm_text or ""),
            flags=re.IGNORECASE,
        )
        if hit:
            parsed["block_reason"] = str(hit.group(1) or "").strip()
    if not parsed["block_reason"] and parsed["unknowns"]:
        parsed["block_reason"] = "关键约束未补齐：" + "；".join(parsed["unknowns"][:2])

    parsed["next_questions"] = parsed["next_questions"][:3]
    return parsed


def _extract_latest_pm_context(history_messages: List[Dict[str, Any]]) -> str:
    for item in reversed(history_messages or []):
        if not isinstance(item, dict):
            continue
        meta = item.get("orchestration_meta")
        if isinstance(meta, dict):
            pm_text = str(meta.get("pm_requirements_raw") or "").strip()
            if pm_text:
                return pm_text
    return ""


def _normalize_orchestration_task(item: Any, index: int) -> Optional[Dict[str, Any]]:
    if not isinstance(item, dict):
        return None
    task_id = str(item.get("task_id") or f"T{index}").strip() or f"T{index}"
    title = str(item.get("title") or f"任务{index}").strip() or f"任务{index}"
    agent = str(item.get("agent") or "engineer_agent").strip() or "engineer_agent"
    if agent not in _ORCHESTRATION_TASK_AGENTS:
        agent = "engineer_agent"

    params = item.get("params")
    if not isinstance(params, dict):
        params = {}
    depends_on = item.get("depends_on")
    if not isinstance(depends_on, list):
        depends_on = []

    return {
        "task_id": task_id,
        "title": title,
        "agent": agent,
        "description": str(item.get("description") or "").strip(),
        "params": params,
        "depends_on": [str(x).strip() for x in depends_on if str(x).strip()],
        "milestone": str(item.get("milestone") or "").strip(),
        "acceptance": str(item.get("acceptance") or "").strip(),
        "risk": str(item.get("risk") or "").strip(),
        "rollback": str(item.get("rollback") or "").strip(),
    }


def _parse_todo_tasks(plan_text: str) -> List[Dict[str, Any]]:
    payload = _extract_json_payload(plan_text)
    raw_tasks: List[Any] = []
    if isinstance(payload, list):
        raw_tasks = payload
    elif isinstance(payload, dict):
        for key in ("tasks", "todo_list", "todos", "todo", "items"):
            value = payload.get(key)
            if isinstance(value, list):
                raw_tasks = value
                break

    normalized: List[Dict[str, Any]] = []
    for idx, item in enumerate(raw_tasks, start=1):
        parsed = _normalize_orchestration_task(item, idx)
        if parsed:
            normalized.append(parsed)
    return normalized


def _infer_task_result_status(task_output: str, failed: bool = False) -> str:
    if failed:
        return "failed"
    text = str(task_output or "")
    hit = re.search(r"result\s*[:：]\s*(done|blocked|failed)", text, flags=re.IGNORECASE)
    if hit:
        return hit.group(1).lower()
    return "done"


def _run_strategy_edit_orchestration(
    *,
    llm_user_content: str,
    history_messages: List[Dict],
    resource_context: str,
    extra_context: str,
    model_name: str,
    temperature: float,
    trace_context: Dict[str, Any],
) -> Tuple[str, str, str, Dict[str, Any], str, Dict[str, Any]]:
    stage_history = list(history_messages or [])
    previous_pm_context = _extract_latest_pm_context(history_messages)
    resolved_model = _canonical_model_name(model_name or (os.getenv("OPENAI_MODEL") or "gpt-4o-mini"))
    provider = "orchestration"
    usage: Dict[str, Any] = {}
    flow_error = ""
    orchestration_meta: Dict[str, Any] = {}

    def _run_stage(
        stage_name: str,
        agent_name: str,
        prompt_text: str,
        extra_payload: Optional[Dict[str, Any]] = None,
    ) -> Tuple[str, str]:
        nonlocal resolved_model, provider, usage, stage_history
        payload = {
            **(trace_context or {}),
            "stage_name": stage_name,
            "agent_name": agent_name,
        }
        if isinstance(extra_payload, dict):
            payload.update(extra_payload)
        _append_agent_log("orchestration_stage_start", payload)
        local_trace = dict(trace_context or {})
        local_trace.update(
            {
                "orchestration_stage": stage_name,
                "agent_name": agent_name,
            }
        )
        try:
            stage_text, stage_model, stage_provider, stage_usage = chat_with_agent(
                agent_name=agent_name,
                user_content=prompt_text,
                history_messages=stage_history[-MAX_HISTORY_MESSAGES:],
                resource_context=resource_context,
                extra_context=extra_context,
                model_name=model_name,
                temperature=temperature,
                trace_context=local_trace,
            )
            if stage_model:
                resolved_model = stage_model
            if stage_provider:
                provider = stage_provider
            if isinstance(stage_usage, dict) and stage_usage:
                usage = stage_usage
            stage_history.append({"role": "user", "content": prompt_text})
            stage_history.append({"role": "assistant", "content": stage_text})
            _append_agent_log(
                "orchestration_stage_done",
                {
                    **payload,
                    "output_chars": len(stage_text or ""),
                },
            )
            return stage_text, ""
        except Exception as exc:
            err = str(exc)
            _append_agent_log(
                "orchestration_stage_error",
                {
                    **payload,
                    "error": err,
                },
                level="error",
            )
            return "", err

    pm_prompt = (
        "你现在处于【需求澄清阶段】。\n"
        "请先判断用户需求是否足够清晰可执行，并必须输出 requirements_ready 字段。\n"
        "若不清晰，请给出 next_questions（最多3条）。\n\n"
        "若 requirements_ready=false，必须给出 block_reason（不能进入架构拆解的原因）。\n\n"
        + (
            "上一轮 PM 需求包（用于延续上下文）：\n"
            f"{_clip_text(previous_pm_context, max_chars=2600)}\n\n"
            if previous_pm_context
            else ""
        )
        +
        "用户原始需求：\n"
        f"{llm_user_content}"
    )
    pm_text, pm_error = _run_stage("pm_discovery", "pm_agent", pm_prompt)
    if pm_error:
        flow_error = pm_error
        return (
            f"策略编排在 PM 阶段失败：{pm_error}",
            resolved_model,
            provider,
            usage,
            flow_error,
            orchestration_meta,
        )

    requirements_pkg = _extract_pm_requirements_package(pm_text)
    requirements_ready = bool(requirements_pkg.get("requirements_ready"))
    orchestration_meta = {
        "requirements_ready": requirements_ready,
        "pm_requirements_raw": pm_text,
        "pm_requirements": {
            "block_reason": requirements_pkg.get("block_reason") or "",
            "next_questions": requirements_pkg.get("next_questions") or [],
            "unknowns": requirements_pkg.get("unknowns") or [],
        },
    }
    _append_agent_log(
        "orchestration_requirements_ready_parsed",
        {
            **(trace_context or {}),
            "requirements_ready": requirements_ready,
            "block_reason": requirements_pkg.get("block_reason") or "",
            "next_questions": requirements_pkg.get("next_questions") or [],
            "pm_output_preview": _clip_text(pm_text, max_chars=600),
        },
    )
    if not requirements_ready:
        reason = (requirements_pkg.get("block_reason") or "").strip() or "关键约束信息不完整。"
        questions = requirements_pkg.get("next_questions") or []
        lines = [
            "【PM需求澄清】",
            f"当前无法进入架构拆解：{reason}",
            "",
            "请补充以下最小必要信息：",
        ]
        if questions:
            for idx, q in enumerate(questions, start=1):
                lines.append(f"{idx}. {q}")
        else:
            lines.append("1. 请补充目标范围、输入输出和验收标准。")
        lines.append("")
        lines.append("补充后我会自动转交 architect_agent 进行任务分解与执行编排。")
        waiting_text = "\n".join(lines).strip()
        return waiting_text, resolved_model, provider, usage, "", orchestration_meta

    plan_prompt = (
        "你现在处于【任务分解阶段】。\n"
        "请基于 PM 输出生成可执行 ToDo，并严格返回 JSON（仅 JSON，不要额外解释）。\n"
        "JSON 结构：\n"
        "{\n"
        "  \"tasks\": [\n"
        "    {\n"
        "      \"task_id\": \"T1\",\n"
        "      \"title\": \"...\",\n"
        "      \"agent\": \"pm_agent|engineer_agent|analyst_agent|visualization_master_agent\",\n"
        "      \"description\": \"...\",\n"
        "      \"params\": {},\n"
        "      \"depends_on\": [],\n"
        "      \"milestone\": \"...\",\n"
        "      \"acceptance\": \"...\",\n"
        "      \"risk\": \"...\",\n"
        "      \"rollback\": \"...\"\n"
        "    }\n"
        "  ]\n"
        "}\n\n"
        "PM 输出如下：\n"
        f"{pm_text}"
    )
    plan_text, plan_error = _run_stage("architect_plan", "architect_agent", plan_prompt)
    if plan_error:
        flow_error = plan_error
        return (
            f"策略编排在 Architect 任务分解阶段失败：{plan_error}",
            resolved_model,
            provider,
            usage,
            flow_error,
            orchestration_meta,
        )

    tasks = _parse_todo_tasks(plan_text)
    if not tasks:
        no_task_text = (
            "【阶段2：Architect任务分解】\n"
            f"{plan_text}\n\n"
            "未解析到可执行 tasks，请让 Architect 按约定 JSON 结构返回 ToDo。"
        )
        return no_task_text, resolved_model, provider, usage, "", orchestration_meta

    task_results: List[Dict[str, Any]] = []
    for idx, task in enumerate(tasks, start=1):
        task_prompt = (
            "你收到一条来自 architect_agent 的执行任务，请严格按任务要求完成。\n"
            "请输出你的标准结果格式，并包含 result: done|blocked|failed。\n\n"
            f"任务序号：{idx}/{len(tasks)}\n"
            "任务定义(JSON)：\n"
            f"{json.dumps(task, ensure_ascii=False, indent=2)}"
        )
        task_text, task_error = _run_stage(
            f"task_execute_{task.get('task_id') or idx}",
            task.get("agent") or "engineer_agent",
            task_prompt,
            {
                "task_id": task.get("task_id"),
                "task_title": task.get("title"),
            },
        )
        status = _infer_task_result_status(task_text, failed=bool(task_error))
        task_results.append(
            {
                "task": task,
                "status": status,
                "error": task_error,
                "output": task_text if not task_error else f"执行失败：{task_error}",
            }
        )

    summary_prompt = (
        "你现在处于【执行总结阶段】。\n"
        "请根据 PM 澄清、ToDo 计划和各子任务执行结果，给出最终可执行总结。\n"
        "输出需包含：MVP完成情况、已完成任务、阻塞任务与风险、下一步行动。\n\n"
        "PM 输出：\n"
        f"{pm_text}\n\n"
        "ToDo(JSON)：\n"
        f"{json.dumps(tasks, ensure_ascii=False, indent=2)}\n\n"
        "任务执行结果(JSON)：\n"
        f"{json.dumps(task_results, ensure_ascii=False, indent=2)}"
    )
    summary_text, summary_error = _run_stage("architect_summary", "architect_agent", summary_prompt)
    if summary_error:
        flow_error = summary_error
        summary_text = f"Architect 总结阶段失败：{summary_error}"

    lines: List[str] = []
    lines.append("## 阶段1：PM需求澄清")
    lines.append(_clip_text(pm_text))
    lines.append("")
    lines.append("## 阶段2：Architect ToDo")
    lines.append("```json")
    lines.append(json.dumps(tasks, ensure_ascii=False, indent=2))
    lines.append("```")
    lines.append("")
    lines.append("## 阶段3：按ToDo逐步执行")
    for item in task_results:
        task = item.get("task") or {}
        lines.append(
            f"- {task.get('task_id', '')} | {task.get('title', '')} | "
            f"agent={task.get('agent', '')} | status={item.get('status', 'unknown')}"
        )
        lines.append(_clip_text(item.get("output"), max_chars=1200))
    lines.append("")
    lines.append("## 阶段4：Architect总结")
    lines.append(_clip_text(summary_text))
    lines.append("")
    if flow_error:
        lines.append(f"流程警告：{flow_error}")

    return "\n".join(lines).strip(), resolved_model, provider, usage, flow_error, orchestration_meta


def _resolve_strategy_name(strategy_id: str) -> str:
    sid = (strategy_id or "").strip()
    if not sid:
        return ""
    payload = _load_strategies_payload()
    strategies = payload.get("strategies", [])
    target = _find_by_id(strategies, sid)
    return (target.get("name") or "").strip() if target else ""


def _is_relpath_referenced(resources: List[Dict], rel_path: str, exclude_resource_id: str = "") -> bool:
    target = (rel_path or "").strip()
    if not target:
        return False
    for item in resources:
        if exclude_resource_id and item.get("id") == exclude_resource_id:
            continue
        if (item.get("source_relpath") or "").strip() == target:
            return True
        if (item.get("markdown_relpath") or "").strip() == target:
            return True
    return False


def _register_crawled_resources(crawl_results: List[Dict]) -> List[str]:
    created_ids: List[str] = []
    payload = _load_resources_payload()
    resources = payload.get("resources", [])
    crawl_group_id = "group_crawled"
    crawl_group_name = "爬虫抓取"

    for item in crawl_results:
        if item.get("status") != "ok" or not item.get("markdown_path"):
            continue

        md_path = Path(item["markdown_path"]).resolve()
        if not md_path.exists():
            continue

        rid = _gen_id("res")
        rel_markdown = md_path.relative_to(BASE_DIR).as_posix()
        published_at = _normalize_datetime_to_iso(item.get("published_at") or "")
        content_time = _normalize_date_only(item.get("published_at") or "")
        record = {
            "id": rid,
            "original_name": item.get("title") or item.get("url"),
            "stored_name": md_path.name,
            "extension": ".wx",
            "size_bytes": md_path.stat().st_size,
            "uploaded_at": _now_iso(),
            "status": "ok",
            "error": "",
            "source_type": "url",
            "source_url": item.get("url", ""),
            "source_relpath": "",
            "markdown_relpath": rel_markdown,
            "group_id": crawl_group_id,
            "group_name": crawl_group_name,
            "published_at": published_at,
            "content_time": content_time,
            "content_time_type": "published" if content_time else "unknown",
            "content_time_confidence": 0.95 if content_time else 0.0,
            "content_time_evidence": ("wechat:published_at" if content_time else ""),
        }
        resources.append(record)
        created_ids.append(rid)

    if created_ids:
        payload["resources"] = resources
        _save_resources_payload(payload)
    return created_ids


def _attach_crawl_relpaths(crawl_results: List[Dict]) -> List[Dict]:
    normalized: List[Dict] = []
    key_map = {
        "markdown_path": "markdown_relpath",
        "pdf_path": "pdf_relpath",
        "docx_path": "docx_relpath",
    }
    for item in crawl_results or []:
        if not isinstance(item, dict):
            continue
        cloned = dict(item)
        for key, rel_key in key_map.items():
            raw_path = cloned.get(key)
            if not raw_path:
                continue
            try:
                rel = Path(raw_path).resolve().relative_to(BASE_DIR).as_posix()
            except Exception:
                continue
            cloned[rel_key] = rel
        normalized.append(cloned)
    return normalized


def _build_crawler_agent_summary(crawl_results: List[Dict]) -> str:
    total = len(crawl_results or [])
    ok_count = sum(1 for item in (crawl_results or []) if (item or {}).get("status") == "ok")
    fail_count = max(total - ok_count, 0)
    lines = [f"抓取任务完成：共 {total} 条，成功 {ok_count} 条，失败 {fail_count} 条。"]
    if not total:
        lines.append("未识别到可抓取链接，请检查输入内容是否包含公众号文章 URL。")
        return "\n".join(lines)
    for idx, item in enumerate(crawl_results, start=1):
        title = (item or {}).get("title") or (item or {}).get("url") or f"抓取文章{idx}"
        status = "成功" if (item or {}).get("status") == "ok" else "失败"
        lines.append(f"{idx}. {title}（{status}）")
    return "\n".join(lines)


def _normalize_group_records(groups: List[Dict]) -> List[Dict]:
    normalized: List[Dict] = []
    seen_ids = set()
    seen_names = set()
    for item in groups or []:
        if not isinstance(item, dict):
            continue
        group_id = (item.get("group_id") or "").strip()
        group_name = (item.get("group_name") or "").strip()
        if not group_id or not group_name:
            continue
        lowered = group_name.lower()
        if group_id in seen_ids or lowered in seen_names:
            continue
        seen_ids.add(group_id)
        seen_names.add(lowered)
        normalized.append(
            {
                "group_id": group_id,
                "group_name": group_name,
                "updated_at": (item.get("updated_at") or "").strip(),
            }
        )
    return normalized


def _summarize_resource_groups(resources: List[Dict], manual_groups: Optional[List[Dict]] = None) -> List[Dict]:
    groups: Dict[str, Dict] = {}
    for item in manual_groups or []:
        group_id = (item.get("group_id") or "").strip()
        group_name = (item.get("group_name") or "").strip()
        if not group_id or not group_name:
            continue
        groups[group_id] = {
            "group_id": group_id,
            "group_name": group_name,
            "count": 0,
            "updated_at": (item.get("updated_at") or "").strip(),
        }
    for item in resources:
        group_id = item.get("group_id") or f"legacy_{DEFAULT_GROUP_NAME}"
        group_name = item.get("group_name") or DEFAULT_GROUP_NAME
        group = groups.get(group_id)
        if not group:
            group = {
                "group_id": group_id,
                "group_name": group_name,
                "count": 0,
                "updated_at": item.get("uploaded_at") or "",
            }
            groups[group_id] = group
        elif group_name:
            group["group_name"] = group_name
        group["count"] += 1
        uploaded_at = item.get("uploaded_at") or ""
        if uploaded_at and uploaded_at > (group.get("updated_at") or ""):
            group["updated_at"] = uploaded_at
    return sorted(groups.values(), key=lambda x: x.get("updated_at", ""), reverse=True)


def _find_group_by_name(resources: List[Dict], group_name: str, groups: Optional[List[Dict]] = None) -> Tuple[str, str]:
    target = (group_name or "").strip()
    if not target:
        return "", ""
    for item in groups or []:
        if (item.get("group_name") or "").strip() == target:
            return (item.get("group_id") or "").strip(), target
    for item in resources:
        if (item.get("group_name") or "").strip() == target:
            return (item.get("group_id") or "").strip(), target
    return "", ""


def _resolve_group_id_by_name(resources: List[Dict], group_name: str, groups: Optional[List[Dict]] = None) -> str:
    target = (group_name or "").strip()
    if not target:
        return ""
    existing_id, _ = _find_group_by_name(resources, target, groups)
    if existing_id:
        return existing_id
    return _gen_id("group")


def _copy_resource_record(item: Dict, target_group_id: str, target_name: str) -> Tuple[Dict, str]:
    new_id = _gen_id("res")
    ext = (item.get("extension") or "").strip() or Path(item.get("original_name") or "").suffix
    original_name = item.get("original_name") or new_id
    safe_stem = secure_filename(Path(original_name).stem) or "file"
    stored_name = f"{new_id}_{safe_stem}{ext}"
    upload_rel = item.get("source_relpath") or ""
    markdown_rel = item.get("markdown_relpath") or ""

    new_upload_rel = ""
    new_markdown_rel = ""

    if upload_rel:
        src_upload = (BASE_DIR / upload_rel).resolve()
        if src_upload.exists() and src_upload.is_file():
            dst_upload = (UPLOAD_DIR / stored_name).resolve()
            dst_upload.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(src_upload, dst_upload)
            new_upload_rel = dst_upload.relative_to(BASE_DIR).as_posix()

    if markdown_rel:
        src_md = (BASE_DIR / markdown_rel).resolve()
        if src_md.exists() and src_md.is_file():
            dst_md = (MARKDOWN_DIR / _markdown_filename(new_id, original_name)).resolve()
            dst_md.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(src_md, dst_md)
            new_markdown_rel = dst_md.relative_to(BASE_DIR).as_posix()

    if not new_markdown_rel and markdown_rel:
        return {}, "markdown 文件不存在"

    record = dict(item)
    record.update(
        {
            "id": new_id,
            "stored_name": stored_name,
            "uploaded_at": _now_iso(),
            "group_id": target_group_id,
            "group_name": target_name,
            "source_relpath": new_upload_rel,
            "markdown_relpath": new_markdown_rel,
            "status": item.get("status") or "ok",
            "error": "",
            "progress": 100,
            "progress_message": "完成",
        }
    )
    if new_upload_rel:
        try:
            record["size_bytes"] = (BASE_DIR / new_upload_rel).stat().st_size
        except Exception:
            record["size_bytes"] = item.get("size_bytes", 0)
    return record, ""


@strategy_watch_bp.route("/api/strategy-watch/runtime", methods=["GET"])
def strategy_watch_runtime():
    _load_env_files()
    model_default = _canonical_model_name(os.getenv("OPENAI_MODEL") or "deepseek-v3.2")
    if model_default not in ALLOWED_RUNTIME_MODEL_VALUES:
        model_default = "deepseek-v3.2"
    portrait_model_default = (
        _canonical_model_name((os.getenv("OPENAI_PORTRAIT_MODEL") or "").strip())
        or model_default
    )
    if portrait_model_default not in ALLOWED_RUNTIME_MODEL_VALUES:
        portrait_model_default = model_default
    return jsonify(
        {
            "success": True,
            "data": {
                "base_url": (os.getenv("OPENAI_BASE_URL") or "https://api.openai.com/v1"),
                "model": model_default,
                "portrait_model": portrait_model_default,
                "model_options": _get_runtime_model_options(),
                "api_key_configured": bool(
                    (os.getenv("OPENAI_API_KEY") or os.getenv("LLM_API_KEY") or "").strip()
                ),
                "agents": list_agent_profiles(),
                "conversation_modes": _build_conversation_mode_runtime(),
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/agents", methods=["GET"])
def strategy_watch_agents():
    return jsonify({"success": True, "data": list_agent_profiles(), "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/resources", methods=["GET"])
def list_strategy_resources():
    with _STORE_LOCK:
        payload = _load_resources_payload()
        resources = payload.get("resources", [])
        groups = payload.get("groups", [])
        active_resource_id = payload.get("active_resource_id") or ""
    resources = sorted(resources, key=lambda x: x.get("uploaded_at", ""), reverse=True)
    return jsonify(
        {
            "success": True,
            "data": {
                "resources": resources,
                "groups": _summarize_resource_groups(resources, groups),
                "active_resource_id": active_resource_id,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resources", methods=["POST"])
def upload_strategy_resources():
    start_ts = time.time()
    try:
        req_len = int(request.content_length or 0)
    except Exception:
        req_len = 0
    try:
        print(f"[strategy-watch] upload request in, content_length={req_len}")
    except Exception:
        pass

    if "files" not in request.files:
        return jsonify({"success": False, "error": "请通过 files 字段上传文件。", "timestamp": _now_iso()}), 400

    files = request.files.getlist("files")
    whisper_model = (request.form.get("whisper_model") or "tiny").strip()
    group_name = (request.form.get("group_name") or "").strip() or DEFAULT_GROUP_NAME
    strategy_id = (request.form.get("strategy_id") or "").strip()
    strategy_name = ""
    if strategy_id:
        strategies_payload = _load_strategies_payload()
        strategy = _find_by_id(strategies_payload.get("strategies", []), strategy_id)
        if not strategy:
            return jsonify({"success": False, "error": "策略不存在。", "timestamp": _now_iso()}), 404
        strategy_name = (strategy.get("name") or "").strip()

    uploaded = []
    jobs = []
    rejected = []
    to_process = []
    with _STORE_LOCK:
        payload = _load_resources_payload()
        resources = payload.get("resources", [])
        groups = payload.get("groups", [])

        existing_group_id = _resolve_group_id_by_name(resources, group_name, groups)
        group_id = existing_group_id or _gen_id("group")
        if not any((item.get("group_id") or "").strip() == group_id for item in groups):
            groups.append({"group_id": group_id, "group_name": group_name, "updated_at": _now_iso()})

        for file_storage in files:
            original_name = (file_storage.filename or "").strip()
            if not original_name:
                continue

            ext = Path(original_name).suffix.lower()
            if ext not in SUPPORTED_EXTENSIONS:
                rejected.append({"filename": original_name, "reason": f"不支持的文件类型: {ext}"})
                continue

            rid = _gen_id("res")
            safe_stem = secure_filename(Path(original_name).stem) or "file"
            stored_name = f"{rid}_{safe_stem}{ext}"
            upload_path = UPLOAD_DIR / stored_name
            markdown_path = MARKDOWN_DIR / _markdown_filename(rid, original_name)

            file_storage.save(upload_path)

            rel_upload = upload_path.relative_to(BASE_DIR).as_posix()
            rel_markdown = markdown_path.relative_to(BASE_DIR).as_posix()
            inferred_content_time, inferred_hit = _extract_first_date(original_name)
            record = {
                "id": rid,
                "original_name": original_name,
                "stored_name": stored_name,
                "extension": ext,
                "size_bytes": upload_path.stat().st_size if upload_path.exists() else 0,
                "uploaded_at": _now_iso(),
                "status": "processing",
                "error": "",
                "source_type": "file",
                "source_relpath": rel_upload,
                "markdown_relpath": rel_markdown,
                "group_id": group_id,
                "group_name": group_name,
                "strategy_id": strategy_id,
                "strategy_name": strategy_name,
                "progress": 0,
                "progress_message": "排队中",
                "published_at": "",
                "content_time": inferred_content_time,
                "content_time_type": "inferred_filename" if inferred_content_time else "unknown",
                "content_time_confidence": 0.70 if inferred_content_time else 0.0,
                "content_time_evidence": (f"filename:{inferred_hit}" if inferred_content_time else ""),
            }
            resources.append(record)
            uploaded.append(record)
            to_process.append(
                {
                    "resource_id": rid,
                    "upload_path": upload_path,
                    "markdown_path": markdown_path,
                    "whisper_model": whisper_model,
                    "original_name": original_name,
                }
            )

        payload["resources"] = resources
        payload["groups"] = groups
        _save_resources_payload(payload)

    for item in to_process:
        job_id = _gen_id("job")
        job_payload = {
            "resource_id": item["resource_id"],
            "original_name": item["original_name"],
            "status": "queued",
            "progress": 0,
            "message": "排队中",
            "strategy_id": strategy_id,
            "strategy_name": strategy_name,
            "created_at": _now_iso(),
        }
        _set_job(job_id, **job_payload)
        job_payload["job_id"] = job_id
        jobs.append(job_payload)
        worker = threading.Thread(
            target=_process_resource_job,
            args=(
                job_id,
                item["resource_id"],
                item["upload_path"],
                item["markdown_path"],
                item["whisper_model"],
            ),
            daemon=True,
        )
        worker.start()

    try:
        cost_ms = int((time.time() - start_ts) * 1000)
        print(
            f"[strategy-watch] upload accepted, files={len(uploaded)}, rejected={len(rejected)}, jobs={len(jobs)}, cost_ms={cost_ms}"
        )
    except Exception:
        pass

    return jsonify(
        {
            "success": True,
            "data": {
                "uploaded": uploaded,
                "rejected": rejected,
                "group_id": group_id,
                "group_name": group_name,
                "jobs": jobs,
                "async": True,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resources/jobs/<string:job_id>", methods=["GET"])
def get_resource_job(job_id: str):
    job = _get_job(job_id)
    if not job:
        return jsonify({"success": False, "error": "任务不存在。", "timestamp": _now_iso()}), 404
    return jsonify({"success": True, "data": job, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/resources/jobs", methods=["GET"])
def list_resource_jobs():
    return jsonify({"success": True, "data": _list_jobs(), "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/agent-logs", methods=["GET"])
def list_agent_logs():
    limit_raw = (request.args.get("limit") or "").strip()
    bucket = (request.args.get("bucket") or request.args.get("name") or "").strip()
    try:
        limit = int(limit_raw) if limit_raw else 200
    except Exception:
        limit = 200
    limit = max(1, min(limit, 2000))
    files: List[str] = []
    if AGENT_LOGS_DIR.exists() and AGENT_LOGS_DIR.is_dir():
        files = sorted(path.stem for path in AGENT_LOGS_DIR.glob("*.jsonl"))
    return jsonify(
        {
            "success": True,
            "data": _read_agent_logs(limit=limit, bucket=bucket),
            "meta": {"bucket": bucket, "files": files},
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resources/<string:resource_id>", methods=["DELETE"])
def delete_strategy_resource(resource_id: str):
    remaining_resources: List[Dict] = []
    with _STORE_LOCK:
        payload = _load_resources_payload()
        resources = payload.get("resources", [])
        target = _find_by_id(resources, resource_id)
        if not target:
            return jsonify({"success": False, "error": "资源不存在。", "timestamp": _now_iso()}), 404

        resources = [item for item in resources if item.get("id") != resource_id]
        remaining_resources = resources
        payload["resources"] = resources
        if payload.get("active_resource_id") == resource_id:
            payload["active_resource_id"] = ""
        _save_resources_payload(payload)

    for rel_key in ("source_relpath", "markdown_relpath"):
        rel_path = target.get(rel_key) or ""
        if not rel_path or "://" in rel_path:
            continue
        if _is_relpath_referenced(remaining_resources, rel_path):
            continue
        file_path = BASE_DIR / rel_path
        try:
            if file_path.exists() and file_path.is_file():
                file_path.unlink()
        except Exception:
            pass

    return jsonify({"success": True, "message": "资源已删除。", "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/resources/<string:resource_id>", methods=["PATCH"])
def rename_strategy_resource(resource_id: str):
    payload = request.get_json(silent=True) or {}
    name = (payload.get("name") or payload.get("original_name") or "").strip()
    if not name:
        return jsonify({"success": False, "error": "资源名称不能为空。", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        resources_payload = _load_resources_payload()
        resources = resources_payload.get("resources", [])
        target = _find_by_id(resources, resource_id)
        if not target:
            return jsonify({"success": False, "error": "资源不存在。", "timestamp": _now_iso()}), 404

        current_ext = (target.get("extension") or "").strip()
        incoming_ext = Path(name).suffix
        if incoming_ext and incoming_ext != current_ext:
            name = f"{Path(name).stem}{current_ext}" if current_ext else name

        target["original_name"] = name
        resources_payload["resources"] = resources
        _save_resources_payload(resources_payload)

    return jsonify({"success": True, "data": target, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/resources/<string:resource_id>/markdown", methods=["GET"])
def get_resource_markdown(resource_id: str):
    with _STORE_LOCK:
        payload = _load_resources_payload()
        resources = payload.get("resources", [])
    target = _find_by_id(resources, resource_id)
    if not target:
        return jsonify({"success": False, "error": "资源不存在。", "timestamp": _now_iso()}), 404

    content = _read_markdown_content(target)
    return jsonify(
        {
            "success": True,
            "data": {
                "resource_id": resource_id,
                "original_name": target.get("original_name"),
                "content": content,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resources/<string:resource_id>/download", methods=["GET"])
def download_resource_file(resource_id: str):
    fmt = (request.args.get("format") or "markdown").strip().lower()
    with _STORE_LOCK:
        payload = _load_resources_payload()
        resources = payload.get("resources", [])
    target = _find_by_id(resources, resource_id)
    if not target:
        return jsonify({"success": False, "error": "资源不存在。", "timestamp": _now_iso()}), 404

    if fmt in {"source", "origin", "original", "markdown", "md"}:
        markdown_path, err = _resolve_markdown_path_from_resource(target)
        if err:
            return jsonify({"success": False, "error": err, "timestamp": _now_iso()}), 404
        return send_file(
            markdown_path,
            as_attachment=True,
            download_name=_build_resource_markdown_download_name(target),
        )

    if fmt in {"ai_summary", "summary"}:
        markdown_path, err = _resolve_markdown_path_from_resource(target)
        if err:
            return jsonify({"success": False, "error": err, "timestamp": _now_iso()}), 404
        summary_path = _resource_ai_summary_path(target)
        regenerate = (request.args.get("regenerate") or "").strip().lower() in {"1", "true", "yes", "on"}
        if regenerate or (not summary_path.exists()):
            try:
                summary_path.parent.mkdir(parents=True, exist_ok=True)
                summarize_markdown_with_llm(markdown_path, summary_path)
            except Exception as exc:
                return jsonify({"success": False, "error": f"生成 AI 总结失败: {exc}", "timestamp": _now_iso()}), 500
        if not summary_path.exists() or not summary_path.is_file():
            return jsonify({"success": False, "error": "AI 总结文件不存在。", "timestamp": _now_iso()}), 404
        return send_file(summary_path, as_attachment=True, download_name=summary_path.name)

    return jsonify({"success": False, "error": "不支持的下载格式。", "timestamp": _now_iso()}), 400


@strategy_watch_bp.route("/api/strategy-watch/resources/<string:resource_id>/ai-summary/download", methods=["GET"])
def download_resource_ai_summary(resource_id: str):
    with _STORE_LOCK:
        payload = _load_resources_payload()
        resources = payload.get("resources", [])
    target = _find_by_id(resources, resource_id)
    if not target:
        return jsonify({"success": False, "error": "资源不存在。", "timestamp": _now_iso()}), 404

    markdown_path, err = _resolve_markdown_path_from_resource(target)
    if err:
        return jsonify({"success": False, "error": err, "timestamp": _now_iso()}), 404

    summary_path = _resource_ai_summary_path(target)
    regenerate = (request.args.get("regenerate") or "").strip().lower() in {"1", "true", "yes", "on"}
    if regenerate or (not summary_path.exists()):
        try:
            summary_path.parent.mkdir(parents=True, exist_ok=True)
            summarize_markdown_with_llm(markdown_path, summary_path)
        except Exception as exc:
            return jsonify({"success": False, "error": f"生成 AI 总结失败: {exc}", "timestamp": _now_iso()}), 500

    if not summary_path.exists() or not summary_path.is_file():
        return jsonify({"success": False, "error": "AI 总结文件不存在。", "timestamp": _now_iso()}), 404

    return send_file(summary_path, as_attachment=True, download_name=summary_path.name)


@strategy_watch_bp.route("/api/strategy-watch/crawled/download", methods=["GET"])
def download_crawled_file():
    relpath = (request.args.get("path") or "").strip()
    if not relpath:
        return jsonify({"success": False, "error": "缺少 path 参数。", "timestamp": _now_iso()}), 400
    if "://" in relpath or Path(relpath).is_absolute():
        return jsonify({"success": False, "error": "非法路径。", "timestamp": _now_iso()}), 400

    file_path = (BASE_DIR / relpath).resolve()
    try:
        file_path.relative_to(CRAWLED_DIR)
    except Exception:
        return jsonify({"success": False, "error": "非法路径。", "timestamp": _now_iso()}), 403
    if not file_path.exists() or not file_path.is_file():
        return jsonify({"success": False, "error": "文件不存在。", "timestamp": _now_iso()}), 404

    return send_file(file_path, as_attachment=True, download_name=file_path.name)


@strategy_watch_bp.route("/api/strategy-watch/resource-groups", methods=["POST"])
def create_strategy_resource_group():
    payload = request.get_json(silent=True) or {}
    group_name = (payload.get("group_name") or payload.get("name") or "").strip()
    if not group_name:
        return jsonify({"success": False, "error": "组名不能为空", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        resources_payload = _load_resources_payload()
        resources = resources_payload.get("resources", [])
        groups = resources_payload.get("groups", [])
        existing_id, _ = _find_group_by_name(resources, group_name, groups)
        if existing_id:
            return jsonify({"success": False, "error": "组名已存在", "timestamp": _now_iso()}), 400

        group_id = _gen_id("group")
        groups.append({"group_id": group_id, "group_name": group_name, "updated_at": _now_iso()})
        resources_payload["groups"] = groups
        _save_resources_payload(resources_payload)

    return jsonify(
        {
            "success": True,
            "data": {"group_id": group_id, "group_name": group_name},
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resource-groups/<string:group_id>", methods=["PATCH"])
def rename_strategy_resource_group(group_id: str):
    payload = request.get_json(silent=True) or {}
    group_name = (payload.get("group_name") or payload.get("name") or "").strip()
    if not group_name:
        return jsonify({"success": False, "error": "组名不能为空", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        resources_payload = _load_resources_payload()
        resources = resources_payload.get("resources", [])
        groups = resources_payload.get("groups", [])
        existing_id, _ = _find_group_by_name(resources, group_name, groups)
        if existing_id and existing_id != group_id:
            return jsonify({"success": False, "error": "组名已存在", "timestamp": _now_iso()}), 400
        updated = 0
        for item in resources:
            if item.get("group_id") == group_id:
                item["group_name"] = group_name
                updated += 1
        group_updated = False
        for item in groups:
            if (item.get("group_id") or "").strip() == group_id:
                item["group_name"] = group_name
                item["updated_at"] = _now_iso()
                group_updated = True
                break
        if not group_updated and updated > 0:
            groups.append({"group_id": group_id, "group_name": group_name, "updated_at": _now_iso()})
            group_updated = True
        if updated == 0 and not group_updated:
            return jsonify({"success": False, "error": "资源组不存在", "timestamp": _now_iso()}), 404
        resources_payload["resources"] = resources
        resources_payload["groups"] = groups
        _save_resources_payload(resources_payload)

    return jsonify(
        {
            "success": True,
            "data": {"group_id": group_id, "group_name": group_name, "updated_count": updated},
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resource-groups/transfer", methods=["POST"])
def transfer_strategy_resource_group():
    payload = request.get_json(silent=True) or {}
    source_group_id = (payload.get("source_group_id") or "").strip()
    target_group_id = (payload.get("target_group_id") or "").strip()
    target_group_name = (payload.get("target_group_name") or "").strip()
    mode = (payload.get("mode") or "move").strip().lower()
    if mode not in {"move", "copy"}:
        return jsonify({"success": False, "error": "mode 必须为 move 或 copy", "timestamp": _now_iso()}), 400
    if not source_group_id:
        return jsonify({"success": False, "error": "缺少 source_group_id", "timestamp": _now_iso()}), 400
    if not target_group_id and not target_group_name:
        return jsonify({"success": False, "error": "缺少目标分组", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        resources_payload = _load_resources_payload()
        resources = resources_payload.get("resources", [])
        groups = resources_payload.get("groups", [])
        source_items = [item for item in resources if item.get("group_id") == source_group_id]
        source_group_exists = any((item.get("group_id") or "").strip() == source_group_id for item in groups)
        if not source_items and not source_group_exists:
            return jsonify({"success": False, "error": "源分组不存在", "timestamp": _now_iso()}), 404

        if target_group_id:
            target_name = next(
                (item.get("group_name") for item in groups if (item.get("group_id") or "").strip() == target_group_id),
                "",
            ).strip()
            if not target_name:
                target_name = next(
                    (item.get("group_name") for item in resources if item.get("group_id") == target_group_id), ""
                ).strip()
            if not target_name:
                return jsonify({"success": False, "error": "目标分组不存在", "timestamp": _now_iso()}), 404
        else:
            existing_id, _ = _find_group_by_name(resources, target_group_name, groups)
            if existing_id:
                target_group_id = existing_id
                target_name = target_group_name
            else:
                target_group_id = _gen_id("group")
                target_name = target_group_name

        if target_group_id == source_group_id:
            return jsonify({"success": False, "error": "源分组与目标分组相同", "timestamp": _now_iso()}), 400

        moved = 0
        copied = 0
        skipped = []

        if mode == "move":
            for item in resources:
                if item.get("group_id") == source_group_id:
                    item["group_id"] = target_group_id
                    item["group_name"] = target_name
                    moved += 1
        else:
            new_records = []
            for item in source_items:
                new_id = _gen_id("res")
                ext = (item.get("extension") or "").strip() or Path(item.get("original_name") or "").suffix
                original_name = item.get("original_name") or new_id
                safe_stem = secure_filename(Path(original_name).stem) or "file"
                stored_name = f"{new_id}_{safe_stem}{ext}"
                upload_rel = item.get("source_relpath") or ""
                markdown_rel = item.get("markdown_relpath") or ""

                new_upload_rel = ""
                new_markdown_rel = ""

                # Copy upload file if exists
                if upload_rel:
                    src_upload = (BASE_DIR / upload_rel).resolve()
                    if src_upload.exists() and src_upload.is_file():
                        dst_upload = (UPLOAD_DIR / stored_name).resolve()
                        dst_upload.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copyfile(src_upload, dst_upload)
                        new_upload_rel = dst_upload.relative_to(BASE_DIR).as_posix()

                # Copy markdown file if exists
                if markdown_rel:
                    src_md = (BASE_DIR / markdown_rel).resolve()
                    if src_md.exists() and src_md.is_file():
                        dst_md = (MARKDOWN_DIR / _markdown_filename(new_id, original_name)).resolve()
                        dst_md.parent.mkdir(parents=True, exist_ok=True)
                        shutil.copyfile(src_md, dst_md)
                        new_markdown_rel = dst_md.relative_to(BASE_DIR).as_posix()

                if not new_markdown_rel and markdown_rel:
                    skipped.append({"resource_id": item.get("id"), "reason": "markdown 文件不存在"})
                    continue

                record = dict(item)
                record.update(
                    {
                        "id": new_id,
                        "stored_name": stored_name,
                        "uploaded_at": _now_iso(),
                        "group_id": target_group_id,
                        "group_name": target_name,
                        "source_relpath": new_upload_rel,
                        "markdown_relpath": new_markdown_rel,
                        "status": item.get("status") or "ok",
                        "error": "",
                        "progress": 100,
                        "progress_message": "完成",
                    }
                )
                if new_upload_rel:
                    try:
                        record["size_bytes"] = (BASE_DIR / new_upload_rel).stat().st_size
                    except Exception:
                        record["size_bytes"] = item.get("size_bytes", 0)
                new_records.append(record)
                copied += 1

            resources.extend(new_records)

        resources_payload["resources"] = resources
        if not any((item.get("group_id") or "").strip() == target_group_id for item in groups):
            groups.append({"group_id": target_group_id, "group_name": target_name, "updated_at": _now_iso()})
        resources_payload["groups"] = groups
        _save_resources_payload(resources_payload)

    return jsonify(
        {
            "success": True,
            "data": {
                "mode": mode,
                "source_group_id": source_group_id,
                "target_group_id": target_group_id,
                "target_group_name": target_name,
                "moved": moved,
                "copied": copied,
                "skipped": skipped,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resources/transfer", methods=["POST"])
def transfer_strategy_resources():
    payload = request.get_json(silent=True) or {}
    resource_ids = payload.get("resource_ids") or []
    if not isinstance(resource_ids, list):
        resource_ids = []
    mode = (payload.get("mode") or "move").strip().lower()
    if mode not in {"move", "copy"}:
        return jsonify({"success": False, "error": "mode 必须为 move 或 copy", "timestamp": _now_iso()}), 400
    if not resource_ids:
        return jsonify({"success": False, "error": "resource_ids 不能为空", "timestamp": _now_iso()}), 400

    target_group_id = (payload.get("target_group_id") or "").strip()
    target_group_name = (payload.get("target_group_name") or "").strip()
    if not target_group_id and not target_group_name:
        return jsonify({"success": False, "error": "缺少目标分组", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        resources_payload = _load_resources_payload()
        resources = resources_payload.get("resources", [])
        groups = resources_payload.get("groups", [])
        resource_map = {item.get("id"): item for item in resources}
        targets = [resource_map.get(rid) for rid in resource_ids if resource_map.get(rid)]
        if not targets:
            return jsonify({"success": False, "error": "未找到资源", "timestamp": _now_iso()}), 404

        if target_group_id:
            target_name = next(
                (
                    item.get("group_name")
                    for item in groups
                    if (item.get("group_id") or "").strip() == target_group_id
                ),
                "",
            )
            if not target_name:
                target_name = next(
                (item.get("group_name") for item in resources if item.get("group_id") == target_group_id),
                "",
                )
            if not target_name:
                target_name = target_group_name or DEFAULT_GROUP_NAME
        else:
            existing_id, _ = _find_group_by_name(resources, target_group_name, groups)
            if existing_id:
                target_group_id = existing_id
                target_name = target_group_name
            else:
                target_group_id = _gen_id("group")
                target_name = target_group_name

        moved = 0
        copied = 0
        skipped: List[Dict] = []

        if mode == "move":
            for item in targets:
                item["group_id"] = target_group_id
                item["group_name"] = target_name
                moved += 1
        else:
            new_records: List[Dict] = []
            for item in targets:
                record, reason = _copy_resource_record(item, target_group_id, target_name)
                if not record:
                    skipped.append({"resource_id": item.get("id"), "reason": reason or "复制失败"})
                    continue
                new_records.append(record)
                copied += 1
            resources.extend(new_records)

        resources_payload["resources"] = resources
        if not any((item.get("group_id") or "").strip() == target_group_id for item in groups):
            groups.append({"group_id": target_group_id, "group_name": target_name, "updated_at": _now_iso()})
        resources_payload["groups"] = groups
        _save_resources_payload(resources_payload)

    return jsonify(
        {
            "success": True,
            "data": {
                "mode": mode,
                "target_group_id": target_group_id,
                "target_group_name": target_name,
                "moved": moved,
                "copied": copied,
                "skipped": skipped,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/resources/active", methods=["PATCH"])
def set_active_strategy_resource():
    payload = request.get_json(silent=True) or {}
    resource_id = (payload.get("resource_id") or "").strip()

    with _STORE_LOCK:
        resources_payload = _load_resources_payload()
        resources = resources_payload.get("resources", [])
        if resource_id and not _find_by_id(resources, resource_id):
            return jsonify({"success": False, "error": "资源不存在。", "timestamp": _now_iso()}), 404
        resources_payload["active_resource_id"] = resource_id
        _save_resources_payload(resources_payload)

    return jsonify(
        {"success": True, "data": {"active_resource_id": resource_id}, "timestamp": _now_iso()}
    )


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles", methods=["GET"])
def list_memory_profiles():
    with _STORE_LOCK:
        payload = _load_memory_profiles_payload()
        profiles = payload.get("profiles", [])
        active_profile_id = payload.get("active_profile_id") or ""
        links = _load_memory_links()
        portraits = _load_memory_portraits()
    data = [_memory_profile_view(item, links, portraits) for item in profiles]
    data.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
    return jsonify(
        {
            "success": True,
            "data": {
                "profiles": data,
                "active_profile_id": active_profile_id,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles", methods=["POST"])
def create_memory_profile():
    payload = request.get_json(silent=True) or {}
    name = (payload.get("name") or "").strip() or f"人格 {datetime.now().strftime('%m-%d %H:%M')}"
    description = (payload.get("description") or "").strip()
    source_blogger = (payload.get("source_blogger") or "").strip()
    profile = {
        "id": _gen_id("mp"),
        "name": name,
        "description": description,
        "source_blogger": source_blogger,
        "group_subscriptions": [],
        "created_at": _now_iso(),
        "updated_at": _now_iso(),
    }
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profiles = profiles_payload.get("profiles", [])
        profiles.append(profile)
        profiles_payload["profiles"] = profiles
        if not profiles_payload.get("active_profile_id"):
            profiles_payload["active_profile_id"] = profile["id"]
        _save_memory_profiles_payload(profiles_payload)
        portraits = _load_memory_portraits()
        portraits[profile["id"]] = _default_memory_portrait(profile["id"])
        _save_memory_portraits(portraits)
    return jsonify({"success": True, "data": profile, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/active", methods=["PATCH"])
def set_active_memory_profile():
    payload = request.get_json(silent=True) or {}
    profile_id = (payload.get("profile_id") or "").strip()
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profiles = profiles_payload.get("profiles", [])
        if profile_id and not _find_by_id(profiles, profile_id):
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        profiles_payload["active_profile_id"] = profile_id
        _save_memory_profiles_payload(profiles_payload)
    return jsonify(
        {"success": True, "data": {"active_profile_id": profile_id}, "timestamp": _now_iso()}
    )


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>", methods=["PATCH"])
def update_memory_profile(profile_id: str):
    payload = request.get_json(silent=True) or {}
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profiles = profiles_payload.get("profiles", [])
        target = _find_by_id(profiles, profile_id)
        if not target:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        if "name" in payload:
            target["name"] = (payload.get("name") or "").strip() or target.get("name")
        if "description" in payload:
            target["description"] = (payload.get("description") or "").strip()
        if "source_blogger" in payload:
            target["source_blogger"] = (payload.get("source_blogger") or "").strip()
        target["updated_at"] = _now_iso()
        profiles_payload["profiles"] = profiles
        _save_memory_profiles_payload(profiles_payload)
    return jsonify({"success": True, "data": target, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>", methods=["DELETE"])
def delete_memory_profile(profile_id: str):
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profiles = profiles_payload.get("profiles", [])
        target = _find_by_id(profiles, profile_id)
        if not target:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        profiles = [item for item in profiles if item.get("id") != profile_id]
        profiles_payload["profiles"] = profiles
        if profiles_payload.get("active_profile_id") == profile_id:
            profiles_payload["active_profile_id"] = ""
        _save_memory_profiles_payload(profiles_payload)

        links = [item for item in _load_memory_links() if item.get("profile_id") != profile_id]
        _save_memory_links(links)
        portraits = _load_memory_portraits()
        if profile_id in portraits:
            portraits.pop(profile_id, None)
            _save_memory_portraits(portraits)

    return jsonify({"success": True, "message": "人格已删除。", "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/bind-resources", methods=["POST"])
def bind_memory_profile_resources(profile_id: str):
    payload = request.get_json(silent=True) or {}
    resource_ids = payload.get("resource_ids") or []
    if not isinstance(resource_ids, list):
        resource_ids = []
    if not resource_ids:
        return jsonify({"success": False, "error": "resource_ids 不能为空。", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profile = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if not profile:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        result = _bind_resources_to_profile(profile_id, resource_ids, bind_source="manual")
        profile["updated_at"] = _now_iso()
        _save_memory_profiles_payload(profiles_payload)

    return jsonify({"success": True, "data": result, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/bind-group", methods=["POST"])
def bind_memory_profile_group(profile_id: str):
    payload = request.get_json(silent=True) or {}
    group_id = (payload.get("group_id") or "").strip()
    if not group_id:
        return jsonify({"success": False, "error": "缺少 group_id。", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profile = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if not profile:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404

        result = _sync_profile_group_resources(profile, group_id)
        profile["updated_at"] = _now_iso()
        _save_memory_profiles_payload(profiles_payload)

    return jsonify(
        {
            "success": True,
            "data": {
                "group_id": group_id,
                "added": result.get("added", 0),
                "skipped": result.get("skipped", []),
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/sync-group", methods=["POST"])
def sync_memory_profile_group(profile_id: str):
    payload = request.get_json(silent=True) or {}
    group_id = (payload.get("group_id") or "").strip()
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profile = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if not profile:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404

        subs = profile.get("group_subscriptions") if isinstance(profile.get("group_subscriptions"), list) else []
        target_group_ids = [group_id] if group_id else [
            (item.get("group_id") or "").strip()
            for item in subs
            if isinstance(item, dict) and (item.get("group_id") or "").strip()
        ]
        if not target_group_ids:
            return jsonify({"success": False, "error": "没有可同步的分组。", "timestamp": _now_iso()}), 400

        total_added = 0
        all_skipped: List[Dict] = []
        sync_details: List[Dict] = []
        for gid in target_group_ids:
            result = _sync_profile_group_resources(profile, gid)
            added = int(result.get("added") or 0)
            skipped = result.get("skipped") or []
            total_added += added
            all_skipped.extend(skipped)
            sync_details.append({"group_id": gid, "added": added, "skipped_count": len(skipped)})

        profile["updated_at"] = _now_iso()
        _save_memory_profiles_payload(profiles_payload)

    return jsonify(
        {
            "success": True,
            "data": {
                "synced_groups": sync_details,
                "added": total_added,
                "skipped": all_skipped,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/extract-portrait-draft", methods=["POST"])
def extract_memory_portrait_draft(profile_id: str):
    payload = request.get_json(silent=True) or {}
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profile = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if not profile:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        portrait_map = _load_memory_portraits()
        portrait = portrait_map.get(profile_id) or _default_memory_portrait(profile_id)

    snippets, used_ids, skipped = _collect_profile_resource_snippets(
        profile_id,
        max_chars=32_000,
        max_items=12,
        per_item_chars=2_400,
    )
    if not snippets:
        return jsonify(
            {
                "success": False,
                "error": "当前人格还没有可用于侧写的资料，请先绑定资料。",
                "timestamp": _now_iso(),
            }
        ), 400

    prompt = (
        "请基于资料生成人物侧写初稿，并返回 JSON（不要输出其他文字）。\n"
        "JSON 字段:\n"
        "{\n"
        '  "methodology": "交易方法论",\n'
        '  "tactics": "交易手法",\n'
        '  "views": "观点",\n'
        '  "operations": "交易操作",\n'
        '  "risk_rules": "风控规则",\n'
        '  "style_constraints": "表达与行为约束",\n'
        '  "evidence_refs": [\n'
        "    {\n"
        '      "resource_id": "res_xxx",\n'
        '      "topic": "观点|操作|方法论|手法",\n'
        '      "quote": "关键证据原文",\n'
        '      "source_time": "YYYY-MM-DD",\n'
        '      "source_time_type": "published|recorded|inferred_filename|inferred_text",\n'
        '      "timecode": "00:00:00-00:00:10"\n'
        "    }\n"
        "  ]\n"
        "}\n"
        "要求：证据尽量可追溯，source_time 优先使用资料内容时间。"
    )

    extra_context = (
        f"人格名称: {(profile.get('name') or '').strip()}\n"
        f"人格说明: {(profile.get('description') or '').strip()}\n"
        f"来源博主: {(profile.get('source_blogger') or '').strip()}"
    )
    requested_model = (payload.get("model") or "").strip()
    portrait_model = _canonical_model_name(
        requested_model
        or (os.getenv("OPENAI_PORTRAIT_MODEL") or "").strip()
        or "deepseek-v3.2"
    )
    try:
        reply, used_model, _provider, _usage = chat_with_agent(
            agent_name="analyst_agent",
            user_content=prompt,
            history_messages=[],
            resource_context=snippets,
            extra_context=extra_context,
            model_name=portrait_model,
            temperature=0.2,
            trace_context={
                "entrypoint": "extract_memory_portrait_draft",
                "profile_id": profile_id,
                "agent_name": "analyst_agent",
            },
        )
    except Exception as exc:
        return jsonify({"success": False, "error": f"侧写初稿生成失败: {exc}", "timestamp": _now_iso()}), 500

    parsed = _normalize_portrait_draft_payload(_extract_json_block(reply), reply)
    if not parsed:
        parsed = {"methodology": reply}

    def _section_to_text(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, list):
            lines = []
            for item in value:
                txt = str(item or "").strip()
                if txt:
                    lines.append(f"- {txt}")
            return "\n".join(lines).strip()
        if isinstance(value, dict):
            return json.dumps(value, ensure_ascii=False, indent=2)
        return str(value).strip()

    evidence_refs_raw = parsed.get("evidence_refs") if isinstance(parsed.get("evidence_refs"), list) else []
    evidence_refs: List[Dict] = []
    for item in evidence_refs_raw:
        normalized = _normalize_evidence_ref(item)
        if normalized:
            evidence_refs.append(normalized)

    portrait.update(
        {
            "methodology": _section_to_text(parsed.get("methodology")),
            "tactics": _section_to_text(parsed.get("tactics")),
            "views": _section_to_text(parsed.get("views")),
            "operations": _section_to_text(parsed.get("operations")),
            "risk_rules": _section_to_text(parsed.get("risk_rules")),
            "style_constraints": _section_to_text(parsed.get("style_constraints")),
            "evidence_refs": evidence_refs,
            "updated_at": _now_iso(),
            "updated_by": "ai_draft",
        }
    )

    with _STORE_LOCK:
        portrait_map = _load_memory_portraits()
        portrait_map[profile_id] = portrait
        _save_memory_portraits(portrait_map)
        profiles_payload = _load_memory_profiles_payload()
        target = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if target:
            target["updated_at"] = _now_iso()
            _save_memory_profiles_payload(profiles_payload)

    return jsonify(
        {
            "success": True,
            "data": {
                "portrait": portrait,
                "used_model": used_model,
                "used_resource_ids": used_ids,
                "skipped_refs": skipped,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/portrait", methods=["GET"])
def get_memory_portrait(profile_id: str):
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profile = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if not profile:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        portrait_map = _load_memory_portraits()
        portrait = portrait_map.get(profile_id) or _default_memory_portrait(profile_id)
    return jsonify({"success": True, "data": portrait, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/portrait", methods=["PATCH"])
def update_memory_portrait(profile_id: str):
    payload = request.get_json(silent=True) or {}
    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profile = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if not profile:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        portrait_map = _load_memory_portraits()
        portrait = portrait_map.get(profile_id) or _default_memory_portrait(profile_id)

        for key in ("methodology", "tactics", "views", "operations", "risk_rules", "style_constraints"):
            if key in payload:
                portrait[key] = (payload.get(key) or "").strip()

        if "evidence_refs" in payload:
            refs_raw = payload.get("evidence_refs")
            refs = []
            if isinstance(refs_raw, list):
                for item in refs_raw:
                    normalized = _normalize_evidence_ref(item)
                    if normalized:
                        refs.append(normalized)
            portrait["evidence_refs"] = refs

        portrait["updated_at"] = _now_iso()
        portrait["updated_by"] = "manual"
        portrait_map[profile_id] = portrait
        _save_memory_portraits(portrait_map)

        profile["updated_at"] = _now_iso()
        _save_memory_profiles_payload(profiles_payload)

    return jsonify({"success": True, "data": portrait, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/portrait/export", methods=["POST"])
def export_memory_portrait(profile_id: str):
    payload = request.get_json(silent=True) or {}
    export_format = (payload.get("format") or request.args.get("format") or "md").strip().lower()
    if export_format not in {"md", "docx", "pdf"}:
        return jsonify({"success": False, "error": "仅支持导出 md/docx/pdf。", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        profiles_payload = _load_memory_profiles_payload()
        profile = _find_by_id(profiles_payload.get("profiles", []), profile_id)
        if not profile:
            return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
        portrait_map = _load_memory_portraits()
        stored_portrait = portrait_map.get(profile_id) or _default_memory_portrait(profile_id)

    portrait = _merge_portrait_override(stored_portrait, payload.get("portrait"))
    markdown_text = _build_memory_portrait_markdown(profile, portrait)

    profile_name = (profile.get("name") or profile_id).strip() or profile_id
    base_name = secure_filename(profile_name) or profile_id
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d")

    try:
        if export_format == "md":
            content = markdown_text.encode("utf-8")
            mimetype = "text/markdown; charset=utf-8"
            ext = "md"
        elif export_format == "docx":
            content = _render_memory_portrait_docx(profile, portrait)
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            content = _render_memory_portrait_pdf(profile, portrait)
            mimetype = "application/pdf"
            ext = "pdf"
    except Exception as exc:
        return jsonify({"success": False, "error": f"导出失败: {exc}", "timestamp": _now_iso()}), 500

    return send_file(
        BytesIO(content),
        as_attachment=True,
        download_name=f"{base_name}_portrait_{stamp}.{ext}",
        mimetype=mimetype,
    )


@strategy_watch_bp.route("/api/strategy-watch/memory-profiles/<string:profile_id>/preview-context", methods=["GET"])
def preview_memory_profile_context(profile_id: str):
    context_text, meta = _build_memory_profile_context(profile_id)
    if not context_text and not meta.get("memory_profile_id"):
        return jsonify({"success": False, "error": "人格不存在。", "timestamp": _now_iso()}), 404
    return jsonify(
        {
            "success": True,
            "data": {
                "memory_profile_id": meta.get("memory_profile_id") or profile_id,
                "memory_profile_name": meta.get("memory_profile_name") or "",
                "context": context_text,
                "used_memory_resource_ids": meta.get("used_memory_resource_ids") or [],
                "skipped_memory_refs": meta.get("skipped_memory_refs") or [],
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/conversations", methods=["GET"])
def list_strategy_conversations():
    with _STORE_LOCK:
        conversations = _load_conversations()

    summaries = []
    for conv in conversations:
        messages = conv.get("messages", [])
        last_message = messages[-1] if messages else None
        summaries.append(
            {
                "id": conv.get("id"),
                "title": conv.get("title"),
                "created_at": conv.get("created_at"),
                "updated_at": conv.get("updated_at"),
                "message_count": len(messages),
                "last_message_preview": (last_message or {}).get("content", "")[:120] if last_message else "",
            }
        )

    summaries.sort(key=lambda x: x.get("updated_at", ""), reverse=True)
    return jsonify({"success": True, "data": summaries, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/conversations", methods=["POST"])
def create_strategy_conversation():
    payload = request.get_json(silent=True) or {}
    title = (payload.get("title") or "").strip() or f"策略会话 {datetime.now().strftime('%m-%d %H:%M')}"

    conversation = {
        "id": _gen_id("conv"),
        "title": title,
        "created_at": _now_iso(),
        "updated_at": _now_iso(),
        "messages": [],
    }

    with _STORE_LOCK:
        conversations = _load_conversations()
        conversations.append(conversation)
        _save_conversations(conversations)

    return jsonify({"success": True, "data": conversation, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/conversations/<string:conversation_id>", methods=["PATCH"])
def rename_strategy_conversation(conversation_id: str):
    payload = request.get_json(silent=True) or {}
    title = (payload.get("title") or "").strip()
    if not title:
        return jsonify({"success": False, "error": "会话标题不能为空。", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        conversations = _load_conversations()
        conv = _find_by_id(conversations, conversation_id)
        if not conv:
            return jsonify({"success": False, "error": "会话不存在。", "timestamp": _now_iso()}), 404
        conv["title"] = title
        conv["updated_at"] = _now_iso()
        _save_conversations(conversations)

    return jsonify({"success": True, "data": conv, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/conversations/<string:conversation_id>", methods=["DELETE"])
def delete_strategy_conversation(conversation_id: str):
    with _STORE_LOCK:
        conversations = _load_conversations()
        target = _find_by_id(conversations, conversation_id)
        if not target:
            return jsonify({"success": False, "error": "会话不存在。", "timestamp": _now_iso()}), 404

        conversations = [item for item in conversations if item.get("id") != conversation_id]
        _save_conversations(conversations)

    return jsonify({"success": True, "message": "会话已删除。", "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/conversations/<string:conversation_id>/messages", methods=["GET"])
def list_strategy_messages(conversation_id: str):
    with _STORE_LOCK:
        conversations = _load_conversations()
    conv = _find_by_id(conversations, conversation_id)
    if not conv:
        return jsonify({"success": False, "error": "会话不存在。", "timestamp": _now_iso()}), 404
    return jsonify(
        {
            "success": True,
            "data": {
                "id": conv.get("id"),
                "title": conv.get("title"),
                "messages": conv.get("messages", []),
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/conversations/<string:conversation_id>/messages", methods=["POST"])
def send_strategy_message(conversation_id: str):
    payload = request.get_json(silent=True) or {}
    content = (payload.get("content") or "").strip()
    if not content:
        return jsonify({"success": False, "error": "消息内容不能为空。", "timestamp": _now_iso()}), 400

    resource_ids = payload.get("resource_ids") or []
    if not isinstance(resource_ids, list):
        resource_ids = []
    strategy_id = (payload.get("strategy_id") or "").strip()
    memory_profile_id = (payload.get("memory_profile_id") or "").strip()
    prompt_template_id = (payload.get("prompt_template_id") or "").strip()
    prompt_template = (payload.get("prompt_template") or "").strip()
    llm_user_content = _compose_user_content_with_prompt_template(content, prompt_template)

    agent_name, conversation_mode, run_crawler_direct = _resolve_conversation_dispatch(payload)

    temperature = payload.get("temperature", 0.2)
    model_name = (payload.get("model") or "").strip()
    crawler_cookie = (
        (payload.get("crawler_cookie") or "")
        or (request.headers.get("X-Wechat-Cookie") or "")
    ).strip()

    with _STORE_LOCK:
        conversations = _load_conversations()
        conv = _find_by_id(conversations, conversation_id)
        if not conv:
            return jsonify({"success": False, "error": "会话不存在。", "timestamp": _now_iso()}), 404

        user_message = {
            "id": _gen_id("msg"),
            "role": "user",
            "content": content,
            "created_at": _now_iso(),
            "resource_ids": resource_ids,
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": strategy_id,
            "memory_profile_id": memory_profile_id,
            "prompt_template_id": prompt_template_id,
            "prompt_template": prompt_template,
        }
        conv.setdefault("messages", []).append(user_message)
        conv["updated_at"] = _now_iso()
        _save_conversations(conversations)
    conversation_title = (conv.get("title") or "").strip()
    strategy_name = _resolve_strategy_name(strategy_id)
    _append_agent_log(
        "conversation_message_received",
        {
            "conversation_id": conversation_id,
            "conversation_title": conversation_title,
            "user_message_id": user_message.get("id"),
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": strategy_id,
            "strategy_name": strategy_name,
            "model": model_name or "",
            "stream": False,
        },
    )

    crawled_resource_ids: List[str] = []
    crawl_results: List[Dict] = []
    if run_crawler_direct:
        crawl_results = crawl_wechat_articles_from_text(
            content,
            output_dir=CRAWLED_DIR,
            cookie_override=crawler_cookie,
        )
        crawl_results = _attach_crawl_relpaths(crawl_results)
        ok_items = [x for x in crawl_results if x.get("status") == "ok"]
        if ok_items:
            with _STORE_LOCK:
                crawled_resource_ids = _register_crawled_resources(ok_items)
        assistant_message = {
            "id": _gen_id("msg"),
            "role": "assistant",
            "content": _build_crawler_agent_summary(crawl_results),
            "created_at": _now_iso(),
            "model": "crawler_direct",
            "provider": "crawler",
            "usage": {},
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": strategy_id,
            "memory_profile_id": memory_profile_id,
            "used_resource_ids": [],
            "skipped_resource_refs": [],
            "used_memory_resource_ids": [],
            "skipped_memory_refs": [],
            "crawled_resource_ids": crawled_resource_ids,
            "crawl_results": crawl_results,
            "error": False,
            "error_message": "",
        }
        with _STORE_LOCK:
            conversations = _load_conversations()
            conv = _find_by_id(conversations, conversation_id)
            if conv:
                conv.setdefault("messages", []).append(assistant_message)
                conv["updated_at"] = _now_iso()
                _save_conversations(conversations)
        _append_agent_log(
            "conversation_crawler_done",
            {
                "conversation_id": conversation_id,
                "conversation_title": conversation_title,
                "user_message_id": user_message.get("id"),
                "assistant_message_id": assistant_message.get("id"),
                "agent_name": agent_name,
                "strategy_id": strategy_id,
                "strategy_name": strategy_name,
                "ok_count": sum(1 for item in crawl_results if (item or {}).get("status") == "ok"),
                "total_count": len(crawl_results or []),
            },
        )
        return jsonify(
            {
                "success": True,
                "data": {
                    "conversation_id": conversation_id,
                    "user_message": user_message,
                    "assistant_message": assistant_message,
                    "crawl_results": crawl_results,
                },
                "timestamp": _now_iso(),
            }
        )

    resource_context, used_resource_ids, skipped_resources = _build_resource_context(resource_ids)
    history_for_conversation = conv.get("messages", [])[-MAX_HISTORY_MESSAGES:]

    strategy_context, resolved_strategy_id = _build_strategy_context(strategy_id)
    extra_context = strategy_context
    runtime_clock_context = _build_runtime_clock_context()
    if runtime_clock_context:
        if extra_context:
            extra_context += "\n\n"
        extra_context += runtime_clock_context
    memory_context, memory_meta = _build_memory_profile_context(memory_profile_id)
    resolved_memory_profile_id = (memory_meta.get("memory_profile_id") or "").strip()
    used_memory_resource_ids = memory_meta.get("used_memory_resource_ids") or []
    skipped_memory_refs = memory_meta.get("skipped_memory_refs") or []
    if memory_context:
        if extra_context:
            extra_context += "\n\n"
        extra_context += "长期记忆上下文：\n" + memory_context
    if conversation_mode in {"dialog", "strategy_edit", "strategy_analysis"}:
        dialog_cache_context = _build_dialog_data_cache_context(llm_user_content)
        if dialog_cache_context:
            if extra_context:
                extra_context += "\n\n"
            extra_context += "数据缓存上下文（data-cache-inspector 自动执行）：\n" + dialog_cache_context

    assistant_text = ""
    error_text = ""
    orchestration_meta: Dict[str, Any] = {}
    resolved_model = _canonical_model_name(model_name or (os.getenv("OPENAI_MODEL") or "gpt-4o-mini"))
    provider = ""
    usage = {}
    trace_context = {
        "conversation_id": conversation_id,
        "conversation_title": conversation_title,
        "user_message_id": user_message.get("id"),
        "agent_name": agent_name,
        "conversation_mode": conversation_mode,
        "strategy_id": resolved_strategy_id or strategy_id,
        "strategy_name": strategy_name or _resolve_strategy_name(resolved_strategy_id),
        "entrypoint": "send_strategy_message",
    }
    try:
        if conversation_mode == "strategy_edit":
            (
                assistant_text,
                resolved_model,
                provider,
                usage,
                error_text,
                orchestration_meta,
            ) = _run_strategy_edit_orchestration(
                llm_user_content=llm_user_content,
                history_messages=history_for_conversation,
                resource_context=resource_context,
                extra_context=extra_context,
                model_name=model_name,
                temperature=temperature,
                trace_context=trace_context,
            )
        else:
            assistant_text, resolved_model, provider, usage = chat_with_agent(
                agent_name=agent_name,
                user_content=llm_user_content,
                history_messages=history_for_conversation,
                resource_context=resource_context,
                extra_context=extra_context,
                model_name=model_name,
                temperature=temperature,
                trace_context=trace_context,
            )
    except Exception as exc:
        error_text = str(exc)
        assistant_text = (
            f"LLM 调用失败：{error_text}\n"
            "请检查 OPENAI_API_KEY / OPENAI_BASE_URL / OPENAI_MODEL 配置。"
        )

    assistant_message = {
        "id": _gen_id("msg"),
        "role": "assistant",
        "content": assistant_text,
        "created_at": _now_iso(),
        "model": resolved_model,
        "provider": provider,
        "usage": usage or {},
        "agent_name": agent_name,
        "conversation_mode": conversation_mode,
        "strategy_id": resolved_strategy_id,
        "memory_profile_id": resolved_memory_profile_id,
        "used_resource_ids": used_resource_ids,
        "skipped_resource_refs": skipped_resources,
        "used_memory_resource_ids": used_memory_resource_ids,
        "skipped_memory_refs": skipped_memory_refs,
        "crawled_resource_ids": crawled_resource_ids,
        "crawl_results": crawl_results,
        "error": bool(error_text),
        "error_message": error_text,
        "orchestration_meta": orchestration_meta if conversation_mode == "strategy_edit" else {},
    }

    with _STORE_LOCK:
        conversations = _load_conversations()
        conv = _find_by_id(conversations, conversation_id)
        if conv:
            conv.setdefault("messages", []).append(assistant_message)
            conv["updated_at"] = _now_iso()
            _save_conversations(conversations)
    _append_agent_log(
        "conversation_agent_done",
        {
            "conversation_id": conversation_id,
            "conversation_title": conversation_title,
            "user_message_id": user_message.get("id"),
            "assistant_message_id": assistant_message.get("id"),
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": resolved_strategy_id or strategy_id,
            "strategy_name": strategy_name or _resolve_strategy_name(resolved_strategy_id),
            "provider": provider,
            "model": resolved_model,
            "error": bool(error_text),
            "error_message": error_text,
        },
    )

    return jsonify(
        {
            "success": True,
            "data": {
                "conversation_id": conversation_id,
                "user_message": user_message,
                "assistant_message": assistant_message,
                "crawl_results": crawl_results,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/conversations/<string:conversation_id>/messages/stream", methods=["POST"])
def send_strategy_message_stream(conversation_id: str):
    payload = request.get_json(silent=True) or {}
    content = (payload.get("content") or "").strip()
    if not content:
        return jsonify({"success": False, "error": "消息内容不能为空。", "timestamp": _now_iso()}), 400

    resource_ids = payload.get("resource_ids") or []
    if not isinstance(resource_ids, list):
        resource_ids = []
    strategy_id = (payload.get("strategy_id") or "").strip()
    memory_profile_id = (payload.get("memory_profile_id") or "").strip()
    prompt_template_id = (payload.get("prompt_template_id") or "").strip()
    prompt_template = (payload.get("prompt_template") or "").strip()
    llm_user_content = _compose_user_content_with_prompt_template(content, prompt_template)

    agent_name, conversation_mode, run_crawler_direct = _resolve_conversation_dispatch(payload)

    temperature = payload.get("temperature", 0.2)
    model_name = (payload.get("model") or "").strip()
    crawler_cookie = (
        (payload.get("crawler_cookie") or "")
        or (request.headers.get("X-Wechat-Cookie") or "")
    ).strip()

    with _STORE_LOCK:
        conversations = _load_conversations()
        conv = _find_by_id(conversations, conversation_id)
        if not conv:
            return jsonify({"success": False, "error": "会话不存在。", "timestamp": _now_iso()}), 404

        user_message = {
            "id": _gen_id("msg"),
            "role": "user",
            "content": content,
            "created_at": _now_iso(),
            "resource_ids": resource_ids,
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": strategy_id,
            "memory_profile_id": memory_profile_id,
            "prompt_template_id": prompt_template_id,
            "prompt_template": prompt_template,
        }
        conv.setdefault("messages", []).append(user_message)
        conv["updated_at"] = _now_iso()
        _save_conversations(conversations)
    conversation_title = (conv.get("title") or "").strip()
    strategy_name = _resolve_strategy_name(strategy_id)
    _append_agent_log(
        "conversation_message_received",
        {
            "conversation_id": conversation_id,
            "conversation_title": conversation_title,
            "user_message_id": user_message.get("id"),
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": strategy_id,
            "strategy_name": strategy_name,
            "model": model_name or "",
            "stream": True,
        },
    )

    crawled_resource_ids: List[str] = []
    crawl_results: List[Dict] = []
    if run_crawler_direct:
        crawl_results = crawl_wechat_articles_from_text(
            content,
            output_dir=CRAWLED_DIR,
            cookie_override=crawler_cookie,
        )
        crawl_results = _attach_crawl_relpaths(crawl_results)
        ok_items = [x for x in crawl_results if x.get("status") == "ok"]
        if ok_items:
            with _STORE_LOCK:
                crawled_resource_ids = _register_crawled_resources(ok_items)

        assistant_message_id = _gen_id("msg")
        assistant_created_at = _now_iso()
        assistant_text = _build_crawler_agent_summary(crawl_results)
        assistant_message = {
            "id": assistant_message_id,
            "role": "assistant",
            "content": assistant_text,
            "created_at": assistant_created_at,
            "model": "crawler_direct",
            "provider": "crawler",
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": strategy_id,
            "memory_profile_id": memory_profile_id,
            "usage": {},
            "used_resource_ids": [],
            "skipped_resource_refs": [],
            "used_memory_resource_ids": [],
            "skipped_memory_refs": [],
            "crawled_resource_ids": crawled_resource_ids,
            "crawl_results": crawl_results,
            "error": False,
            "error_message": "",
        }

        def _crawler_sse(data: Dict) -> str:
            return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"

        def generate_crawler():
            yield _crawler_sse(
                {
                    "type": "meta",
                    "conversation_id": conversation_id,
                    "user_message": user_message,
                    "assistant_message": {
                        "id": assistant_message_id,
                        "role": "assistant",
                        "content": "",
                        "created_at": assistant_created_at,
                        "model": "crawler_direct",
                        "provider": "crawler",
                        "agent_name": agent_name,
                        "conversation_mode": conversation_mode,
                        "strategy_id": strategy_id,
                        "memory_profile_id": memory_profile_id,
                        "used_resource_ids": [],
                        "skipped_resource_refs": [],
                        "used_memory_resource_ids": [],
                        "skipped_memory_refs": [],
                        "crawled_resource_ids": crawled_resource_ids,
                        "crawl_results": crawl_results,
                    },
                }
            )
            if assistant_text:
                yield _crawler_sse({"type": "delta", "text": assistant_text})
            with _STORE_LOCK:
                conversations = _load_conversations()
                conv = _find_by_id(conversations, conversation_id)
                if conv:
                    conv.setdefault("messages", []).append(assistant_message)
                    conv["updated_at"] = _now_iso()
                    _save_conversations(conversations)
            _append_agent_log(
                "conversation_crawler_done",
                {
                    "conversation_id": conversation_id,
                    "conversation_title": conversation_title,
                    "user_message_id": user_message.get("id"),
                    "assistant_message_id": assistant_message_id,
                    "agent_name": agent_name,
                    "strategy_id": strategy_id,
                    "strategy_name": strategy_name,
                    "ok_count": sum(1 for item in crawl_results if (item or {}).get("status") == "ok"),
                    "total_count": len(crawl_results or []),
                    "stream": True,
                },
            )
            yield _crawler_sse({"type": "done"})

        return Response(
            stream_with_context(generate_crawler()),
            mimetype="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
            },
        )

    resource_context, used_resource_ids, skipped_resources = _build_resource_context(resource_ids)
    history_for_conversation = conv.get("messages", [])[-MAX_HISTORY_MESSAGES:]

    strategy_context, resolved_strategy_id = _build_strategy_context(strategy_id)
    extra_context = strategy_context
    runtime_clock_context = _build_runtime_clock_context()
    if runtime_clock_context:
        if extra_context:
            extra_context += "\n\n"
        extra_context += runtime_clock_context
    memory_context, memory_meta = _build_memory_profile_context(memory_profile_id)
    resolved_memory_profile_id = (memory_meta.get("memory_profile_id") or "").strip()
    used_memory_resource_ids = memory_meta.get("used_memory_resource_ids") or []
    skipped_memory_refs = memory_meta.get("skipped_memory_refs") or []
    if memory_context:
        if extra_context:
            extra_context += "\n\n"
        extra_context += "长期记忆上下文：\n" + memory_context
    if conversation_mode in {"dialog", "strategy_edit", "strategy_analysis"}:
        dialog_cache_context = _build_dialog_data_cache_context(llm_user_content)
        if dialog_cache_context:
            if extra_context:
                extra_context += "\n\n"
            extra_context += "数据缓存上下文（data-cache-inspector 自动执行）：\n" + dialog_cache_context

    assistant_message_id = _gen_id("msg")
    assistant_created_at = _now_iso()

    def _sse(data: Dict) -> str:
        return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"

    def _save_assistant_message(assistant_message: Dict) -> None:
        with _STORE_LOCK:
            conversations = _load_conversations()
            conv = _find_by_id(conversations, conversation_id)
            if conv:
                conv.setdefault("messages", []).append(assistant_message)
                conv["updated_at"] = _now_iso()
                _save_conversations(conversations)

    def generate():
        def _friendly_llm_error(raw_error: str) -> str:
            text = (raw_error or "").strip() or "unknown error"
            if "Max context tokens exceeded" in text or "-10023" in text:
                return (
                    "LLM 调用失败：上下文过长，已超过模型限制。"
                    "请减少资料数量、清理历史对话，或切换到更大上下文模型。"
                )
            if "Parameter error" in text or "-10003" in text:
                return (
                    "LLM 调用失败：请求参数不被当前模型/网关接受。"
                    "请检查模型名（例如 deepseek-v3.2 / qwen-max）和流式参数兼容性。"
                )
            return f"LLM 调用失败：{text}"

        error_text = ""
        resolved_model = _canonical_model_name(model_name or (os.getenv("OPENAI_MODEL") or "gpt-4o-mini"))
        provider = ""
        full_text = ""
        usage = {}

        usage_collector = {}
        trace_context = {
            "conversation_id": conversation_id,
            "conversation_title": conversation_title,
            "user_message_id": user_message.get("id"),
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": resolved_strategy_id or strategy_id,
            "strategy_name": strategy_name or _resolve_strategy_name(resolved_strategy_id),
            "entrypoint": "send_strategy_message_stream",
            "stream": True,
        }

        if conversation_mode == "strategy_edit":
            assistant_text = ""
            orchestration_meta: Dict[str, Any] = {}
            try:
                (
                    assistant_text,
                    resolved_model,
                    provider,
                    usage,
                    error_text,
                    orchestration_meta,
                ) = _run_strategy_edit_orchestration(
                    llm_user_content=llm_user_content,
                    history_messages=history_for_conversation,
                    resource_context=resource_context,
                    extra_context=extra_context,
                    model_name=model_name,
                    temperature=temperature,
                    trace_context=trace_context,
                )
            except Exception as exc:
                error_text = str(exc)
                assistant_text = (
                    _friendly_llm_error(error_text)
                    + "\n请检查 OPENAI_API_KEY / OPENAI_BASE_URL / OPENAI_MODEL 配置。"
                )

            meta = {
                "type": "meta",
                "conversation_id": conversation_id,
                "user_message": user_message,
                "assistant_message": {
                    "id": assistant_message_id,
                    "role": "assistant",
                    "content": "",
                    "created_at": assistant_created_at,
                    "model": resolved_model,
                    "provider": provider,
                    "agent_name": agent_name,
                    "conversation_mode": conversation_mode,
                    "strategy_id": resolved_strategy_id,
                    "memory_profile_id": resolved_memory_profile_id,
                    "used_resource_ids": used_resource_ids,
                    "skipped_resource_refs": skipped_resources,
                    "used_memory_resource_ids": used_memory_resource_ids,
                    "skipped_memory_refs": skipped_memory_refs,
                    "crawled_resource_ids": crawled_resource_ids,
                    "crawl_results": crawl_results,
                },
            }
            yield _sse(meta)

            if assistant_text:
                chunk_size = 240
                for i in range(0, len(assistant_text), chunk_size):
                    yield _sse({"type": "delta", "text": assistant_text[i : i + chunk_size]})

            assistant_message = {
                "id": assistant_message_id,
                "role": "assistant",
                "content": assistant_text,
                "created_at": assistant_created_at,
                "model": resolved_model,
                "provider": provider,
                "agent_name": agent_name,
                "conversation_mode": conversation_mode,
                "strategy_id": resolved_strategy_id,
                "memory_profile_id": resolved_memory_profile_id,
                "usage": usage or {},
                "used_resource_ids": used_resource_ids,
                "skipped_resource_refs": skipped_resources,
                "used_memory_resource_ids": used_memory_resource_ids,
                "skipped_memory_refs": skipped_memory_refs,
                "crawled_resource_ids": crawled_resource_ids,
                "crawl_results": crawl_results,
                "error": bool(error_text),
                "error_message": error_text,
                "orchestration_meta": orchestration_meta,
            }
            _save_assistant_message(assistant_message)
            _append_agent_log(
                "conversation_agent_done",
                {
                    "conversation_id": conversation_id,
                    "conversation_title": conversation_title,
                    "user_message_id": user_message.get("id"),
                    "assistant_message_id": assistant_message_id,
                    "agent_name": agent_name,
                    "conversation_mode": conversation_mode,
                    "strategy_id": resolved_strategy_id or strategy_id,
                    "strategy_name": strategy_name or _resolve_strategy_name(resolved_strategy_id),
                    "provider": provider,
                    "model": resolved_model,
                    "stream": True,
                    "error": bool(error_text),
                    "error_message": error_text,
                },
            )
            yield _sse({"type": "done"})
            return

        try:
            stream, resolved_model, provider, usage_collector = chat_with_agent_stream(
                agent_name=agent_name,
                user_content=llm_user_content,
                history_messages=history_for_conversation,
                resource_context=resource_context,
                extra_context=extra_context,
                model_name=model_name,
                temperature=temperature,
                trace_context=trace_context,
            )
        except Exception as exc:
            error_text = str(exc)
            assistant_text = (
                _friendly_llm_error(error_text)
                + "\n请检查 OPENAI_API_KEY / OPENAI_BASE_URL / OPENAI_MODEL 配置。"
            )
            assistant_message = {
                "id": assistant_message_id,
                "role": "assistant",
                "content": assistant_text,
                "created_at": assistant_created_at,
                "model": resolved_model,
                "provider": provider,
                "agent_name": agent_name,
                "conversation_mode": conversation_mode,
                "strategy_id": resolved_strategy_id,
                "memory_profile_id": resolved_memory_profile_id,
                "usage": usage or {},
                "used_resource_ids": used_resource_ids,
                "skipped_resource_refs": skipped_resources,
                "used_memory_resource_ids": used_memory_resource_ids,
                "skipped_memory_refs": skipped_memory_refs,
                "crawled_resource_ids": crawled_resource_ids,
                "crawl_results": crawl_results,
                "error": True,
                "error_message": error_text,
            }
            _save_assistant_message(assistant_message)
            _append_agent_log(
                "conversation_agent_done",
                {
                    "conversation_id": conversation_id,
                    "conversation_title": conversation_title,
                    "user_message_id": user_message.get("id"),
                    "assistant_message_id": assistant_message_id,
                    "agent_name": agent_name,
                    "conversation_mode": conversation_mode,
                    "strategy_id": resolved_strategy_id or strategy_id,
                    "strategy_name": strategy_name or _resolve_strategy_name(resolved_strategy_id),
                    "provider": provider,
                    "model": resolved_model,
                    "stream": True,
                    "error": True,
                    "error_message": error_text,
                },
            )
            yield _sse({"type": "error", "message": assistant_text})
            yield _sse({"type": "done"})
            return

        meta = {
            "type": "meta",
            "conversation_id": conversation_id,
            "user_message": user_message,
            "assistant_message": {
                "id": assistant_message_id,
                "role": "assistant",
                "content": "",
                "created_at": assistant_created_at,
                "model": resolved_model,
                "provider": provider,
                "agent_name": agent_name,
                "conversation_mode": conversation_mode,
                "strategy_id": resolved_strategy_id,
                "memory_profile_id": resolved_memory_profile_id,
                "used_resource_ids": used_resource_ids,
                "skipped_resource_refs": skipped_resources,
                "used_memory_resource_ids": used_memory_resource_ids,
                "skipped_memory_refs": skipped_memory_refs,
                "crawled_resource_ids": crawled_resource_ids,
                "crawl_results": crawl_results,
            },
        }
        yield _sse(meta)

        try:
            for chunk in stream:
                if not chunk:
                    continue
                full_text += chunk
                yield _sse({"type": "delta", "text": chunk})
        except Exception as exc:
            error_text = str(exc)
            assistant_text = _friendly_llm_error(error_text)
            yield _sse({"type": "error", "message": assistant_text})
            if full_text:
                full_text += f"\n\n[{assistant_text}]"
            else:
                full_text = assistant_text

        if isinstance(usage_collector, dict):
            usage = usage_collector.get("usage") or usage

        assistant_message = {
            "id": assistant_message_id,
            "role": "assistant",
            "content": full_text,
            "created_at": assistant_created_at,
            "model": resolved_model,
            "provider": provider,
            "agent_name": agent_name,
            "conversation_mode": conversation_mode,
            "strategy_id": resolved_strategy_id,
            "memory_profile_id": resolved_memory_profile_id,
            "usage": usage or {},
            "used_resource_ids": used_resource_ids,
            "skipped_resource_refs": skipped_resources,
            "used_memory_resource_ids": used_memory_resource_ids,
            "skipped_memory_refs": skipped_memory_refs,
            "crawled_resource_ids": crawled_resource_ids,
            "crawl_results": crawl_results,
            "error": bool(error_text),
            "error_message": error_text,
        }
        _save_assistant_message(assistant_message)
        _append_agent_log(
            "conversation_agent_done",
            {
                "conversation_id": conversation_id,
                "conversation_title": conversation_title,
                "user_message_id": user_message.get("id"),
                "assistant_message_id": assistant_message_id,
                "agent_name": agent_name,
                "conversation_mode": conversation_mode,
                "strategy_id": resolved_strategy_id or strategy_id,
                "strategy_name": strategy_name or _resolve_strategy_name(resolved_strategy_id),
                "provider": provider,
                "model": resolved_model,
                "stream": True,
                "error": bool(error_text),
                "error_message": error_text,
            },
        )
        yield _sse({"type": "done"})

    return Response(
        stream_with_context(generate()),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@strategy_watch_bp.route("/api/strategy-watch/strategies", methods=["GET"])
def list_strategy_watch_strategies():
    with _STORE_LOCK:
        payload = _load_strategies_payload()
        strategies = payload.get("strategies", [])
        active_strategy_id = payload.get("active_strategy_id") or ""
    strategies = sorted(strategies, key=lambda x: x.get("updated_at", ""), reverse=True)
    return jsonify(
        {
            "success": True,
            "data": {
                "strategies": strategies,
                "active_strategy_id": active_strategy_id,
            },
            "timestamp": _now_iso(),
        }
    )


@strategy_watch_bp.route("/api/strategy-watch/strategies", methods=["POST"])
def create_strategy_watch_strategy():
    payload = request.get_json(silent=True) or {}
    name = (payload.get("name") or "").strip() or f"看盘策略 {datetime.now().strftime('%m-%d %H:%M')}"
    description = (payload.get("description") or "").strip()
    view_type = (payload.get("view_type") or DEFAULT_STRATEGY_VIEW).strip() or DEFAULT_STRATEGY_VIEW
    config = payload.get("config") or {}

    strategy = {
        "id": _gen_id("strategy"),
        "name": name,
        "description": description,
        "view_type": view_type,
        "config": config,
        "created_at": _now_iso(),
        "updated_at": _now_iso(),
    }

    with _STORE_LOCK:
        payload = _load_strategies_payload()
        strategies = payload.get("strategies", [])
        strategies.append(strategy)
        payload["strategies"] = strategies
        if not payload.get("active_strategy_id"):
            payload["active_strategy_id"] = strategy["id"]
        _save_strategies_payload(payload)

    return jsonify({"success": True, "data": strategy, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/strategies/<string:strategy_id>", methods=["PATCH"])
def update_strategy_watch_strategy(strategy_id: str):
    payload = request.get_json(silent=True) or {}
    with _STORE_LOCK:
        data = _load_strategies_payload()
        strategies = data.get("strategies", [])
        target = _find_by_id(strategies, strategy_id)
        if not target:
            return jsonify({"success": False, "error": "策略不存在。", "timestamp": _now_iso()}), 404

        if "name" in payload:
            target["name"] = (payload.get("name") or "").strip() or target.get("name")
        if "description" in payload:
            target["description"] = (payload.get("description") or "").strip()
        if "view_type" in payload:
            target["view_type"] = (payload.get("view_type") or DEFAULT_STRATEGY_VIEW).strip() or DEFAULT_STRATEGY_VIEW
        if "config" in payload and isinstance(payload.get("config"), dict):
            target["config"] = payload.get("config")
        target["updated_at"] = _now_iso()

        data["strategies"] = strategies
        _save_strategies_payload(data)

    return jsonify({"success": True, "data": target, "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/strategies/<string:strategy_id>", methods=["DELETE"])
def delete_strategy_watch_strategy(strategy_id: str):
    with _STORE_LOCK:
        data = _load_strategies_payload()
        strategies = data.get("strategies", [])
        target = _find_by_id(strategies, strategy_id)
        if not target:
            return jsonify({"success": False, "error": "策略不存在。", "timestamp": _now_iso()}), 404
        strategies = [item for item in strategies if item.get("id") != strategy_id]
        data["strategies"] = strategies
        if data.get("active_strategy_id") == strategy_id:
            data["active_strategy_id"] = ""
        _save_strategies_payload(data)

    return jsonify({"success": True, "message": "策略已删除。", "timestamp": _now_iso()})


@strategy_watch_bp.route("/api/strategy-watch/strategies/active", methods=["PATCH"])
def set_active_strategy_watch_strategy():
    payload = request.get_json(silent=True) or {}
    strategy_id = (payload.get("strategy_id") or "").strip()

    with _STORE_LOCK:
        data = _load_strategies_payload()
        strategies = data.get("strategies", [])
        if strategy_id and not _find_by_id(strategies, strategy_id):
            return jsonify({"success": False, "error": "策略不存在。", "timestamp": _now_iso()}), 404
        data["active_strategy_id"] = strategy_id
        _save_strategies_payload(data)

    return jsonify(
        {"success": True, "data": {"active_strategy_id": strategy_id}, "timestamp": _now_iso()}
    )


@strategy_watch_bp.route("/api/strategy-watch/strategies/<string:strategy_id>/generate-view", methods=["POST"])
def generate_strategy_watch_view(strategy_id: str):
    payload = request.get_json(silent=True) or {}
    goal = (payload.get("goal") or "").strip()
    if not goal:
        return jsonify({"success": False, "error": "目标描述不能为空。", "timestamp": _now_iso()}), 400

    with _STORE_LOCK:
        data = _load_strategies_payload()
        strategies = data.get("strategies", [])
        target = _find_by_id(strategies, strategy_id)
        if not target:
            return jsonify({"success": False, "error": "策略不存在。", "timestamp": _now_iso()}), 404
        target_config = target.get("config") if isinstance(target.get("config"), dict) else {}
        existing_widgets = target_config.get("widgets") if isinstance(target_config.get("widgets"), list) else []

    prompt = _build_visualization_prompt(
        goal,
        view_type=target.get("view_type", ""),
        existing_widgets=existing_widgets,
    )
    try:
        reply, resolved_model, provider, _usage = chat_with_agent(
            agent_name="visualization_master_agent",
            user_content=prompt,
            history_messages=[],
            resource_context="",
            extra_context="",
            model_name="",
            temperature=0.2,
            trace_context={
                "entrypoint": "generate_strategy_watch_view",
                "strategy_id": strategy_id,
                "agent_name": "visualization_master_agent",
            },
        )
    except Exception as exc:
        return jsonify(
            {"success": False, "error": f"生成视图失败: {exc}", "timestamp": _now_iso()}), 500

    view_config = _extract_json_block(reply)
    if not view_config:
        return jsonify(
            {"success": False, "error": "未能解析Agent输出的JSON配置。", "timestamp": _now_iso()}), 500

    with _STORE_LOCK:
        data = _load_strategies_payload()
        strategies = data.get("strategies", [])
        target = _find_by_id(strategies, strategy_id)
        if not target:
            return jsonify({"success": False, "error": "策略不存在。", "timestamp": _now_iso()}), 404

        target["name"] = (view_config.get("name") or target.get("name") or "").strip() or target.get("name")
        target["description"] = (view_config.get("description") or target.get("description") or "").strip()
        target["view_type"] = (
            (view_config.get("view_type") or target.get("view_type") or DEFAULT_STRATEGY_VIEW).strip()
        )
        widgets = view_config.get("widgets")
        if isinstance(widgets, list):
            current_config = target.get("config") if isinstance(target.get("config"), dict) else {}
            current_widgets = current_config.get("widgets") if isinstance(current_config.get("widgets"), list) else []
            merged_widgets = _merge_strategy_widgets(current_widgets, widgets)
            next_config = dict(current_config)
            next_config["widgets"] = merged_widgets
            target["config"] = next_config
        target["updated_at"] = _now_iso()
        data["strategies"] = strategies
        _save_strategies_payload(data)

    return jsonify(
        {
            "success": True,
            "data": {
                "strategy_id": strategy_id,
                "view_config": view_config,
                "model": resolved_model,
                "provider": provider,
            },
            "timestamp": _now_iso(),
        }
    )

