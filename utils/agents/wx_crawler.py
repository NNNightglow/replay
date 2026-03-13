#!/usr/bin/env python
# -*- coding: utf-8 -*-

import hashlib
import os
import re
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup, NavigableString, Tag

WX_URL_PATTERN = re.compile(r"https?://mp\.weixin\.qq\.com/[^\s)>\]\"']+")
DESKTOP_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
)
MOBILE_WECHAT_UA = (
    "Mozilla/5.0 (iPhone; CPU iPhone OS 17_0 like Mac OS X) "
    "AppleWebKit/605.1.15 (KHTML, like Gecko) Mobile/15E148 "
    "MicroMessenger/8.0.51(0x1800332b) NetType/WIFI Language/zh_CN"
)
DEFAULT_HEADERS = {
    "Accept": (
        "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,"
        "*/*;q=0.8"
    ),
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
    "Upgrade-Insecure-Requests": "1",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "none",
    "User-Agent": (
        DESKTOP_UA
    )
}


def extract_wechat_urls(text: str) -> List[str]:
    if not text:
        return []
    urls = WX_URL_PATTERN.findall(text)
    dedup = []
    seen = set()
    for u in urls:
        cleaned = u.rstrip(".,;!?")
        if cleaned not in seen:
            seen.add(cleaned)
            dedup.append(cleaned)
    return dedup


def _build_headers(url: str, mobile: bool = False, cookie: str = "") -> Dict[str, str]:
    headers = dict(DEFAULT_HEADERS)
    headers["Referer"] = "https://mp.weixin.qq.com/"
    if mobile:
        headers["User-Agent"] = MOBILE_WECHAT_UA
        headers["Sec-Fetch-Site"] = "cross-site"
    if cookie:
        headers["Cookie"] = cookie
    return headers


def _is_wechat_blocked(html: str) -> bool:
    if not html:
        return True
    lowered = html.lower()
    block_markers = (
        "环境异常",
        "访问过于频繁",
        "暂时无法访问",
        "请在微信客户端打开链接",
        "frequency control",
    )
    if any(marker in html for marker in block_markers):
        return True
    if "js_content" in html:
        return False
    # 微信也可能只在脚本中返回正文，这里保守地认为非目标页
    return "mp.weixin.qq.com/s" in lowered and "msg_title" not in lowered


def _safe_datetime_from_unix(raw: str) -> str:
    try:
        ts = int((raw or "").strip())
    except Exception:
        return ""
    if ts <= 0:
        return ""
    try:
        return datetime.fromtimestamp(ts, timezone.utc).isoformat()
    except Exception:
        return ""


def _safe_datetime_from_text(raw: str) -> str:
    value = (raw or "").strip().replace("/", "-")
    if not value:
        return ""
    fmts = [
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
    ]
    for fmt in fmts:
        try:
            dt_obj = datetime.strptime(value, fmt)
            if fmt == "%Y-%m-%d":
                dt_obj = dt_obj.replace(hour=0, minute=0, second=0)
            return dt_obj.replace(tzinfo=timezone.utc).isoformat()
        except Exception:
            continue
    return ""


def _extract_published_at(html: str, soup: BeautifulSoup) -> str:
    meta_candidates = [
        soup.find("meta", attrs={"property": "article:published_time"}),
        soup.find("meta", attrs={"name": "publishdate"}),
        soup.find("meta", attrs={"name": "PublishDate"}),
    ]
    for node in meta_candidates:
        if not node:
            continue
        value = _safe_datetime_from_text(node.get("content", ""))
        if value:
            return value

    patterns = [
        r"(?:var\s+ct|window\.__ct)\s*=\s*\"?(\d{10})\"?",
        r"\"publish_time\"\\s*:\\s*\"([0-9:\\-\\s/]+)\"",
        r"\"ori_create_time\"\\s*:\\s*\"?(\d{10})\"?",
    ]
    for pattern in patterns:
        match = re.search(pattern, html or "")
        if not match:
            continue
        raw = (match.group(1) or "").strip()
        if raw.isdigit() and len(raw) == 10:
            value = _safe_datetime_from_unix(raw)
        else:
            value = _safe_datetime_from_text(raw)
        if value:
            return value

    return ""


def _load_cookie_override(cookie_override: str = "") -> str:
    if cookie_override and cookie_override.strip():
        return cookie_override.strip()
    for key in ("WX_CRAWLER_COOKIE", "WECHAT_COOKIE"):
        val = (os.getenv(key) or "").strip()
        if val:
            return val
    return ""


def _request_with_fallback(url: str, timeout: int, cookie: str = "") -> requests.Response:
    session = requests.Session()
    attempts = [
        {"mobile": False, "cookie": ""},
        {"mobile": True, "cookie": ""},
    ]
    if cookie:
        attempts.append({"mobile": False, "cookie": cookie})
        attempts.append({"mobile": True, "cookie": cookie})

    last_resp = None
    last_error = ""
    for idx, conf in enumerate(attempts, start=1):
        try:
            resp = session.get(
                url,
                headers=_build_headers(url, mobile=conf["mobile"], cookie=conf["cookie"]),
                timeout=timeout,
                allow_redirects=True,
            )
            last_resp = resp
            resp.raise_for_status()
            if not _is_wechat_blocked(resp.text):
                return resp
            last_error = "触发微信访问限制"
        except Exception as exc:
            last_error = str(exc)
        if idx < len(attempts):
            time.sleep(0.7)

    if last_resp is not None:
        raise RuntimeError(
            f"微信页面访问受限（HTTP {last_resp.status_code}）: {last_error}. "
            "可在请求体传 crawler_cookie，或设置环境变量 WX_CRAWLER_COOKIE/WECHAT_COOKIE。"
        )
    raise RuntimeError(
        f"请求失败: {last_error}. "
        "可在请求体传 crawler_cookie，或设置环境变量 WX_CRAWLER_COOKIE/WECHAT_COOKIE。"
    )


def fetch_wechat_article(url: str, timeout: int = 30, cookie_override: str = "") -> Dict:
    cookie = _load_cookie_override(cookie_override)
    resp = _request_with_fallback(url=url, timeout=timeout, cookie=cookie)

    html = resp.text
    soup = BeautifulSoup(html, "html.parser")

    title = ""
    title_node = soup.select_one("#activity-name")
    if title_node:
        title = title_node.get_text(" ", strip=True)
    if not title:
        meta_title = soup.find("meta", attrs={"property": "og:title"})
        if meta_title and meta_title.get("content"):
            title = meta_title.get("content", "").strip()
    if not title:
        title = "微信公众号文章"

    content_node = soup.select_one("#js_content")
    if content_node:
        content_text = content_node.get_text("\n", strip=True)
        content_html = str(content_node)
    else:
        content_text = soup.get_text("\n", strip=True)
        content_html = ""
    published_at = _extract_published_at(html, soup)

    return {
        "url": url,
        "title": title,
        "content": content_text,
        "content_html": content_html,
        "fetched_at_utc": datetime.now(timezone.utc).isoformat(),
        "published_at": published_at,
    }


def _guess_image_ext(image_url: str, content_type: str = "") -> str:
    if content_type:
        if "jpeg" in content_type or "jpg" in content_type:
            return ".jpg"
        if "png" in content_type:
            return ".png"
        if "gif" in content_type:
            return ".gif"
        if "webp" in content_type:
            return ".webp"
    try:
        path = urlparse(image_url).path
        ext = Path(path).suffix.lower()
        if ext in {".jpg", ".jpeg", ".png", ".gif", ".webp"}:
            return ext
    except Exception:
        pass
    return ".jpg"


def _download_images(image_urls: List[str], images_dir: Path, timeout: int = 30) -> Dict[str, Path]:
    images_dir.mkdir(parents=True, exist_ok=True)
    saved_paths: Dict[str, Path] = {}
    for idx, image_url in enumerate(image_urls, start=1):
        if not image_url:
            continue
        try:
            resp = requests.get(image_url, headers=DEFAULT_HEADERS, timeout=timeout)
            resp.raise_for_status()
            ext = _guess_image_ext(image_url, resp.headers.get("Content-Type", ""))
            file_name = f"img_{idx:02d}{ext}"
            target = images_dir / file_name
            target.write_bytes(resp.content)
            saved_paths[image_url] = target
        except Exception:
            continue
    return saved_paths


def _extract_content_blocks(content_html: str, fallback_text: str = "") -> List[Dict]:
    if not content_html:
        text = (fallback_text or "").strip()
        return [{"type": "text", "text": text}] if text else []

    soup = BeautifulSoup(content_html, "html.parser")
    blocks: List[Dict] = []
    text_buf: List[str] = []
    block_tags = {
        "p",
        "section",
        "div",
        "li",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "blockquote",
    }

    def flush_text():
        if not text_buf:
            return
        text = "".join(text_buf)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        text_buf.clear()
        if text:
            blocks.append({"type": "text", "text": text})

    def walk(node):
        if isinstance(node, NavigableString):
            txt = str(node)
            if txt.strip():
                text_buf.append(txt)
            return
        if not isinstance(node, Tag):
            return
        if node.name == "img":
            flush_text()
            src = (node.get("data-src") or node.get("data-original") or node.get("src") or "").strip()
            if src:
                blocks.append({"type": "image", "url": src})
            return
        if node.name == "br":
            text_buf.append("\n")
            return

        is_block = node.name in block_tags
        if is_block and text_buf:
            text_buf.append("\n")
        for child in node.children:
            walk(child)
        if is_block:
            text_buf.append("\n")

    root = soup.body or soup
    for child in root.children:
        walk(child)
    flush_text()
    return blocks


def save_article_as_markdown(article: Dict, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    url = article.get("url", "")
    title = article.get("title", "微信公众号文章")
    content = article.get("content", "")
    fetched_at = article.get("fetched_at_utc", "")
    published_at = article.get("published_at", "")
    content_html = article.get("content_html", "")

    digest = hashlib.md5(url.encode("utf-8")).hexdigest()[:12]
    safe_name = re.sub(r"[\\/:*?\"<>|]+", "_", title).strip("_")[:80] or "wx_article"
    target = output_dir / f"{safe_name}_{digest}.md"

    blocks = _extract_content_blocks(content_html, fallback_text=content)
    image_urls = [b.get("url") for b in blocks if b.get("type") == "image"]
    image_map: Dict[str, Path] = {}
    if image_urls:
        images_dir = output_dir / "wx_images" / f"{safe_name}_{digest}"
        image_map = _download_images(image_urls, images_dir)

    body_lines: List[str] = []
    for block in blocks:
        if block.get("type") == "text":
            text = (block.get("text") or "").strip()
            if text:
                body_lines.append(text)
                body_lines.append("")
        elif block.get("type") == "image":
            img_path = image_map.get(block.get("url", ""))
            if img_path:
                rel_path = img_path.relative_to(output_dir).as_posix()
                body_lines.append(f"![image]({rel_path})")
                body_lines.append("")

    markdown = "\n".join(
        [
            "---",
            f'title: "{title}"',
            f'source_url: "{url}"',
            f"fetched_at_utc: {fetched_at}",
            f"published_at: {published_at}",
            f"content_time: {published_at}",
            "content_time_type: published",
            f'content_time_evidence: "{("wechat:published_at" if published_at else "")}"',
            "status: ok",
            "---",
            "",
            f"# {title}",
            "",
            *body_lines,
        ]
    )
    target.write_text(markdown, encoding="utf-8")
    return target


def _resolve_pdf_font() -> Tuple[str, Optional[Path]]:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/msyh.ttf"),
        Path("C:/Windows/Fonts/msyhbd.ttc"),
        Path("C:/Windows/Fonts/msyhbd.ttf"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return "WXFont", path
    return "Helvetica", None


def _write_simple_pdf(text: str, target: Path) -> None:
    # Minimal PDF writer without external deps (Helvetica only, single page).
    def _escape(s: str) -> str:
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    lines = [ln.rstrip() for ln in (text or "").splitlines()]
    if not lines:
        lines = [""]

    content_lines = [
        "BT",
        "/F1 12 Tf",
        "72 720 Td",
        "14 TL",
    ]
    for line in lines:
        content_lines.append(f"({_escape(line)}) Tj")
        content_lines.append("T*")
    content_lines.append("ET")
    content_stream = "\n".join(content_lines).encode("utf-8")

    objects = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    objects.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
    objects.append(b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
                   b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    objects.append(b"<< /Length " + str(len(content_stream)).encode("utf-8") + b" >>\nstream\n" +
                   content_stream + b"\nendstream")

    xref = [b"0000000000 65535 f "]
    output = bytearray()
    output.extend(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{i} 0 obj\n".encode("utf-8"))
        output.extend(obj)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects)+1}\n".encode("utf-8"))
    output.extend(xref[0] + b"\n")
    for off in offsets:
        output.extend(f"{off:010d} 00000 n \n".encode("utf-8"))
    output.extend(b"trailer\n")
    output.extend(f"<< /Size {len(objects)+1} /Root 1 0 R >>\n".encode("utf-8"))
    output.extend(b"startxref\n")
    output.extend(f"{xref_offset}\n".encode("utf-8"))
    output.extend(b"%%EOF")

    target.write_bytes(output)


def _resolve_output_target(base_target: Path) -> Path:
    if not base_target.exists():
        return base_target
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%S")
    return base_target.with_name(f"{base_target.stem}_{stamp}{base_target.suffix}")


def save_article_as_pdf(article: Dict, output_dir: Path) -> Path:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Image as RLImage
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception:
        # Fallback: minimal PDF without external deps
        output_dir.mkdir(parents=True, exist_ok=True)
        url = article.get("url", "")
        title = article.get("title", "微信公众号文章")
        content = article.get("content", "")
        fetched_at = article.get("fetched_at_utc", "")
        digest = hashlib.md5(url.encode("utf-8")).hexdigest()[:12]
        safe_name = re.sub(r"[\\/:*?\"<>|]+", "_", title).strip("_")[:80] or "wx_article"
        target = _resolve_output_target(output_dir / f"{safe_name}_{digest}.pdf")
        plain = f"{title}\n来源: {url}\n抓取时间(UTC): {fetched_at}\n\n{content}"
        _write_simple_pdf(plain, target)
        return target

    output_dir.mkdir(parents=True, exist_ok=True)
    url = article.get("url", "")
    title = article.get("title", "微信公众号文章")
    content = article.get("content", "")
    fetched_at = article.get("fetched_at_utc", "")
    content_html = article.get("content_html", "")

    digest = hashlib.md5(url.encode("utf-8")).hexdigest()[:12]
    safe_name = re.sub(r"[\\/:*?\"<>|]+", "_", title).strip("_")[:80] or "wx_article"
    target = _resolve_output_target(output_dir / f"{safe_name}_{digest}.pdf")

    font_name, font_path = _resolve_pdf_font()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
        except Exception:
            font_name = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "wx_title",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        leading=20,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "wx_meta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        leading=12,
        textColor="#555555",
        spaceAfter=10,
    )
    body_style = ParagraphStyle(
        "wx_body",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=16,
    )

    story = [
        Paragraph(title or "微信公众号文章", title_style),
        Paragraph(f"来源：{url}", meta_style),
        Paragraph(f"抓取时间(UTC)：{fetched_at}", meta_style),
        Spacer(1, 8),
    ]

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
            continue
        safe_text = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe_text, body_style))

    doc = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=36,
        bottomMargin=36,
        title=title,
        author="wx_crawler",
    )
    blocks = _extract_content_blocks(content_html, fallback_text=content)
    image_urls = [b.get("url") for b in blocks if b.get("type") == "image"]
    image_map: Dict[str, Path] = {}
    if image_urls:
        images_dir = output_dir / "wx_images" / f"{safe_name}_{digest}"
        image_map = _download_images(image_urls, images_dir)

    max_width = A4[0] - doc.leftMargin - doc.rightMargin
    for block in blocks:
        if block.get("type") == "text":
            text = (block.get("text") or "").strip()
            if not text:
                continue
            for line in text.split("\n"):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 8))
                    continue
                safe_text = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_text, body_style))
        elif block.get("type") == "image":
            img_path = image_map.get(block.get("url", ""))
            if not img_path:
                continue
            try:
                img = RLImage(str(img_path))
                if img.imageWidth and img.imageHeight:
                    scale = min(1.0, max_width / float(img.imageWidth))
                    img.drawWidth = img.imageWidth * scale
                    img.drawHeight = img.imageHeight * scale
                story.append(Spacer(1, 6))
                story.append(img)
            except Exception:
                continue

    doc.build(story)
    return target


def save_article_as_docx(article: Dict, output_dir: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:
        raise RuntimeError("缺少Word依赖，请先安装 python-docx") from exc

    output_dir.mkdir(parents=True, exist_ok=True)
    url = article.get("url", "")
    title = article.get("title", "微信公众号文章")
    content = article.get("content", "")
    fetched_at = article.get("fetched_at_utc", "")
    content_html = article.get("content_html", "")

    digest = hashlib.md5(url.encode("utf-8")).hexdigest()[:12]
    safe_name = re.sub(r"[\\/:*?\"<>|]+", "_", title).strip("_")[:80] or "wx_article"
    target = _resolve_output_target(output_dir / f"{safe_name}_{digest}.docx")

    doc = Document()
    doc.add_heading(title or "微信公众号文章", level=1)
    doc.add_paragraph(f"来源：{url}")
    doc.add_paragraph(f"抓取时间(UTC)：{fetched_at}")
    doc.add_paragraph("")
    blocks = _extract_content_blocks(content_html, fallback_text=content)
    image_urls = [b.get("url") for b in blocks if b.get("type") == "image"]
    image_map: Dict[str, Path] = {}
    if image_urls:
        images_dir = output_dir / "wx_images" / f"{safe_name}_{digest}"
        image_map = _download_images(image_urls, images_dir)

    for block in blocks:
        if block.get("type") == "text":
            text = (block.get("text") or "").strip()
            if not text:
                continue
            for line in text.split("\n"):
                doc.add_paragraph(line)
        elif block.get("type") == "image":
            img_path = image_map.get(block.get("url", ""))
            if not img_path:
                continue
            try:
                doc.add_picture(str(img_path), width=Inches(5.8))
            except Exception:
                continue
    doc.save(str(target))
    return target


def crawl_wechat_articles_from_text(
    text: str,
    output_dir: Path,
    cookie_override: str = "",
) -> List[Dict]:
    urls = extract_wechat_urls(text)
    results: List[Dict] = []
    cookie_override = _load_cookie_override(cookie_override)

    for url in urls:
        try:
            article = fetch_wechat_article(url, cookie_override=cookie_override)
            md_path = save_article_as_markdown(article, output_dir)
            pdf_path = None
            pdf_error = ""
            try:
                pdf_path = save_article_as_pdf(article, output_dir)
            except Exception as exc:
                pdf_error = str(exc)
            docx_path = None
            docx_error = ""
            try:
                docx_path = save_article_as_docx(article, output_dir)
            except Exception as exc:
                docx_error = str(exc)
            results.append(
                {
                    "url": url,
                    "title": article.get("title", ""),
                    "published_at": article.get("published_at", ""),
                    "fetched_at_utc": article.get("fetched_at_utc", ""),
                    "status": "ok",
                    "markdown_path": str(md_path),
                    "pdf_path": str(pdf_path) if pdf_path else None,
                    "pdf_error": pdf_error,
                    "docx_path": str(docx_path) if docx_path else None,
                    "docx_error": docx_error,
                    "error": "",
                }
            )
        except Exception as exc:
            results.append(
                {
                    "url": url,
                    "title": "",
                    "status": "failed",
                    "markdown_path": None,
                    "pdf_path": None,
                    "pdf_error": "",
                    "docx_path": None,
                    "docx_error": "",
                    "error": str(exc),
                }
            )
    return results

